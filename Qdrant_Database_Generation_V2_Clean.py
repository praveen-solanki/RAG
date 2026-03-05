"""
ADVANCED RAG INGESTION SYSTEM V2 (CLEAN)
=========================================
Improvements over V1:
- pypdfium2 for PDF text extraction (proper word spacing, no concatenated table words)
- TOC / boilerplate page detection and skipping
  * Ratio threshold raised 0.30 → 0.50 (need majority of lines to match)
  * Minimum 5 non-empty lines required before ratio check (prevents false positives
    on short section-header pages like "Version ......... 3.0.0")
  * Max-chars guard: pages > 800 chars are never discarded regardless of ratio
  * Skipped page numbers are logged for auditability
- Near-duplicate chunk removal via Jaccard similarity (85% threshold)
- Larger chunk overlap (256 chars instead of 128) to reduce boundary failures
- Embedding robustness: pre-truncate chunks > MAX_EMBED_CHARS before sending to
  Ollama; on HTTP 500, truncate and retry instead of silently dropping the chunk;
  per-file summary log of any remaining skips
- All V1 features preserved: section-aware chunking, BM25, hybrid Qdrant upload,
  OllamaBGEM3Embedder with retry, deduplication by file hash, dimension-mismatch check
"""

import argparse
import json
import numpy as np
import os
import re
import hashlib
import uuid
import time
from typing import Optional, List, Dict, Tuple
from dataclasses import dataclass
from pathlib import Path
import logging

import pypdfium2 as pdfium
import pdfplumber          # kept for PDF table-count metadata only (not text extraction)
import docx
import requests
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import (
    VectorParams,
    Distance,
    Filter,
    FieldCondition,
    MatchValue,
    PointStruct,
    SparseVector,
    SparseVectorParams,
    SparseIndexParams,
)
from rank_bm25 import BM25Okapi
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize

# Download NLTK data if needed
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    try:
        nltk.download('punkt_tab', quiet=True)
    except Exception:
        nltk.download('punkt', quiet=True)

# ================= CONFIG =================

DATA_DIR = r"/home/olj3kor/praveen/Image_dataset_generation/pdfs/standards"
COLLECTION = "test_63pdf"
QDRANT_URL = "http://localhost:7333"

# Embedding options
USE_OLLAMA_BGE_M3 = True
OLLAMA_URL = "http://localhost:11434"
OLLAMA_MODEL = "bge-m3:latest"

# Fallback to SentenceTransformer if Ollama unavailable
FALLBACK_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

# Chunking parameters
CHUNK_SIZE = 512
CHUNK_OVERLAP = 256          # increased from 128 to reduce boundary-split failures
MIN_CHUNK_SIZE = 100
MAX_CHUNK_SIZE = 1024

# Section detection
ENABLE_SECTION_AWARE = True
SECTION_PATTERNS = [
    r'^#{1,6}\s+(.+)$',       # Markdown headers
    r'^([A-Z][^.!?]*):$',     # Title case with colon
    r'^\d+\.\s+([A-Z].+)$',   # Numbered sections
    r'^([A-Z\s]{3,})$',       # All-caps headers (min 3 chars)
]

BM25_OUTPUT = "bm25_index.json"

# Near-duplicate suppression
JACCARD_DEDUP_THRESHOLD = 0.85   # chunks with >85% Jaccard similarity are dropped
JACCARD_WINDOW = 20              # compare each new chunk against last N accepted chunks

# TOC detection
TOC_LINE_RATIO = 0.50            # >50% of non-empty lines look like TOC entries → skip page
TOC_MIN_CONTENT_CHARS = 50       # pages with fewer characters are classified as TOC/boilerplate
TOC_MIN_LINE_COUNT = 5           # don't ratio-classify pages with fewer non-empty lines (section headers)
TOC_MAX_CONTENT_CHARS = 800      # never classify dense pages as TOC regardless of ratio

# Embedding safety
MAX_EMBED_CHARS = 4000           # truncate chunk text to this before sending to Ollama (fits any context window)

# =========================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class DocumentSection:
    """Represents a document section with hierarchy."""
    title: str
    content: str
    level: int
    page_number: Optional[int] = None
    section_type: str = "text"
    section_hierarchy: Optional[List[str]] = None

    def __post_init__(self):
        if self.section_hierarchy is None:
            self.section_hierarchy = [self.title]


@dataclass
class EnrichedChunk:
    """Chunk with rich metadata."""
    text: str
    section_title: str
    section_hierarchy: List[str]
    page_number: Optional[int]
    chunk_type: str
    word_count: int
    sentence_count: int
    start_char: int
    end_char: int


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def jaccard_similarity(text_a: str, text_b: str) -> float:
    """Compute word-level Jaccard similarity between two text strings."""
    words_a = set(text_a.lower().split())
    words_b = set(text_b.lower().split())
    if not words_a or not words_b:
        return 0.0
    return len(words_a & words_b) / len(words_a | words_b)


def file_hash(path: str) -> str:
    """Generate SHA-256 hash of a file using block-wise reading."""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def already_indexed(client: QdrantClient, collection: str, file_hash_value: str) -> bool:
    """Return True if the file hash is already present in the collection."""
    try:
        filt = Filter(
            must=[FieldCondition(key="file_hash", match=MatchValue(value=file_hash_value))]
        )
        points, _ = client.scroll(collection_name=collection, scroll_filter=filt, limit=1)
        return len(points) > 0
    except Exception as e:
        logger.warning(f"Could not check if file is already indexed: {e}")
        return False


# ---------------------------------------------------------------------------
# Embedding models
# ---------------------------------------------------------------------------

class OllamaBGEM3Embedder:
    """BGE-M3 embedder using Ollama with retry logic."""

    def __init__(self, base_url: str = "http://localhost:11434", model: str = "bge-m3"):
        self.base_url = base_url
        self.model = model
        self.dimension = 1024
        self._test_connection()

    def _test_connection(self):
        """Test Ollama connection and verify model availability."""
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=5)
            if response.status_code == 200:
                logger.info(f"✓ Connected to Ollama at {self.base_url}")
                available_models = [m.get("name", "") for m in response.json().get("models", [])]
                if self.model not in available_models:
                    logger.warning(
                        f"Model '{self.model}' not found in Ollama. "
                        f"Available: {available_models}"
                    )
            else:
                raise ConnectionError("Ollama not responding")
        except Exception as e:
            logger.error(f"✗ Cannot connect to Ollama: {e}")
            raise

    def encode(self, texts: List[str], batch_size: int = 8,
               show_progress_bar: bool = False) -> List[Optional[List[float]]]:
        """
        Encode texts using Ollama BGE-M3 with per-text retry and automatic
        pre-truncation for oversized inputs.

        BGE-M3 in Ollama returns HTTP 500 when the prompt exceeds the model's
        context window.  Every text is pre-truncated to MAX_EMBED_CHARS before
        the first request — this eliminates context-window 500s entirely without
        consuming a retry slot.  The three retry attempts are reserved for
        genuine transient server or network errors.

        On any failure the Ollama response body is included in the warning log
        so the root cause is visible without inspecting Ollama's own log file.
        A per-batch summary warns if any chunks end up missing from the index.
        """
        embeddings: List[Optional[List[float]]] = []
        skipped_count = 0

        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]

            for text in batch:
                # Pre-truncate BEFORE the first request.  MAX_EMBED_CHARS (~1 000 tokens)
                # is well within BGE-M3's 8192-token context window, so this never
                # loses meaningful content and prevents all context-overflow 500s.
                safe_text = text if len(text) <= MAX_EMBED_CHARS else text[:MAX_EMBED_CHARS]

                embedding = None
                for attempt in range(3):
                    try:
                        response = requests.post(
                            f"{self.base_url}/api/embeddings",
                            json={"model": self.model, "prompt": safe_text},
                            timeout=60
                        )
                        if response.status_code == 200:
                            data = response.json()
                            if "embedding" in data:
                                embedding = data["embedding"]
                                break
                            else:
                                # Ollama API changed key name or returned an error body
                                logger.warning(
                                    f"Embedding attempt {attempt + 1}: "
                                    f"unexpected response format (no 'embedding' key): "
                                    f"{response.text[:120]}"
                                )
                        else:
                            logger.warning(
                                f"Embedding attempt {attempt + 1} failed with status "
                                f"{response.status_code}: {response.text[:120]}"
                            )
                    except Exception as e:
                        logger.warning(f"Embedding attempt {attempt + 1} error: {e}")
                    if attempt < 2:
                        time.sleep(2 ** attempt)

                if embedding is None:
                    skipped_count += 1
                    logger.warning(
                        f"Failed to embed chunk after 3 attempts "
                        f"(sent {len(safe_text)} chars, original {len(text)} chars). "
                        f"Total skipped so far: {skipped_count}"
                    )
                embeddings.append(embedding)

            if show_progress_bar and (i // batch_size) % 10 == 0:
                logger.info(f"Encoded {min(i + batch_size, len(texts))}/{len(texts)} texts")

        if skipped_count:
            logger.warning(
                f"⚠  {skipped_count}/{len(texts)} chunks could not be embedded and "
                f"will be missing from the index. Check Ollama logs for details."
            )

        return embeddings


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

class SectionAwareChunker:
    """Advanced chunker that respects document section structure."""

    def __init__(self, chunk_size: int = 512, overlap: int = 256):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.section_patterns = [re.compile(p, re.MULTILINE) for p in SECTION_PATTERNS]

    def detect_sections(self, text: str) -> List[DocumentSection]:
        """Detect document sections and hierarchy from plain text."""
        sections = []
        lines = text.split('\n')
        current_section: Dict = {"title": "Introduction", "content": "", "level": 0}
        section_stack = [current_section]

        for line in lines:
            line_stripped = line.strip()

            if not line_stripped:
                current_section["content"] += "\n"
                continue

            is_header = False
            header_level = 0

            for pattern in self.section_patterns:
                match = pattern.match(line_stripped)
                if match:
                    is_header = True
                    if line_stripped.startswith('#'):
                        header_level = len(line_stripped) - len(line_stripped.lstrip('#'))
                    elif line_stripped.isupper():
                        header_level = 1
                    else:
                        header_level = 2
                    break

            if is_header and len(line_stripped) < 200:
                if current_section["content"].strip():
                    hierarchy = [s["title"] for s in section_stack if s["title"]]
                    section_type = (
                        "table"
                        if current_section["title"].startswith("[Table ")
                        else "text"
                    )
                    sections.append(DocumentSection(
                        title=current_section["title"],
                        content=current_section["content"].strip(),
                        level=current_section["level"],
                        section_type=section_type,
                        section_hierarchy=list(hierarchy),
                    ))

                while len(section_stack) > 1 and section_stack[-1]["level"] >= header_level:
                    section_stack.pop()

                current_section = {
                    "title": line_stripped.strip('#: ').strip(),
                    "content": "",
                    "level": header_level
                }
                section_stack.append(current_section)
            else:
                current_section["content"] += line + "\n"

        if current_section["content"].strip():
            hierarchy = [s["title"] for s in section_stack if s["title"]]
            section_type = (
                "table"
                if current_section["title"].startswith("[Table ")
                else "text"
            )
            sections.append(DocumentSection(
                title=current_section["title"],
                content=current_section["content"].strip(),
                level=current_section["level"],
                section_type=section_type,
                section_hierarchy=list(hierarchy),
            ))

        return sections if sections else [DocumentSection("Document", text, 0)]

    def _split_long_sentence(self, sentence: str) -> List[str]:
        """Split a sentence that exceeds MAX_CHUNK_SIZE at word boundaries."""
        if len(sentence) <= MAX_CHUNK_SIZE:
            return [sentence]
        words = sentence.split()
        parts: List[str] = []
        current = ""
        for word in words:
            if len(current) + len(word) + 1 > MAX_CHUNK_SIZE and current:
                parts.append(current.strip())
                current = word
            else:
                current = current + " " + word if current else word
        if current.strip():
            parts.append(current.strip())
        return parts

    def chunk_with_sentences(
        self,
        text: str,
        section_title: str = "",
        section_hierarchy: Optional[List[str]] = None
    ) -> List[EnrichedChunk]:
        """Chunk text respecting sentence boundaries with configurable overlap."""
        if not text or len(text) < MIN_CHUNK_SIZE:
            return []

        if section_hierarchy is None:
            section_hierarchy = [section_title]

        raw_sentences = sent_tokenize(text)
        sentences: List[str] = []
        for s in raw_sentences:
            sentences.extend(self._split_long_sentence(s))

        chunks: List[EnrichedChunk] = []
        current_chunk = ""
        current_start = 0

        for sentence in sentences:
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                chunks.append(EnrichedChunk(
                    text=current_chunk.strip(),
                    section_title=section_title,
                    section_hierarchy=list(section_hierarchy),
                    page_number=None,
                    chunk_type="text",
                    word_count=len(word_tokenize(current_chunk)),
                    sentence_count=len(sent_tokenize(current_chunk)),
                    start_char=current_start,
                    end_char=current_start + len(current_chunk)
                ))

                old_chunk_len = len(current_chunk)
                overlap_text = (
                    current_chunk[-self.overlap:]
                    if len(current_chunk) > self.overlap
                    else current_chunk
                )
                current_chunk = overlap_text + " " + sentence
                current_start += old_chunk_len - len(overlap_text)
            else:
                current_chunk += " " + sentence if current_chunk else sentence

        if current_chunk.strip():
            chunks.append(EnrichedChunk(
                text=current_chunk.strip(),
                section_title=section_title,
                section_hierarchy=list(section_hierarchy),
                page_number=None,
                chunk_type="text",
                word_count=len(word_tokenize(current_chunk)),
                sentence_count=len(sent_tokenize(current_chunk)),
                start_char=current_start,
                end_char=current_start + len(current_chunk)
            ))

        return chunks

    def chunk_sections(self, sections: List[DocumentSection]) -> List[EnrichedChunk]:
        """Chunk all document sections while preserving section context."""
        all_chunks: List[EnrichedChunk] = []

        for section in sections:
            hierarchy = (
                section.section_hierarchy
                if hasattr(section, 'section_hierarchy')
                else [section.title]
            )
            section_chunks = self.chunk_with_sentences(section.content, section.title, hierarchy)

            for chunk in section_chunks:
                chunk.page_number = section.page_number
                chunk.chunk_type = section.section_type

            all_chunks.extend(section_chunks)

        return all_chunks


# ---------------------------------------------------------------------------
# BM25 index
# ---------------------------------------------------------------------------

class BM25Index:
    """BM25 sparse vector index."""

    def __init__(self):
        self.bm25 = None
        self.tokenized_corpus: List[List[str]] = []
        self.vocabulary: Dict[str, int] = {}
        self.token_idf: Dict[str, float] = {}

    def fit(self, texts: List[str]):
        """Build BM25 index from a list of text strings."""
        self.tokenized_corpus = [self._tokenize(t) for t in texts]
        self.bm25 = BM25Okapi(self.tokenized_corpus)

        all_tokens = sorted({t for doc in self.tokenized_corpus for t in doc})
        self.vocabulary = {token: idx for idx, token in enumerate(all_tokens)}

        N = len(self.tokenized_corpus)
        for token in self.vocabulary:
            df = sum(1 for doc in self.tokenized_corpus if token in doc)
            self.token_idf[token] = np.log((N - df + 0.5) / (df + 0.5) + 1)

        logger.info(f"  ✓ BM25 vocabulary size: {len(self.vocabulary)}")

    def _tokenize(self, text: str) -> List[str]:
        return [token.lower() for token in word_tokenize(text) if token.isalnum()]

    def get_sparse_vector(self, text: str) -> SparseVector:
        """Return a BM25 TF-IDF sparse vector for the given text."""
        tokens = self._tokenize(text)
        total = len(tokens)
        token_counts: Dict[str, int] = {}

        for token in tokens:
            if token in self.vocabulary:
                token_counts[token] = token_counts.get(token, 0) + 1

        indices: List[int] = []
        values: List[float] = []

        for token, count in token_counts.items():
            tf = count / total if total else 0.0
            idf = self.token_idf.get(token, 1.0)
            indices.append(self.vocabulary[token])
            values.append(float(tf * idf))

        return SparseVector(indices=indices, values=values)


# ---------------------------------------------------------------------------
# Document loader (V2: pypdfium2 for PDF text, pdfplumber for table counts)
# ---------------------------------------------------------------------------

class AdvancedDocumentLoader:
    """Enhanced document loader with pypdfium2-based PDF extraction."""

    # TOC line pattern: text … dots … page number
    _TOC_LINE_RE = re.compile(r'(\.\s*){2,}.*\d+\s*$')

    @staticmethod
    def _is_toc_page(page_text: str) -> bool:
        """
        Return True if this page looks like a Table-of-Contents / boilerplate
        page that adds no retrievable content to the index.

        Classification rules (all must pass to be considered TOC):

        1. Near-empty pages (< TOC_MIN_CONTENT_CHARS chars) are always dropped —
           these are blank separators or single-line chapter dividers.
        2. Content-dense pages (> TOC_MAX_CONTENT_CHARS chars) are NEVER dropped
           regardless of ratio, preventing false positives on real specification
           pages that happen to contain a version annotation like
           "Version ......... 3.0.0".
        3. Pages with fewer than TOC_MIN_LINE_COUNT non-empty lines are NOT
           ratio-classified — they are section headings, not TOC pages.
        4. A page is classified as TOC only when MORE THAN TOC_LINE_RATIO of its
           non-empty lines match the leader-dot pattern  (text . . . number).
           The threshold is 0.50 (majority), not 0.30, to avoid false positives
           from pages with a single dotted version line among real content.
        """
        stripped = page_text.strip()

        # Rule 1 – near-empty page: blank separator, single-line chapter title
        if len(stripped) < TOC_MIN_CONTENT_CHARS:
            return True

        # Rule 2 – content-dense page: too much text to be a pure TOC page
        if len(stripped) > TOC_MAX_CONTENT_CHARS:
            return False

        # At this point stripped is between 50 and 800 chars of non-whitespace content.
        # splitlines() on such a string always yields at least one non-empty line,
        # so no explicit empty-list guard is needed here.
        non_empty_lines = [l for l in stripped.splitlines() if l.strip()]

        # Rule 3 – too few lines to make a reliable ratio decision
        if len(non_empty_lines) < TOC_MIN_LINE_COUNT:
            return False

        # Rule 4 – majority of lines must be TOC-style entries
        toc_line_count = sum(
            1 for line in non_empty_lines
            if AdvancedDocumentLoader._TOC_LINE_RE.search(line)
        )
        return (toc_line_count / len(non_empty_lines)) >= TOC_LINE_RATIO

    @staticmethod
    def extract_metadata(path: str) -> Dict:
        """Extract filesystem metadata."""
        file_stat = os.stat(path)
        return {
            "file_size_bytes": file_stat.st_size,
            "created_timestamp": file_stat.st_ctime,
            "modified_timestamp": file_stat.st_mtime,
            "file_extension": Path(path).suffix.lower(),
        }

    @staticmethod
    def load_pdf(path: str) -> Tuple[str, Dict]:
        """
        Load PDF text using pypdfium2 (proper word spacing) and gather table-count
        metadata using pdfplumber (table extraction only, not for text content).

        TOC pages and near-empty pages are skipped automatically.
        All PDF handles are closed in finally blocks to prevent file-descriptor
        leaks even when a page is corrupted or raises an exception mid-loop.
        """
        metadata: Dict = {"num_pages": 0, "has_tables": False, "tables_count": 0}
        text_parts: List[str] = []
        skipped_pages: List[int] = []

        # --- text extraction via pypdfium2 ---
        pdf = pdfium.PdfDocument(path)
        try:
            num_pages = len(pdf)
            metadata["num_pages"] = num_pages

            for page_num in range(num_pages):
                page = pdf[page_num]
                page_text = ""          # initialise so the variable is always bound
                try:
                    textpage = page.get_textpage()
                    try:
                        page_text = textpage.get_text_bounded()
                    finally:
                        textpage.close()
                except Exception as page_err:
                    # Corrupted or unreadable page — skip it and continue with the rest
                    logger.warning(
                        f"  ⚠ Could not extract text from page {page_num + 1}: {page_err}"
                    )
                finally:
                    page.close()

                if not page_text:
                    continue

                if AdvancedDocumentLoader._is_toc_page(page_text):
                    skipped_pages.append(page_num + 1)
                    continue

                text_parts.append(f"\n[Page {page_num + 1}]\n{page_text}\n")
        finally:
            pdf.close()

        if skipped_pages:
            logger.info(
                f"  ⊘ Skipped {len(skipped_pages)} TOC/boilerplate pages: "
                f"{skipped_pages[:10]}{'...' if len(skipped_pages) > 10 else ''}"
            )

        # --- table count metadata via pdfplumber (no text used) ---
        try:
            with pdfplumber.open(path) as plumber_pdf:
                for page in plumber_pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        metadata["has_tables"] = True
                        metadata["tables_count"] += len(tables)
        except Exception as e:
            logger.warning(f"  pdfplumber table-count failed (non-fatal): {e}")

        return "".join(text_parts), metadata

    @staticmethod
    def load_docx(path: str) -> Tuple[str, Dict]:
        """Load DOCX with metadata."""
        doc = docx.Document(path)
        text_parts: List[str] = []
        metadata: Dict = {"num_paragraphs": 0, "has_tables": False, "tables_count": 0}

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                metadata["num_paragraphs"] += 1

        if doc.tables:
            metadata["has_tables"] = True
            metadata["tables_count"] = len(doc.tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)

        return "\n".join(text_parts), metadata

    @staticmethod
    def load_txt(path: str) -> Tuple[str, Dict]:
        """Load TXT file."""
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        metadata = {
            "num_lines": len(text.split('\n')),
            "char_count": len(text)
        }
        return text, metadata

    @classmethod
    def load(cls, path: str) -> Tuple[Optional[str], Dict]:
        """Universal loader: dispatches to the right format handler."""
        try:
            base_metadata = cls.extract_metadata(path)

            if path.lower().endswith(".pdf"):
                text, doc_metadata = cls.load_pdf(path)
            elif path.lower().endswith(".docx"):
                text, doc_metadata = cls.load_docx(path)
            elif path.lower().endswith(".txt"):
                text, doc_metadata = cls.load_txt(path)
            else:
                return None, {}

            base_metadata.update(doc_metadata)
            return text, base_metadata

        except Exception as e:
            logger.error(f"Error loading {path}: {e}")
            return None, {}


# ---------------------------------------------------------------------------
# Main ingestion pipeline
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Advanced RAG Ingestion System V2")
    parser.add_argument('--data-dir', type=str, default=DATA_DIR,
                        help='Directory containing documents')
    parser.add_argument('--collection', type=str, default=COLLECTION,
                        help='Qdrant collection name')
    parser.add_argument('--qdrant-url', type=str, default=QDRANT_URL,
                        help='Qdrant server URL')
    parser.add_argument('--ollama-url', type=str, default=OLLAMA_URL,
                        help='Ollama server URL')
    parser.add_argument('--ollama-model', type=str, default=OLLAMA_MODEL,
                        help='Ollama embedding model name')
    parser.add_argument('--chunk-size', type=int, default=CHUNK_SIZE,
                        help='Chunk size in characters')
    parser.add_argument('--chunk-overlap', type=int, default=CHUNK_OVERLAP,
                        help='Chunk overlap in characters')
    parser.add_argument('--bm25-output', type=str, default=BM25_OUTPUT,
                        help='Path to save BM25 index JSON')
    args = parser.parse_args()

    data_dir = args.data_dir
    collection = args.collection
    qdrant_url = args.qdrant_url
    ollama_url = args.ollama_url
    ollama_model = args.ollama_model
    chunk_size = args.chunk_size
    chunk_overlap = args.chunk_overlap
    bm25_output = args.bm25_output

    logger.info("=" * 80)
    logger.info("ADVANCED RAG INGESTION SYSTEM V2")
    logger.info("=" * 80)

    # --- Qdrant client ---
    client = QdrantClient(url=qdrant_url)

    # --- Embedding model ---
    if USE_OLLAMA_BGE_M3:
        try:
            logger.info(f"Using Ollama BGE-M3 model at {ollama_url}")
            embedder = OllamaBGEM3Embedder(ollama_url, ollama_model)
            embedding_dim = embedder.dimension
        except Exception:
            logger.warning(f"Ollama unavailable, falling back to {FALLBACK_MODEL}")
            embedder = SentenceTransformer(FALLBACK_MODEL)
            embedding_dim = 384
    else:
        logger.info(f"Using SentenceTransformer: {FALLBACK_MODEL}")
        embedder = SentenceTransformer(FALLBACK_MODEL)
        embedding_dim = 384

    # --- Collection setup ---
    existing = [c.name for c in client.get_collections().collections]
    if collection not in existing:
        client.create_collection(
            collection_name=collection,
            vectors_config={
                "dense": VectorParams(size=embedding_dim, distance=Distance.COSINE)
            },
            sparse_vectors_config={
                "bm25": SparseVectorParams(index=SparseIndexParams(on_disk=False))
            },
        )
        logger.info(f"✓ Created collection: {collection}")
    else:
        col_info = client.get_collection(collection)
        existing_dim = None
        if hasattr(col_info.config.params.vectors, "get"):
            dense_cfg = col_info.config.params.vectors.get("dense")
            if dense_cfg is not None:
                existing_dim = dense_cfg.size
        if existing_dim is not None and existing_dim != embedding_dim:
            logger.error(
                f"✗ Dimension mismatch: collection '{collection}' has {existing_dim}-dim vectors, "
                f"but current embedder produces {embedding_dim}-dim vectors. Aborting."
            )
            return
        logger.info(f"✓ Collection exists: {collection}")

    # --- Pipeline components ---
    chunker = SectionAwareChunker(chunk_size, chunk_overlap)
    bm25_index = BM25Index()

    total_chunks = 0
    total_files = 0
    all_points: List[Tuple[Dict, str]] = []

    logger.info("\nPhase 1: Loading and chunking documents...")

    for root, _, files in os.walk(data_dir):
        for file in files:
            if not file.lower().endswith((".pdf", ".docx", ".txt")):
                continue

            path = os.path.join(root, file)
            logger.info(f"\n→ Processing: {file}")

            try:
                h = file_hash(path)

                if already_indexed(client, collection, h):
                    logger.info("  ⊘ Already indexed")
                    continue

                text, doc_metadata = AdvancedDocumentLoader.load(path)
                if not text or len(text.strip()) < MIN_CHUNK_SIZE:
                    logger.info("  ⊘ Empty or too short after extraction")
                    continue

                # Section-aware chunking
                if ENABLE_SECTION_AWARE:
                    sections = chunker.detect_sections(text)
                    raw_chunks = chunker.chunk_sections(sections)
                    logger.info(
                        f"  ✓ Detected {len(sections)} sections → {len(raw_chunks)} raw chunks"
                    )
                else:
                    raw_chunks = chunker.chunk_with_sentences(text, "Document")
                    logger.info(f"  ✓ Created {len(raw_chunks)} raw chunks")

                if not raw_chunks:
                    continue

                # --- Near-duplicate chunk removal ---
                accepted_chunks = []
                dedup_window: List[str] = []   # sliding window of last JACCARD_WINDOW texts

                for chunk in raw_chunks:
                    is_dup = False
                    for prev_text in dedup_window[-JACCARD_WINDOW:]:
                        if jaccard_similarity(chunk.text, prev_text) > JACCARD_DEDUP_THRESHOLD:
                            is_dup = True
                            break
                    if not is_dup:
                        accepted_chunks.append(chunk)
                        dedup_window.append(chunk.text)

                dedup_removed = len(raw_chunks) - len(accepted_chunks)
                if dedup_removed:
                    logger.info(f"  ⊘ Removed {dedup_removed} near-duplicate chunks")

                chunks = accepted_chunks
                if not chunks:
                    continue

                chunk_texts = [chunk.text for chunk in chunks]

                # --- Embeddings ---
                logger.info(f"  ⚡ Generating embeddings for {len(chunk_texts)} chunks...")
                if isinstance(embedder, OllamaBGEM3Embedder):
                    embeddings = embedder.encode(chunk_texts, batch_size=8)
                else:
                    raw_embeddings = embedder.encode(
                        chunk_texts,
                        batch_size=32,
                        show_progress_bar=False,
                        convert_to_numpy=True
                    )
                    embeddings = [emb.tolist() for emb in raw_embeddings]

                # --- Build Qdrant point records ---
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    if embedding is None:
                        logger.warning(f"  ⚠ Skipping chunk {i} — embedding failed")
                        continue
                    point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{h}_{i}"))
                    point = {
                        "id": point_id,
                        "vector": embedding,
                        "payload": {
                            "content": chunk.text,
                            "source_path": path,
                            "filename": file,
                            "folder": os.path.relpath(root, data_dir),
                            "file_type": Path(file).suffix.lower(),
                            "file_hash": h,
                            "chunk_id": i,
                            "total_chunks": len(chunks),
                            "chunk_length": len(chunk.text),
                            "word_count": chunk.word_count,
                            "sentence_count": chunk.sentence_count,
                            "start_char": chunk.start_char,
                            "end_char": chunk.end_char,
                            "section_title": chunk.section_title,
                            "section_hierarchy": chunk.section_hierarchy,
                            "chunk_type": chunk.chunk_type,
                            **doc_metadata
                        }
                    }
                    all_points.append((point, chunk.text))
                    total_chunks += 1

                total_files += 1

            except Exception as e:
                logger.error(f"  ✗ Error processing {file}: {e}")
                continue

    if not all_points:
        logger.warning("No documents to index!")
        return

    # --- Phase 2: BM25 ---
    all_texts_for_bm25 = [text for _, text in all_points]
    logger.info(f"\nPhase 2: Building BM25 index for {len(all_texts_for_bm25)} chunks...")
    bm25_index.fit(all_texts_for_bm25)

    bm25_data = {
        "vocabulary": bm25_index.vocabulary,
        "token_idf": {k: float(v) for k, v in bm25_index.token_idf.items()},
    }
    Path(bm25_output).parent.mkdir(parents=True, exist_ok=True)
    with open(bm25_output, "w") as f:
        json.dump(bm25_data, f)
    logger.info(f"  ✓ Saved BM25 index to {bm25_output}")

    # --- Phase 3: Upload to Qdrant ---
    logger.info("\nPhase 3: Adding sparse vectors and uploading to Qdrant...")

    points_to_upload: List[PointStruct] = []
    for point_data, chunk_text in all_points:
        sparse_vector = bm25_index.get_sparse_vector(chunk_text)
        points_to_upload.append(PointStruct(
            id=point_data["id"],
            vector={
                "dense": point_data["vector"],
                "bm25": sparse_vector
            },
            payload=point_data["payload"]
        ))

    batch_size = 100
    num_batches = (len(points_to_upload) - 1) // batch_size + 1
    for i in range(0, len(points_to_upload), batch_size):
        batch = points_to_upload[i:i + batch_size]
        client.upsert(collection_name=collection, points=batch, wait=True)
        logger.info(f"  ✓ Uploaded batch {i // batch_size + 1}/{num_batches}")

    # --- Summary ---
    logger.info("\n" + "=" * 80)
    logger.info("INGESTION COMPLETE")
    logger.info("=" * 80)
    logger.info(f"✓ Files processed:      {total_files}")
    logger.info(f"✓ Total chunks indexed: {total_chunks}")
    logger.info(
        f"✓ Avg chunks/file:     "
        f"{total_chunks / total_files if total_files > 0 else 0:.1f}"
    )
    logger.info(f"✓ Collection:           {collection}")
    logger.info(f"✓ Embedding dimension:  {embedding_dim}")
    logger.info(f"✓ BM25 vocabulary size: {len(bm25_index.vocabulary)}")
    logger.info("=" * 80)


if __name__ == "__main__":
    main()
