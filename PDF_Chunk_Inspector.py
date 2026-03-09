#!/usr/bin/env python3
"""
PDF_Chunk_Inspector.py
======================
Standalone debugging / inspection tool for the V3 RAG ingestion pipeline.

Purpose
-------
This script takes a single PDF file, runs the full V3 page-classification and
chunking pipeline on it, and writes a detailed human-readable report so you can
see exactly:

  • How the classifier (heuristic + optional Ollama vision) labels each page.
  • Which pages are skipped (TOC / boilerplate) and why.
  • How sections are detected inside each page's extracted text.
  • How many parent chunks and child chunks come from each page individually,
    including their character ranges and text previews.
  • Per-page-type statistics and a summary table at the end.
  • Anomaly detection: pages with extracted text but zero chunks.

This file does NOT write anything to Qdrant or modify any existing code.
All pipeline classes and configuration are self-contained inside this single
file — no external module is required.  The logic mirrors the real ingestion
pipeline exactly, so what you see here is exactly what would be stored in the
vector database.

Design note on per-page vs global chunking
-------------------------------------------
The real ingestion pipeline calls detect_sections + chunk on the FULL concatenated
document text (so chunks can span page boundaries, which is intentional for better
recall).  This inspector runs the same chunker on each PAGE INDIVIDUALLY so that
the report can clearly show "page N produced X parent + Y child chunks".
The per-page numbers are therefore close to — but not identical to — the numbers
the real ingestion produces.  The final summary section reports the GLOBAL run
(same as real ingestion) for comparison.

Usage
-----
    python PDF_Chunk_Inspector.py --pdf path/to/document.pdf

Options
    --pdf               Path to the PDF to inspect  (required)
    --output            Path for the output report file
                        (default: <pdf-stem>_inspection.txt next to the PDF)
    --no-vision         Skip Ollama vision classifier; use heuristics only
    --ollama-url        Ollama base URL (default: http://localhost:11434)
    --vision-model      Ollama vision model for classification
    --vision-desc-model Ollama model for image/diagram description
    --top-k-preview N   Chars of chunk text to preview in the report (default: 300)
"""

# ── All classes and constants inlined from gen.py (fully standalone) ──────────

import argparse
import base64
import io
import json
import time
import os
import re
import sys
import shutil
import uuid
import collections
import statistics
from typing import Optional, List, Dict, Tuple, Set
from dataclasses import dataclass, field
from pathlib import Path
import logging

import numpy as np
import pdfplumber
import pypdfium2 as pdfium
import docx
import requests
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize

# Optional OCR imports (graceful degradation if not installed)
try:
    from pdf2image import convert_from_path as pdf_to_images
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Optional Pillow (needed for rendering PDF pages to images for Ollama vision)
try:
    from PIL import Image as PILImage, ImageDraw as PILImageDraw, ImageFont as PILImageFont
    PILLOW_AVAILABLE = True
except ImportError:
    PILImage = None
    PILImageDraw = None
    PILImageFont = None
    PILLOW_AVAILABLE = False

# Optional sklearn for semantic chunking
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity as sk_cosine
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

# Optional PyMuPDF for layout analysis (vector drawings, tables, figures)
try:
    import fitz as _fitz   # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    _fitz = None
    PYMUPDF_AVAILABLE = False

# ─── NLTK bootstrap ─────────────────────────────────────────────────────────
for _resource in ("tokenizers/punkt", "tokenizers/punkt_tab",
                  "corpora/stopwords"):
    try:
        nltk.data.find(_resource)
    except LookupError:
        _name = _resource.split("/")[-1]
        try:
            nltk.download(_name, quiet=True)
        except Exception:
            pass

try:
    from nltk.corpus import stopwords as _sw
    _STOP_WORDS: Set[str] = set(_sw.words("english"))
except Exception:
    _STOP_WORDS: Set[str] = set()


# ═══════════════════════════════════════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════════════════════════════════════

OLLAMA_URL        = "http://localhost:11434"

# ── Page Classification (Ollama Vision) ─────────────────────────────────────
ENABLE_PAGE_CLASSIFICATION = True
# Fast lightweight model for page type classification
# Recommended: "moondream:latest" (< 2B params, very fast)
# Alternatives:  "llava:7b", "llava:13b", "llama3.2-vision:11b"
OLLAMA_VISION_MODEL        = "qwen2.5vl:7b"
# Richer vision model for describing images and flow diagrams
# Recommended: "llava:7b"   Alternatives: "llava:13b", "llama3.2-vision:11b"
OLLAMA_VISION_DESC_MODEL   = "qwen2.5vl:32b-q4_K_M"

# Text generation model for LLM-mode context enrichment
# Recommended: "llama3.2:3b" (fast), "llama3:8b" (higher quality)

# ── OCR ─────────────────────────────────────────────────────────────────────
# "pytesseract" — fast local OCR; requires tesseract binary
# "ollama_vision" — uses vision LLM to transcribe (no binary needed)
OCR_ENGINE = "ollama_vision"

# ── Vision classifier settings ───────────────────────────────────────────────
# Heuristic confidence below this triggers a vision LLM call for final decision
VISION_HEURISTIC_CONFIDENCE_THRESHOLD = 0.65
# Scale for rendering PDF pages (higher = better quality, slower)
PAGE_RENDER_SCALE         = 3.0
PAGE_RENDER_JPEG_QUALITY  = 85
MAX_VISION_RESPONSE_CHARS = 4000
VISION_API_TIMEOUT        = 120

# ── Chunking — tuned for AUTOSAR specifications ──────────────────────────────
# AUTOSAR documents contain dense requirement blocks (200-500 chars each) grouped
# inside numbered sub-sections.  A child size of ~800 chars captures 1-3 full
# requirement items; a parent size of ~3 500 chars covers a complete sub-section.
CHILD_CHUNK_SIZE    = 1000          # child chunk target size (chars)
CHILD_CHUNK_OVERLAP = 250          # enough to overlap one requirement item
PARENT_CHUNK_SIZE   = 3500         # parent chunk target size (chars)
PARENT_CHUNK_OVERLAP= 350          # ~10 % of parent
# Per-content-type chunk sizes — override defaults per PageType
# These dicts are intentionally empty here; add overrides as needed.
# The inspector falls back to CHILD_CHUNK_SIZE / PARENT_CHUNK_SIZE when a
# PageType key is missing.
CHILD_CHUNK_SIZES: Dict[str, int] = {}
PARENT_CHUNK_SIZES: Dict[str, int] = {}

MIN_CHUNK_SIZE      = 60           # don't discard short requirement tags / IDs

# Semantic chunking
# AUTOSAR paragraphs are highly cohesive; use a lower threshold so only clear
# topic shifts (e.g. concept description → rationale table) trigger a split.
ENABLE_SEMANTIC_CHUNKING   = True
SEMANTIC_SPLIT_THRESHOLD   = 0.20  # cosine sim below this → semantic boundary
SEMANTIC_WINDOW_SIZE       = 4     # sentences per window for sim computation


# ── PDF extraction ───────────────────────────────────────────────────────────
ENABLE_OCR_FALLBACK   = True       # use OCR_ENGINE if tier-1/2 yield sparse text
# AUTOSAR diagram/figure pages often contain very little text alongside an image.
# Raise the threshold slightly so OCR is triggered for those pages.
OCR_TRIGGER_CHARS     = 300        # page chars below this triggers tier-2 / tier-3

# TOC detection
# AUTOSAR Table-of-Contents pages can be very long (many deeply-nested sections).
# Raise TOC_MAX_CONTENT_CHARS so long TOC pages are still detected and skipped.
TOC_LINE_RATIO        = 0.50
TOC_MIN_CONTENT_CHARS = 50
TOC_MIN_LINE_COUNT    = 5
TOC_MAX_CONTENT_CHARS = 5000       # AUTOSAR TOCs can span ~100 entries (~5 000 chars)

# Embedding safety

# ── Smart Section Detection ───────────────────────────────────────────────────
# Model used for LLM-fallback section detection (Tier 3).
# Should be a fast text-only model (NOT a vision model).
# Recommended: \"qwen2.5:3b\", \"llama3.2:3b\"
OLLAMA_SECTION_MODEL          = "qwen2.5:3b"

# Minimum ratio of font size to body-text size to be treated as a heading candidate.
# 1.05 means "at least 5% larger than body text"
FONT_HEADING_SIZE_RATIO       = 1.05

# When font clustering finds zero headings on a page, fall back to LLM detection.
ENABLE_LLM_SECTION_FALLBACK   = True

# LLM section detection confidence threshold — lines with LLM confidence below
# this are treated as body text.
LLM_SECTION_CONFIDENCE        = 0.0   # accept all LLM-identified headings

# Max chars of page text to send to the section-detection LLM (keep prompt small).
LLM_SECTION_MAX_CHARS         = 3000

SECTION_PATTERNS = [
    # AUTOSAR deep-numbered section headers: "10.3.4.2 Module Overview"
    # or "10.3.4.2 [SWS_Os_00042] Some Requirement Title".
    # (?:[A-Z]|\[) is explicit: the section title starts with a capital letter
    # OR an opening bracket (requirement-ID-prefixed section titles).
    r'^\d+(\.\d+){1,5}\s+(?:[A-Z]|\[)',
    # Single-level AUTOSAR numbered headings: "1 Introduction", "4 Fundamentals"
    # (no dot — these are top-level chapter headings)
    r'^\d{1,2}\s+[A-Z][A-Za-z]',
    # AUTOSAR tagged requirement IDs used as standalone header lines:
    # [SWS_Os_00042], [SRS_ETHTSYN_00001], [ECUC_Com_00012], [AP_SomeModule_00001] …
    # Module names are CamelCase (Os, LinIf) or ALL_CAPS (ETHTSYN).
    # Underscores within the module portion are intentionally allowed to cover
    # compound module names such as [SWS_Lin_Interface_00001].
    r'^\[(?:SWS|SRS|RS|ECUC|TPS|CP|AP|ASWS|TR|EXP|MOD|SRLG|PRS|CONC|BSW|FUNC|COM|SPEC)_[A-Za-z0-9_]+\]',
    # Markdown headers (kept for any Markdown-exported AUTOSAR docs)
    r'^#{1,6}\s+(.+)$',
    # Single-level numbered sections with dot: "1. Introduction"
    r'^\d+\.\s+([A-Z].+)$',
    # Title-case field label with colon: "Description:", "Rationale:", "Use Case:"
    # Matches 3–41 total chars (1 uppercase initial + 2–40 more alphanumeric/space).
    r'^([A-Z][A-Za-z ]{2,40}):$',
    # All-caps headings (5–50 chars to avoid matching short enumerated constants).
    # Minimum of 5 total chars: [A-Z] + 4 more uppercase/space chars.
    r'^([A-Z][A-Z\s]{4,49})$',
    # Inline table markers injected by _extract_table_structured / _extract_tables_as_text.
    # Both variants must be listed because extraction uses different casing:
    #   _extract_table_structured → "[TABLE N]"  (all-caps)
    #   _extract_tables_as_text   → "[Table N]"  (title-case)
    # These patterns make detect_sections() split each table into its own
    # DocumentSection(section_type="table") so it is chunked row-by-row.
    r'^\[TABLE \d+\]$',
    r'^\[Table \d+\]$',
    # Visual-content markers injected by _extract_images_as_text().
    # Each embedded raster image gets [IMAGE N] and each vector diagram /
    # vision-described figure gets [DIAGRAM N].  detect_sections() turns them
    # into separate DocumentSection objects (section_type="image"/"diagram")
    # so visual content is never merged into surrounding prose chunks.
    r'^\[IMAGE \d+\]$',
    r'^\[DIAGRAM \d+\]$',
]

# Regex matching AUTOSAR running page headers that repeat on every page.
# These are stripped before section detection to avoid polluting the first
# section's content and hierarchy.
# Covers patterns like:
#   "Adaptive Platform Machine Configuration"
#   "AUTOSAR AP R25-11"   /  "AUTOSAR CP R22-11"
#   "Specification of Somemodule"
_AUTOSAR_RUNNING_HEADER_RE = re.compile(
    r'^(?:'
    r'AUTOSAR\s+(?:AP|CP|FO|TR)\s+R\d{2}-\d{2}'       # "AUTOSAR AP R25-11"
    r'|AUTOSAR\s+Release\s+\d'                           # "AUTOSAR Release 4.x"
    r'|(?:Specification|Requirements|Explanation|'
    r'Adaptive Platform|Classic Platform)\s+of\s+\S'    # "Specification of ..."
    r')',
    re.IGNORECASE,
)


# ═══════════════════════════════════════════════════════════════════════════
# Logging
# ═══════════════════════════════════════════════════════════════════════════

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════
# Page Types
# ═══════════════════════════════════════════════════════════════════════════

class PageType:
    """
    Canonical page-level content types for AUTOSAR-style technical standards.
    Covers both generic document structure and AUTOSAR-specific content patterns.
    Used throughout the pipeline for classification, extraction routing, and
    filtering decisions before vector indexing.
    """
    # --- Structural pages (filtered from indexing) ---
    COVER              = "cover"              # Title / cover page
    TOC                = "toc"                # Table of contents
    REVISION_HISTORY   = "revision_history"   # Document change history / release logs
    DISCLAIMER         = "disclaimer"         # Legal notices, liability, copyright text

    # --- Core content pages ---
    TEXT               = "text"               # Dense prose / standard paragraphs
    LIST               = "list"               # Bullet lists / enumerations
    TABLE              = "table"              # General tabular data
    DIAGRAM            = "diagram"            # Flow chart, block diagram, UML, schematic
    IMAGE              = "image"              # Photograph or raster graphic
    EQUATION           = "equation"           # Mathematical equations / formulae
    MIXED              = "mixed"              # Combination of multiple content types

    # --- AUTOSAR-specific pages ---
    CODE_SNIPPET       = "code_snippet"       # C headers, API signatures, macros, typedef blocks
    CLASS_REFERENCE    = "class_reference"    # UML class tables, meta-model definitions, attributes
    SPECIFICATION_ITEM = "specification_item" # Traceable constraints / requirements ([constr_xxx], [SWS_xxx])
    GLOSSARY           = "glossary"           # Acronyms, terms, and abbreviation definitions

    # --- Fallback ---
    UNKNOWN            = "unknown"            # Classifier was unable to determine type


# ═══════════════════════════════════════════════════════════════════════════
# Data classes
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class DocumentSection:
    title: str
    content: str
    level: int
    page_number: Optional[int] = None
    section_type: str = "text"
    section_hierarchy: Optional[List[str]] = None

    def __post_init__(self):
        if self.section_hierarchy is None:
            self.section_hierarchy = [self.title]


@dataclass
class ParentChunk:
    """Large context window stored in the parents collection."""
    text: str
    section_title: str
    section_hierarchy: List[str]
    page_number: Optional[int]
    chunk_type: str
    page_type: str = PageType.TEXT          # classified page type (PageType.*)
    word_count: int = 0
    start_char: int = 0
    end_char: int = 0
    parent_id: str = field(default_factory=lambda: str(uuid.uuid4()))


@dataclass
class ChildChunk:
    """Small precision chunk stored in the children collection."""
    text: str
    enriched_text: str          # text with contextual prefix (used for embedding)
    section_title: str
    section_hierarchy: List[str]
    page_number: Optional[int]
    chunk_type: str
    page_type: str = PageType.TEXT          # classified page type (PageType.*)
    word_count: int = 0
    sentence_count: int = 0
    start_char: int = 0
    end_char: int = 0
    parent_id: str = ""         # FK -> ParentChunk.parent_id
    child_index: int = 0        # position within parent


@dataclass
class PageClassificationResult:
    """Holds the classification outcome for a single PDF page."""
    page_type: str                  # one of PageType.*
    confidence: float               # 0-1 heuristic or 0/1 from LLM
    description: str                # short natural-language description of the page
    has_text: bool
    has_tables: bool
    has_images: bool
    text_density: float             # chars / estimated_page_area
    used_vision_model: bool = False # True when the Ollama vision LLM was consulted


# ═══════════════════════════════════════════════════════════════════════════
# Ollama Vision Classifier
# ═══════════════════════════════════════════════════════════════════════════

class OllamaVisionClassifier:
    """
    Two-tier AI-powered page classifier.

    Tier 1 (heuristic) — fast, no model call needed:
        Inspects text density, table count, and embedded-image presence
        using metadata already extracted by pypdfium2 / pdfplumber.
        Returns a (PageType, confidence) pair.

    Tier 2 (vision LLM) — called when confidence < VISION_HEURISTIC_CONFIDENCE_THRESHOLD
        or when the page appears to contain images/diagrams:
        Renders the page to a JPEG, encodes it as base64, and sends it to
        the Ollama vision model (default: moondream:latest).
        The model is asked to classify the page and give a brief description.

    All model names are configurable via CONFIG variables.
    """

    # Prompt used to classify page type with a vision model
    _CLASSIFY_PROMPT = (
        "You are a document analysis assistant specializing in AUTOSAR and technical standards PDFs. "
        "Look at this page and classify it. Choose EXACTLY ONE type from this list:\n"
        "  text, list, table, diagram, image, equation, mixed,\n"
        "  cover, toc, revision_history, disclaimer,\n"
        "  code_snippet, class_reference, specification_item, glossary\n\n"
        "Definitions:\n"
        "  text               - dense prose paragraphs\n"
        "  list               - primarily bullet points or numbered enumerations\n"
        "  table              - rows and columns of structured data\n"
        "  diagram            - flow chart, block diagram, UML, schematic, architecture diagram\n"
        "  image              - photograph, raster graphic, screenshot\n"
        "  equation           - mathematical or chemical equations dominate the page\n"
        "  mixed              - two or more content types side by side\n"
        "  cover              - title page / cover page\n"
        "  toc                - table of contents\n"
        "  revision_history   - document change log, version history table\n"
        "  disclaimer         - legal notices, liability statements, copyright text\n"
        "  code_snippet       - C headers, API signatures, macros, typedef blocks, source code\n"
        "  class_reference    - UML class table, meta-model attribute definitions\n"
        "  specification_item - formalized traceable requirements or constraints (e.g. [SWS_xxx], [constr_xxx])\n"
        "  glossary           - acronyms, terms, and abbreviation definitions\n\n"
        "Reply with a JSON object like: "
        '{"type": "<type>", "description": "<one sentence>"}\n'
        "Reply with JSON only, no markdown, no extra text."
    )

    # Prompt used to get a rich description of an image/diagram page
    _DESCRIBE_IMAGE_PROMPT = (
        "Describe this image from a technical document in detail. "
        "What does it show? Extract any text visible in the image. "
        "If it is a diagram or flow chart, describe the flow, components, and relationships. "
        "Be precise and comprehensive. Reply in plain text."
    )

    _DESCRIBE_DIAGRAM_PROMPT = (
        "This is a technical diagram from an engineering or software document. "
        "Describe it thoroughly: what type of diagram is it, what are the main components, "
        "what process or architecture does it represent, and what are the key relationships "
        "or data flows shown? Extract all visible text labels. "
        "Format your response as structured text that can be indexed for retrieval."
    )

    def __init__(
        self,
        ollama_url: str = OLLAMA_URL,
        classify_model: str = OLLAMA_VISION_MODEL,
        describe_model: str = OLLAMA_VISION_DESC_MODEL,
    ):
        self.ollama_url     = ollama_url
        self.classify_model = classify_model
        self.describe_model = describe_model
        self._available     = self._check_ollama()

    def _check_ollama(self) -> bool:
        try:
            r = requests.get(f"{self.ollama_url}/api/tags", timeout=3)
            if r.status_code == 200:
                logger.info(f"✓ OllamaVisionClassifier connected to {self.ollama_url}")
                return True
        except Exception:
            pass
        logger.warning(
            "OllamaVisionClassifier: Ollama unavailable — will use heuristics only."
        )
        return False

    # ── helpers ─────────────────────────────────────────────────────────────

    @staticmethod
    def _render_page_to_jpeg(pdf_path: str, page_index: int,
                              scale: float = PAGE_RENDER_SCALE,
                              quality: int = PAGE_RENDER_JPEG_QUALITY) -> Optional[str]:
        """
        Render a single PDF page to a JPEG and return it as a base64 string.
        Returns None if pypdfium2 or Pillow is unavailable.
        """
        if not PILLOW_AVAILABLE:
            return None
        try:
            doc = pdfium.PdfDocument(pdf_path)
            try:
                page = doc[page_index]
                try:
                    bitmap = page.render(scale=scale)
                    # pil_img = bitmap.to_pil()
                    pil_img = bitmap.to_pil().convert("RGB")
                    buf = io.BytesIO()
                    pil_img.save(buf, format="JPEG", quality=quality)
                    return base64.b64encode(buf.getvalue()).decode("utf-8")
                finally:
                    page.close()
            finally:
                doc.close()
        except Exception as e:
            logger.debug(f"Page render error for page {page_index+1}: {e}")
            return None

    def _call_vision(self, prompt: str, b64_image: str,
                     model: Optional[str] = None,
                     timeout: int = VISION_API_TIMEOUT) -> str:
        """
        Call an Ollama vision model with an image and a text prompt.
        Returns the response text, or "" on failure.
        """
        if not self._available or not b64_image:
            return ""
        m = model or self.classify_model
        try:
            resp = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": m,
                    "prompt": prompt,
                    "images": [b64_image],
                    "stream": False,
                    "options": {"temperature": 0.0, "num_predict": 256},
                },
                timeout=timeout,
            )
            if resp.status_code == 200:
                return resp.json().get("response", "").strip()[:MAX_VISION_RESPONSE_CHARS]
            logger.debug(f"Vision API HTTP {resp.status_code}: {resp.text[:80]}")
        except Exception as e:
            logger.debug(f"Vision API error: {e}")
        return ""

    # ── Tier 1: fast heuristic ───────────────────────────────────────────────

    @staticmethod
    def _heuristic_classify(
        tier1_text: str,
        plumber_page,          # pdfplumber page object (may be None)
        page_index: int,
    ) -> Tuple[str, float]:
        """
        Returns (PageType, confidence) using lightweight metadata signals only.
        Confidence is a value in [0, 1]; values below
        VISION_HEURISTIC_CONFIDENCE_THRESHOLD will trigger a vision LLM call.
        """
        text = tier1_text.strip()
        text_len = len(text)

        has_tables = False
        table_count = 0
        has_images = False
        img_count = 0

        if plumber_page is not None:
            try:
                tables = plumber_page.extract_tables()
                table_count = len(tables) if tables else 0
                has_tables = table_count > 0
            except Exception:
                pass
            try:
                imgs = plumber_page.images
                img_count = len(imgs) if imgs else 0
                has_images = img_count > 0
            except Exception:
                pass

        # Cover page: very short, often contains a single title and date
        if text_len < 150 and not has_tables and not has_images:
            return PageType.COVER, 0.65

        # Disclaimer: short page dominated by legal/copyright keywords
        if text_len < 1500 and not has_tables:
            lower = text.lower()
            disclaimer_hits = sum(1 for kw in (
                "copyright", "all rights reserved", "liability", "no warranty",
                "reproduction prohibited", "proprietary", "confidential",
                "legal notice", "disclaimer"
            ) if kw in lower)
            if disclaimer_hits >= 2:
                return PageType.DISCLAIMER, 0.82

        # Revision history: table with version/date/author columns
        if has_tables:
            lower = text.lower()
            rev_hits = sum(1 for kw in (
                "revision", "version", "author", "change", "date", "release", "history"
            ) if kw in lower)
            if rev_hits >= 3:
                return PageType.REVISION_HISTORY, 0.85

        # TOC: already handled by _is_toc_page; mark as low-confidence unknown
        # so the main pipeline may skip it.
        if text_len < TOC_MAX_CONTENT_CHARS:
            toc_re = re.compile(r"(\.\s*){2,}.*\d+\s*$")
            non_empty = [l for l in text.splitlines() if l.strip()]
            if non_empty:
                toc_frac = sum(1 for l in non_empty if toc_re.search(l)) / len(non_empty)
                if toc_frac > 0.35:
                    return PageType.TOC, 0.80

        # Pure image page (almost no text, has embedded images)
        if text_len < OCR_TRIGGER_CHARS and has_images and not has_tables:
            return PageType.IMAGE, 0.72

        # Diagram page heuristic: no text layer but has images with certain aspect ratios
        if text_len < OCR_TRIGGER_CHARS and has_images:
            return PageType.DIAGRAM, 0.55   # low confidence → will trigger vision

        # AUTOSAR-specific: code snippet — C headers, typedefs, macros
        code_lines = [l for l in text.splitlines() if re.search(
            r'^\s*(#include|#define|typedef|void |uint\d+_t|Std_ReturnType|FUNC\(|P2VAR\(|CONST\()', l)]
        if len(code_lines) >= 3:
            return PageType.CODE_SNIPPET, 0.85

        # AUTOSAR-specific: specification items — traceable requirement IDs
        spec_lines = [l for l in text.splitlines() if re.search(
            r'\[(SWS|RS|ECUC|constr|TPS|BSW|SRS)_\w+\]', l)]
        if len(spec_lines) >= 2:
            return PageType.SPECIFICATION_ITEM, 0.85

        # AUTOSAR-specific: class reference — UML class/attribute tables
        if has_tables:
            class_hits = sum(1 for kw in (
                "attribute", "multiplicity", "type", "category", "stereotype",
                "aggregation", "class", "meta", "role"
            ) if kw in text.lower())
            if class_hits >= 3:
                return PageType.CLASS_REFERENCE, 0.78

        # Glossary: lines of the form "ACRONYM  Definition of term..."
        non_empty_lines = [l for l in text.splitlines() if l.strip()]
        if non_empty_lines:
            glossary_re = re.compile(r'^[A-Z][A-Z0-9_\-]{1,15}\s{2,}.{10,}')
            gloss_frac = sum(1 for l in non_empty_lines if glossary_re.match(l.strip())) / len(non_empty_lines)
            if gloss_frac > 0.30 and text_len < 4000:
                return PageType.GLOSSARY, 0.80

        # Table-bearing page.
        if has_tables and table_count >= 1:
            if table_count > 1:
                return PageType.TABLE, 0.85
            if text_len < 600:
                return PageType.TABLE, 0.85
            return PageType.MIXED, 0.70

        # List-dominant page: many bullet / numbered lines
        if text_len > 100:
            list_re = re.compile(r'^\s*([-•*]|\d+[.)]) ')
            list_frac = sum(1 for l in non_empty_lines if list_re.match(l)) / max(len(non_empty_lines), 1)
            if list_frac > 0.45:
                return PageType.LIST, 0.82

        # Dense text page
        if text_len > 500 and not has_images:
            return PageType.TEXT, 0.90

        if text_len > 200:
            return PageType.TEXT, 0.75

        # Low-text, no structural signals
        if has_images:
            return PageType.IMAGE, 0.55    # likely image, but trigger vision
        return PageType.UNKNOWN, 0.40

    # ── Tier 2: vision LLM ──────────────────────────────────────────────────

    def _vision_classify(self, pdf_path: str, page_index: int) -> Tuple[str, str]:
        """
        Classify page type with an Ollama vision model.
        Returns (PageType, description) or (PageType.UNKNOWN, "") on failure.
        """
        b64 = self._render_page_to_jpeg(pdf_path, page_index)
        if not b64:
            return PageType.UNKNOWN, ""

        raw = self._call_vision(self._CLASSIFY_PROMPT, b64, self.classify_model)
        if not raw:
            return PageType.UNKNOWN, ""

        # Parse JSON response
        try:
            # Strip markdown code fences if present
            cleaned = re.sub(r"```[a-z]*\n?", "", raw).strip().strip("`")
            data = json.loads(cleaned)
            ptype = data.get("type", "unknown").lower().strip()
            desc  = data.get("description", "")
            valid = {
                PageType.TEXT, PageType.LIST, PageType.TABLE, PageType.DIAGRAM,
                PageType.IMAGE, PageType.EQUATION, PageType.MIXED,
                PageType.COVER, PageType.TOC, PageType.REVISION_HISTORY, PageType.DISCLAIMER,
                PageType.CODE_SNIPPET, PageType.CLASS_REFERENCE,
                PageType.SPECIFICATION_ITEM, PageType.GLOSSARY,
            }
            if ptype not in valid:
                ptype = PageType.UNKNOWN
            return ptype, desc
        except Exception:
            # Fallback: look for type keyword in freeform response
            low = raw.lower()
            for pt in [PageType.SPECIFICATION_ITEM, PageType.CLASS_REFERENCE,
                       PageType.CODE_SNIPPET, PageType.GLOSSARY,
                       PageType.REVISION_HISTORY, PageType.DISCLAIMER,
                       PageType.TABLE, PageType.IMAGE, PageType.DIAGRAM,
                       PageType.EQUATION, PageType.COVER, PageType.TOC,
                       PageType.LIST, PageType.MIXED, PageType.TEXT]:
                if pt in low:
                    return pt, raw[:200]
            return PageType.UNKNOWN, raw[:200]

    # ── Public API ──────────────────────────────────────────────────────────

    def classify(
        self,
        pdf_path: str,
        page_index: int,
        tier1_text: str,
        plumber_page,
    ) -> PageClassificationResult:
        """
        Classify a single PDF page.  Returns a PageClassificationResult.

        Steps:
        1. Run fast heuristic (no model call).
        2. If confidence < threshold OR page has images, call vision LLM.
        """
        text = tier1_text.strip()
        has_tables = False
        has_images = False

        if plumber_page is not None:
            try:
                has_tables = bool(plumber_page.extract_tables())
            except Exception:
                pass
            try:
                has_images = bool(plumber_page.images)
            except Exception:
                pass

        text_density = len(text) / max(1, 600 * 800)   # rough page area chars/pixel

        h_type, h_conf = self._heuristic_classify(tier1_text, plumber_page, page_index)
        description = ""
        used_vision = False

        needs_vision = (
            ENABLE_PAGE_CLASSIFICATION
            and self._available
            and PILLOW_AVAILABLE
            and (
                h_conf < VISION_HEURISTIC_CONFIDENCE_THRESHOLD
                or has_images
                or h_type in (PageType.DIAGRAM, PageType.IMAGE, PageType.UNKNOWN)
            )
        )

        if needs_vision:
            v_type, description = self._vision_classify(pdf_path, page_index)
            used_vision = True
            # Vision model wins when it disagrees with low-confidence heuristic
            if v_type != PageType.UNKNOWN:
                final_type = v_type
            else:
                final_type = h_type if h_type != PageType.UNKNOWN else PageType.TEXT
        else:
            final_type = h_type if h_type != PageType.UNKNOWN else PageType.TEXT

        return PageClassificationResult(
            page_type=final_type,
            confidence=1.0 if used_vision else h_conf,
            description=description,
            has_text=len(text) > 50,
            has_tables=has_tables,
            has_images=has_images,
            text_density=text_density,
            used_vision_model=used_vision,
        )


# ═══════════════════════════════════════════════════════════════════════════
# Tier 1a — ToC (Bookmark Tree) Extractor
# ═══════════════════════════════════════════════════════════════════════════

class ToCExtractor:
    """
    Extract the PDF bookmark tree using pypdfium2.
    Returns a sorted list of (title, page_number_1based, level) tuples.
    Level 1 = top-level chapter, 2 = section, 3 = subsection, etc.

    The bookmark tree is ground-truth structure authored by the document
    creator.  When present it is the most reliable signal for section
    boundaries — no model or heuristic is needed.
    """

    @staticmethod
    def extract(pdf_path: str) -> List[Tuple[str, int, int]]:
        """Return sorted list of (title, page_1based, level). Empty if no ToC."""
        bookmarks: List[Tuple[str, int, int]] = []
        try:
            doc = pdfium.PdfDocument(pdf_path)
            try:
                for item in doc.get_toc():
                    title     = (item.title or "").strip()
                    page_idx  = item.page_index   # 0-based; None if dest unknown
                    level     = (item.level or 0) + 1   # 1-based
                    if page_idx is not None and title:
                        bookmarks.append((title, page_idx + 1, level))
            finally:
                doc.close()
        except Exception as e:
            logger.debug(f"ToCExtractor: {e}")

        bookmarks.sort(key=lambda x: (x[1], x[2]))
        logger.info(f"ToCExtractor: found {len(bookmarks)} bookmark entries")
        return bookmarks

    @staticmethod
    def build_page_section_map(
        bookmarks: List[Tuple[str, int, int]],
        total_pages: int,
    ) -> Dict[int, List[Tuple[str, int]]]:
        """
        For each page number (1-based) build a list of bookmark entries whose
        *start* page equals that page.  Used by SmartSectionDetector to inject
        heading lines at the right position in the text stream.

        Returns: {page_num: [(title, level), ...]}
        """
        page_map: Dict[int, List[Tuple[str, int]]] = {}
        for title, page_num, level in bookmarks:
            page_map.setdefault(page_num, []).append((title, level))
        return page_map

    @staticmethod
    def hierarchy_at_page(
        bookmarks: List[Tuple[str, int, int]],
        page_num: int,
    ) -> List[str]:
        """
        Return the full active section hierarchy for *page_num* by walking
        the bookmark list and keeping the most recent entry at each level.
        """
        active: Dict[int, str] = {}
        for title, bm_page, level in bookmarks:
            if bm_page > page_num:
                break
            # When a shallower heading appears, clear all deeper cached levels
            for deeper in [l for l in active if l > level]:
                del active[deeper]
            active[level] = title

        if not active:
            return ["Document"]
        return ["Document"] + [active[l] for l in sorted(active)]


# ═══════════════════════════════════════════════════════════════════════════
# Tier 1b — Font-Size Heading Detector
# ═══════════════════════════════════════════════════════════════════════════

class FontSectionDetector:
    """
    Detect section headings from pdfplumber's character-level font metadata.

    Strategy
    --------
    1. Collect (fontname, size) for every character on the page via
       pdfplumber's `chars` stream.
    2. The *body text* size is the modal (most frequent) font size.
    3. Any line where ALL characters are >= body_size * FONT_HEADING_SIZE_RATIO
       AND at least one character is bold or a larger-than-body font is
       treated as a heading candidate.
    4. Heading level is derived from relative font size:
         largest distinct size  → level 1
         next larger           → level 2
         …etc.

    This works without any model and handles unnumbered, title-case, or
    bold headings that SECTION_PATTERNS regexes would miss entirely.
    """

    def __init__(self, size_ratio: float = FONT_HEADING_SIZE_RATIO):
        self.size_ratio = size_ratio

    def detect(
        self,
        plumber_page,
    ) -> List[Tuple[str, int]]:
        """
        Returns list of (line_text, level) for heading lines on this page.
        Level 1 is the most prominent heading.  Empty list if no headings found
        or if pdfplumber char metadata is unavailable.
        """
        if plumber_page is None:
            return []

        try:
            chars = plumber_page.chars
        except Exception:
            return []

        if not chars:
            return []

        # ── Step 1: determine body font size (modal) ──────────────────────
        size_counts: Dict[float, int] = {}
        for ch in chars:
            sz = round(float(ch.get("size", 0) or 0), 1)
            if sz > 0:
                size_counts[sz] = size_counts.get(sz, 0) + 1

        if not size_counts:
            return []

        body_size = max(size_counts, key=lambda s: size_counts[s])

        # Collect distinct heading sizes (larger than body * ratio)
        heading_sizes = sorted(
            {sz for sz in size_counts if sz >= body_size * self.size_ratio and sz > body_size},
            reverse=True,   # largest first
        )

        if not heading_sizes:
            return []

        # Build level map: largest size → level 1
        size_to_level: Dict[float, int] = {
            sz: i + 1 for i, sz in enumerate(heading_sizes)
        }

        # ── Step 2: group chars into lines by top-coordinate ──────────────
        # Round top to nearest 2 pts to group chars on the same visual line.
        lines_map: Dict[float, List[dict]] = {}
        for ch in chars:
            top = round(float(ch.get("top", 0) or 0) / 2) * 2
            lines_map.setdefault(top, []).append(ch)

        # ── Step 3: classify each line ────────────────────────────────────
        headings: List[Tuple[str, int]] = []
        for top in sorted(lines_map):
            line_chars = lines_map[top]
            line_text  = "".join(ch.get("text", "") for ch in line_chars).strip()
            if not line_text or len(line_text) > 200:
                continue

            # All char sizes on this line
            line_sizes = [
                round(float(ch.get("size", 0) or 0), 1)
                for ch in line_chars
                if ch.get("text", "").strip()
            ]
            if not line_sizes:
                continue

            dominant_size = max(set(line_sizes), key=line_sizes.count)

            if dominant_size in size_to_level:
                headings.append((line_text, size_to_level[dominant_size]))

        return headings


# ═══════════════════════════════════════════════════════════════════════════
# Tier 3 — LLM Section Detector (Ollama fallback)
# ═══════════════════════════════════════════════════════════════════════════

class LLMSectionDetector:
    """
    Use a local Ollama text model to identify section headings in page text.

    Only invoked when both Tier 1a (bookmark tree) and Tier 1b (font
    clustering) find no headings on a page.  A fast 3B text model is
    sufficient — no vision capability needed.
    """

    _PROMPT_TEMPLATE = """\
You are a document structure analyser. Given the raw text of one page from a technical PDF, identify all section heading lines and their nesting level.

Rules:
- A heading is a short line (< 120 chars) that introduces a new section of content.
- Level 1 = top-level chapter heading (e.g. "1 Introduction")
- Level 2 = section heading (e.g. "1.1 Purpose")
- Level 3+ = deeper subsections
- Body text, table rows, figure captions, and requirement IDs are NOT headings.
- If there are no headings on this page, return an empty list.

Reply ONLY with a JSON array. No explanation. No markdown. Example:
[{{"heading": "4 Crypto Stack Overview", "level": 1}}, {{"heading": "4.1 Stack Architecture", "level": 2}}]

PAGE TEXT:
\"\"\"
{page_text}
\"\"\"
"""

    def __init__(
        self,
        ollama_url: str = OLLAMA_URL,
        model: str = OLLAMA_SECTION_MODEL,
        timeout: int = 30,
    ):
        self.ollama_url = ollama_url
        self.model      = model
        self.timeout    = timeout
        self._available = self._check()

    def _check(self) -> bool:
        try:
            r = requests.get(f"{self.ollama_url}/api/tags", timeout=3)
            return r.status_code == 200
        except Exception:
            return False

    def detect(self, page_text: str) -> List[Tuple[str, int]]:
        """
        Returns list of (heading_text, level).  Empty on failure or no headings.
        """
        if not self._available or not page_text.strip():
            return []

        truncated = page_text[:LLM_SECTION_MAX_CHARS]
        prompt    = self._PROMPT_TEMPLATE.format(page_text=truncated)

        try:
            resp = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model":  self.model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {"temperature": 0.0, "num_predict": 512},
                },
                timeout=self.timeout,
            )
            if resp.status_code != 200:
                return []

            raw = resp.json().get("response", "").strip()
            # Strip markdown fences if present
            raw = re.sub(r"```[a-z]*\n?", "", raw).strip().strip("`")
            # Find the JSON array
            m = re.search(r"\[.*\]", raw, re.DOTALL)
            if not m:
                return []

            data = json.loads(m.group(0))
            results: List[Tuple[str, int]] = []
            for item in data:
                if isinstance(item, dict):
                    h = str(item.get("heading", "")).strip()
                    l = int(item.get("level", 2))
                    if h:
                        results.append((h, max(1, min(l, 6))))
            return results

        except Exception as e:
            logger.debug(f"LLMSectionDetector: {e}")
            return []


# ═══════════════════════════════════════════════════════════════════════════
# Smart Section Detector — orchestrates all three tiers
# ═══════════════════════════════════════════════════════════════════════════

class SmartSectionDetector:
    """
    Three-tier section detection orchestrator.

    Tier 1a  ToC bookmark tree       — ground truth, zero cost
    Tier 1b  Font-size clustering    — fast, no model needed
    Tier 3   LLM text model (Ollama) — fallback for poorly-structured pages

    The regex SECTION_PATTERNS still exists as a final safety net inside
    SemanticChunker.detect_sections() but is only reached when all three
    tiers produce nothing.
    """

    def __init__(
        self,
        toc_entries: List[Tuple[str, int, int]],
        font_detector: FontSectionDetector,
        llm_detector:  LLMSectionDetector,
        plumber_pages: Dict[int, object],   # 0-based page_index → plumber page
    ):
        self.toc_entries   = toc_entries
        self.font_detector = font_detector
        self.llm_detector  = llm_detector
        self.plumber_pages = plumber_pages

        # Pre-build page → bookmark list for fast lookup
        self.toc_page_map  = ToCExtractor.build_page_section_map(
            toc_entries, total_pages=99999
        )
        self._toc_available = bool(toc_entries)

    # ── Public API ────────────────────────────────────────────────────────

    def headings_for_page(
        self,
        page_num: int,       # 1-based
        page_text: str,
    ) -> List[Tuple[str, int]]:
        """
        Return confirmed headings as (text, level) for *page_num*.

        Priority:
          1. Bookmarks that START on this page (Tier 1a)
          2. Font-size analysis of this page (Tier 1b)
          3. LLM analysis of page text (Tier 3) — if ENABLE_LLM_SECTION_FALLBACK
        """
        # Tier 1a: bookmarks that START on this page
        if self._toc_available:
            toc_hits = self.toc_page_map.get(page_num, [])
            if toc_hits:
                logger.debug(f"  Page {page_num}: {len(toc_hits)} headings from ToC bookmark")
                return toc_hits   # (title, level) already

        # Tier 1b: font clustering
        plumber_pg = self.plumber_pages.get(page_num - 1)   # convert to 0-based
        font_hits  = self.font_detector.detect(plumber_pg)
        if font_hits:
            logger.debug(f"  Page {page_num}: {len(font_hits)} headings from font analysis")
            return font_hits

        # Tier 3: LLM fallback
        if ENABLE_LLM_SECTION_FALLBACK and page_text.strip():
            llm_hits = self.llm_detector.detect(page_text)
            if llm_hits:
                logger.debug(f"  Page {page_num}: {len(llm_hits)} headings from LLM")
                return llm_hits

        return []

    def hierarchy_at_page(self, page_num: int) -> List[str]:
        """Return the inherited section hierarchy for *page_num* from the ToC."""
        if self._toc_available:
            return ToCExtractor.hierarchy_at_page(self.toc_entries, page_num)
        return ["Document"]


# ═══════════════════════════════════════════════════════════════════════════
# Content-Type Extractor
# ═══════════════════════════════════════════════════════════════════════════

class ContentTypeExtractor:
    """
    Routes each classified PDF page to the appropriate extraction method.

    PageType  →  Strategy
    ─────────────────────────────────────────────────────────────────────────
    TEXT      →  pypdfium2 text (already extracted) — returned as-is
    TABLE     →  pdfplumber structured table (headers + rows)
    IMAGE     →  pytesseract OCR  OR  Ollama vision transcription
    DIAGRAM   →  Ollama llava vision description (structured prompt)
    EQUATION  →  pypdfium2 text + vision description fallback
    MIXED     →  combine text + table + vision approaches
    COVER     →  pypdfium2 text (usually short)
    TOC       →  skip (return empty string)

    All model names are read from CONFIG variables.
    """

    def __init__(
        self,
        ollama_url: str = OLLAMA_URL,
        classify_model: str = OLLAMA_VISION_MODEL,
        describe_model: str = OLLAMA_VISION_DESC_MODEL,
        ocr_engine: str = OCR_ENGINE,
    ):
        self.ollama_url     = ollama_url
        self.classify_model = classify_model
        self.describe_model = describe_model
        self.ocr_engine     = ocr_engine
        self._available     = self._check_ollama()

    def _check_ollama(self) -> bool:
        try:
            r = requests.get(f"{self.ollama_url}/api/tags", timeout=3)
            return r.status_code == 200
        except Exception:
            return False

    # ── OCR strategies ──────────────────────────────────────────────────────

    @staticmethod
    def _ocr_pytesseract(pdf_path: str, page_index: int) -> str:
        """Use pytesseract to OCR a single page. Returns text or ""."""
        if not OCR_AVAILABLE:
            return ""
        try:
            images = pdf_to_images(
                pdf_path, first_page=page_index + 1, last_page=page_index + 1, dpi=200
            )
            if images:
                return pytesseract.image_to_string(images[0])
        except Exception as e:
            logger.debug(f"pytesseract OCR page {page_index+1}: {e}")
        return ""

    def _ocr_ollama_vision(self, pdf_path: str, page_index: int) -> str:
        """Use Ollama vision LLM to transcribe a page. Returns text or ""."""
        if not self._available or not PILLOW_AVAILABLE:
            return ""
        b64 = OllamaVisionClassifier._render_page_to_jpeg(pdf_path, page_index)
        if not b64:
            return ""
        prompt = (
            "Transcribe all text visible on this page exactly as it appears. "
            "Preserve formatting, numbers, and special characters. "
            "If there are tables, reproduce them in pipe-delimited format. "
            "Reply with the transcribed text only."
        )
        try:
            resp = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": self.classify_model,
                    "prompt": prompt,
                    "images": [b64],
                    "stream": False,
                    "options": {"temperature": 0.0, "num_predict": 1024},
                },
                timeout=VISION_API_TIMEOUT,
            )
            if resp.status_code == 200:
                return resp.json().get("response", "").strip()
        except Exception as e:
            logger.debug(f"OCR via Ollama vision page {page_index+1}: {e}")
        return ""

    def _ocr(self, pdf_path: str, page_index: int) -> str:
        """Dispatch OCR to the configured engine."""
        if self.ocr_engine == "ollama_vision":
            text = self._ocr_ollama_vision(pdf_path, page_index)
            if not text:
                text = self._ocr_pytesseract(pdf_path, page_index)
        else:
            text = self._ocr_pytesseract(pdf_path, page_index)
            if not text:
                text = self._ocr_ollama_vision(pdf_path, page_index)
        return text

    # ── Vision description ───────────────────────────────────────────────────

    def _describe_diagram(self, pdf_path: str, page_index: int) -> str:
        """Ask Ollama llava to give a structured description of a diagram."""
        if not self._available or not PILLOW_AVAILABLE:
            return ""
        b64 = OllamaVisionClassifier._render_page_to_jpeg(pdf_path, page_index)
        if not b64:
            return ""
        prompt = (
            "This is a technical diagram from an engineering or software document. "
            "Describe it thoroughly: what type of diagram is it, what are the main components, "
            "what process or architecture does it represent, and what are the key relationships "
            "or data flows shown? Extract all visible text labels. "
            "Format your response as structured text suitable for retrieval."
        )
        try:
            resp = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": self.describe_model,
                    "prompt": prompt,
                    "images": [b64],
                    "stream": False,
                    "options": {"temperature": 0.0, "num_predict": 800},
                },
                timeout=VISION_API_TIMEOUT,
            )
            if resp.status_code == 200:
                return resp.json().get("response", "").strip()
        except Exception as e:
            logger.debug(f"Diagram description page {page_index+1}: {e}")
        return ""

    def _describe_image(self, pdf_path: str, page_index: int) -> str:
        """Ask Ollama llava to describe an image page."""
        if not self._available or not PILLOW_AVAILABLE:
            return ""
        b64 = OllamaVisionClassifier._render_page_to_jpeg(pdf_path, page_index)
        if not b64:
            return ""
        prompt = (
            "Describe this image from a technical document in detail. "
            "What does it show? Extract any text visible in the image. "
            "If it contains a chart, table, or graph, describe the data and labels precisely. "
            "Be comprehensive so the description can be used for information retrieval."
        )
        try:
            resp = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": self.describe_model,
                    "prompt": prompt,
                    "images": [b64],
                    "stream": False,
                    "options": {"temperature": 0.0, "num_predict": 600},
                },
                timeout=VISION_API_TIMEOUT,
            )
            if resp.status_code == 200:
                return resp.json().get("response", "").strip()
        except Exception as e:
            logger.debug(f"Image description page {page_index+1}: {e}")
        return ""

    # ── Table extraction ────────────────────────────────────────────────────

    # Minimum pixel size thresholds for meaningful raster images:
    #   _MIN_IMAGE_DIMENSION: each side must exceed 20 px to exclude 1×1
    #     tracking pixels and small PDF-form decoration glyphs.
    #   _MIN_IMAGE_AREA: total pixel area must exceed 1 000 px² to exclude
    #     narrow-but-tall or wide-but-short rule-line images.
    _MIN_IMAGE_DIMENSION: int = 20
    _MIN_IMAGE_AREA: int = 1000

    @staticmethod
    def _is_meaningful_image(img: dict) -> bool:
        """
        Return True if the pdfplumber image object represents a substantial
        visual element worth indexing as its own chunk.

        Tiny images (tracking pixels, small icons, PDF form decorations) are
        filtered out to avoid noise.  We check the source resolution in pixels
        because PDF point-space dimensions can be misleading.
        """
        try:
            # pdfplumber stores source resolution as 'srcsize': (width_px, height_px)
            src = img.get("srcsize") or ()
            w_px = src[0] if len(src) > 0 else img.get("width", 0)
            h_px = src[1] if len(src) > 1 else img.get("height", 0)
            min_d = ContentTypeExtractor._MIN_IMAGE_DIMENSION
            min_a = ContentTypeExtractor._MIN_IMAGE_AREA
            return bool(w_px and h_px
                        and w_px > min_d and h_px > min_d
                        and w_px * h_px > min_a)
        except Exception:
            return False

    @staticmethod
    def _has_vector_drawings(plumber_page) -> bool:
        """
        Return True when the page contains a non-trivial network of vector paths
        that is likely a diagram or flow chart.

        We look for:
        • ≥ 3 curve objects (Bezier / arc segments — essentially never produced
          by simple table borders or page frames).
        • ≥ 10 non-degenerate line segments (a dense grid of lines is a strong
          indicator of a flow chart, block diagram, or schematic).

        Degenerate lines (zero-length stubs) are excluded because pdfplumber
        sometimes reports them for invisible form-field anchors.
        """
        if plumber_page is None:
            return False
        try:
            if len(plumber_page.curves or []) >= 3:
                return True
            non_degenerate = [
                ln for ln in (plumber_page.lines or [])
                if ln.get("width", 0) != 0 or ln.get("height", 0) != 0
            ]
            return len(non_degenerate) >= 10
        except Exception:
            return False

    def _extract_images_as_text(
        self,
        pdf_path: str,
        page_index: int,
        plumber_page,
    ) -> List[str]:
        """
        For each meaningful visual element on a page, produce a tagged text
        block ([IMAGE N] or [DIAGRAM N]) so that detect_sections() creates a
        separate DocumentSection per visual element.

        Strategy
        --------
        1. Raster images (detected by pdfplumber as 'images'):
           • Each meaningful image gets its own [IMAGE N] block.
           • All images on the page share one full-page vision description
             (individual image cropping would require complex PDF rendering).
             A size annotation distinguishes each entry so the chunk is not
             a duplicate of its siblings.
           • After the image blocks, if the page also carries significant
             vector drawings (curves / dense line networks), a separate
             [DIAGRAM 1] block is emitted with the diagram-specific prompt so
             that vector-drawn flow charts are not silently dropped.
        2. Vector diagrams / figures (no raster images but the page was
           classified as having visual content):
           • Emit one [DIAGRAM 1] block with the vision description.
           • If no vision model is available, skip — we cannot describe vector
             art without one, and an empty chunk adds no retrieval value.

        Returns a list of tagged strings (one per visual element).
        """
        results: List[str] = []
        meaningful_imgs: List[dict] = []
        try:
            raw_imgs = plumber_page.images if plumber_page else []
            meaningful_imgs = [img for img in (raw_imgs or [])
                               if self._is_meaningful_image(img)]
        except Exception:
            pass

        if meaningful_imgs:
            # One full-page description is shared by all raster images because
            # _describe_image() renders the entire page — individual crops are
            # not feasible without complex PDF rendering.
            shared_desc = self._describe_image(pdf_path, page_index)
            for idx, img in enumerate(meaningful_imgs, 1):
                srcsize = img.get("srcsize", ("?", "?"))
                size_str = f"{srcsize[0]}×{srcsize[1]} px"
                if shared_desc:
                    # All images on the page share the same vision description;
                    # include a size annotation so each chunk is distinguishable.
                    content = (
                        f"{shared_desc}\n"
                        f"[Image {idx}/{len(meaningful_imgs)} on page "
                        f"{page_index + 1}, size {size_str}]"
                    )
                else:
                    content = (
                        f"[Embedded image {idx}/{len(meaningful_imgs)}, "
                        f"page {page_index + 1}, size {size_str}]"
                    )
                results.append(f"\n[IMAGE {idx}]\n{content}\n")

            # Also emit a DIAGRAM block when significant vector drawings exist
            # alongside the raster images — a page may contain both an embedded
            # bitmap and a vector-drawn flow chart, and the vector content would
            # otherwise be silently dropped.
            if self._has_vector_drawings(plumber_page):
                desc = self._describe_diagram(pdf_path, page_index)
                if desc:
                    results.append(f"\n[DIAGRAM 1]\n{desc}\n")
        else:
            # No raster images → the visual content is vector-drawn (shapes,
            # paths).  Only emit a DIAGRAM block when the vision model is
            # reachable (an empty block wastes storage and ranking budget).
            desc = self._describe_diagram(pdf_path, page_index)
            if desc:
                results.append(f"\n[DIAGRAM 1]\n{desc}\n")

        return results

    @staticmethod
    def _is_meaningful_table(table: list) -> bool:
        """
        Return True if the table has enough non-empty cells to be worth indexing.
        pdfplumber sometimes detects diagram or figure borders as table cells,
        producing nearly-empty tables.  We skip those.
        """
        if not table:
            return False
        total = sum(len(row) for row in table)
        if total == 0:
            return False
        non_empty = sum(
            1 for row in table for cell in row
            if cell and str(cell).strip()
        )
        return non_empty / total >= 0.20  # at least 20% of cells must have content

    @staticmethod
    def _extract_table_structured(plumber_page) -> str:
        """
        Extract tables from a pdfplumber page as structured pipe-delimited text,
        including header detection (first row of each table treated as header).
        Near-empty tables (likely diagram-border artifacts) are skipped.
        """
        if plumber_page is None:
            return ""
        parts: List[str] = []
        tbl_counter = 0
        try:
            tables = plumber_page.extract_tables()
            for table in tables:
                if not table or not ContentTypeExtractor._is_meaningful_table(table):
                    continue
                tbl_counter += 1
                rows = []
                for ri, row in enumerate(table):
                    cells = [str(c).strip() if c else "" for c in row]
                    row_str = " | ".join(cells)
                    if ri == 0:
                        # Mark first row as header
                        rows.append(row_str)
                        rows.append("-" * len(row_str))
                    else:
                        rows.append(row_str)
                parts.append(
                    f"\n[TABLE {tbl_counter}]\n" + "\n".join(rows) + "\n"
                )
        except Exception as e:
            logger.debug(f"Structured table extraction error: {e}")
        return "".join(parts)

    # ── Main dispatch ────────────────────────────────────────────────────────

    def extract(
        self,
        page_type: str,
        pdf_path: str,
        page_index: int,
        tier1_text: str,
        plumber_page,
    ) -> str:
        """
        Extract content from a page based on its classified type.
        Returns extracted text suitable for embedding.
        """
        t = page_type

        if t == PageType.TOC:
            return ""   # TOC pages are not indexed

        if t == PageType.COVER:
            return tier1_text.strip()

        if t in (PageType.REVISION_HISTORY, PageType.DISCLAIMER):
            return ""   # structural / boilerplate — not indexed

        if t == PageType.TEXT:
            return tier1_text.strip()

        if t == PageType.LIST:
            return tier1_text.strip()

        if t == PageType.GLOSSARY:
            # Return as-is; SemanticChunker will split on definition boundaries
            return tier1_text.strip()

        if t == PageType.SPECIFICATION_ITEM:
            # Preserve exact text — requirement IDs and traceability tags must not be altered
            return tier1_text.strip()

        if t == PageType.CODE_SNIPPET:
            # Return verbatim; no table/image extraction needed
            return tier1_text.strip()

        if t == PageType.CLASS_REFERENCE:
            # Prefer structured table extraction to preserve column semantics
            table_text = self._extract_table_structured(plumber_page)
            if table_text:
                return table_text
            # pdfplumber found no structured tables — fall back to tier1 text.
            # Strip control characters that confuse sent_tokenize on UML tables.
            cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', ' ', tier1_text)
            return cleaned.strip()

        if t == PageType.EQUATION:
            # Try text extraction first; fall back to vision description
            if tier1_text.strip():
                return tier1_text.strip()
            return self._describe_image(pdf_path, page_index)

        if t == PageType.TABLE:
            table_text = self._extract_table_structured(plumber_page)
            if table_text:
                return table_text
            # pdfplumber found no structured tables — try OCR (image-embedded table)
            if tier1_text.strip():
                return tier1_text.strip()
            ocr_text = self._ocr(pdf_path, page_index)
            if ocr_text.strip():
                return ocr_text.strip()
            # Last resort: vision description of the table layout
            return self._describe_image(pdf_path, page_index)

        if t == PageType.IMAGE:
            # OCR first; if OCR yields nothing, ask vision LLM to describe
            ocr_text = self._ocr(pdf_path, page_index)
            if ocr_text.strip():
                return ocr_text.strip()
            desc = self._describe_image(pdf_path, page_index)
            return desc or tier1_text.strip()

        if t == PageType.DIAGRAM:
            desc = self._describe_diagram(pdf_path, page_index)
            if desc:
                # Prepend any extracted text layer (labels, captions)
                if tier1_text.strip():
                    return f"{tier1_text.strip()}\n\n[Diagram Description]\n{desc}"
                return f"[Diagram Description]\n{desc}"
            # Fall back to OCR if vision is unavailable
            ocr_text = self._ocr(pdf_path, page_index)
            return ocr_text or tier1_text.strip()

        if t == PageType.MIXED:
            parts = []
            if tier1_text.strip():
                parts.append(tier1_text.strip())
            table_text = self._extract_table_structured(plumber_page)
            if table_text:
                parts.append(table_text)
            img_parts = self._extract_images_as_text(pdf_path, page_index, plumber_page)
            parts.extend(img_parts)
            return "\n\n".join(filter(None, parts))

        # Fallback
        return tier1_text.strip()

# ═══════════════════════════════════════════════════════════════════════════
# Helper utilities
# ═══════════════════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════════════════
# Enhanced PDF Loader (3-tier)
# ═══════════════════════════════════════════════════════════════════════════

class EnhancedPDFLoader:
    """
    Three-tier PDF text extraction:
    Tier 1 — pypdfium2 (fast, correct word spacing)
    Tier 2 — pdfplumber (better for some text-layer PDFs; used when Tier 1
              yields fewer than OCR_TRIGGER_CHARS per page)
    Tier 3 — pytesseract OCR (scanned PDFs; used when Tiers 1+2 both fail)
    Tables are extracted with pdfplumber and formatted as pipe-delimited text.
    """

    _TOC_LINE_RE = re.compile(r'(\.\s*){2,}.*\d+\s*$')

    @staticmethod
    def _is_toc_page(text: str) -> bool:
        stripped = text.strip()
        if len(stripped) < TOC_MIN_CONTENT_CHARS:
            return False   # too short to be a TOC page
        if len(stripped) > TOC_MAX_CONTENT_CHARS:
            return False
        lines = [l for l in stripped.splitlines() if l.strip()]
        if len(lines) < TOC_MIN_LINE_COUNT:
            return False
        toc = sum(
            1 for l in lines
            if EnhancedPDFLoader._TOC_LINE_RE.search(l)
        )
        return (toc / len(lines)) >= TOC_LINE_RATIO

    @staticmethod
    def _extract_tables_as_text(plumber_page) -> str:
        """
        Extract tables from a pdfplumber page as pipe-delimited text.
        Near-empty tables (pdfplumber diagram-border artifacts) are skipped.
        The first row of each table is treated as a header and followed by a
        separator line so that detect_sections() can recognize the header row.
        """
        parts: List[str] = []
        tbl_counter = 0
        try:
            tables = plumber_page.extract_tables()
            for table in tables:
                if not table or not ContentTypeExtractor._is_meaningful_table(table):
                    continue
                tbl_counter += 1
                rows = []
                for ri, row in enumerate(table):
                    cells = [str(c).strip() if c else "" for c in row]
                    row_str = " | ".join(cells)
                    if ri == 0:
                        rows.append(row_str)
                        rows.append("-" * len(row_str))
                    else:
                        rows.append(row_str)
                if rows:
                    parts.append(
                        f"\n[Table {tbl_counter}]\n" + "\n".join(rows) + "\n"
                    )
        except Exception as e:
            logger.debug(f"Table extraction error (non-fatal): {e}")
        return "".join(parts)

    @classmethod
    def load(
        cls,
        path: str,
        classifier: "OllamaVisionClassifier" = None,
        content_extractor: "ContentTypeExtractor" = None,
    ) -> Tuple[str, Dict]:
        """
        Load PDF using 3-tier extraction + optional AI page classification.

        When `classifier` is provided, every page is classified (text / table /
        image / diagram / mixed / cover / toc) and routed to the appropriate
        extraction strategy via `content_extractor`.

        page_type annotations are embedded in the text as
          [Page N | type=<type>]
        and also accumulated into metadata["page_types"] for downstream use.
        """
        metadata: Dict = {
            "num_pages": 0,
            "has_tables": False,
            "tables_count": 0,
            "page_types": {},     # {page_number: PageType}
        }
        page_texts: List[str] = []
        skipped_pages: List[int] = []
        type_counts: Dict[str, int] = {}

        # ── Tier 1: pypdfium2 ──────────────────────────────────────────
        pdf = pdfium.PdfDocument(path)
        tier1_results: Dict[int, str] = {}
        try:
            metadata["num_pages"] = len(pdf)
            for pn in range(len(pdf)):
                page = pdf[pn]
                try:
                    tp = page.get_textpage()
                    try:
                        tier1_results[pn] = tp.get_text_bounded()
                    finally:
                        tp.close()
                except Exception as e:
                    logger.warning(f"pypdfium2 page {pn+1} error: {e}")
                    tier1_results[pn] = ""
                finally:
                    page.close()
        finally:
            pdf.close()

        # ── Tier 2 + 3 with pdfplumber (also handles tables) ──────────
        plumber_pages: Dict[int, object] = {}
        plumber_doc = None
        try:
            plumber_doc = pdfplumber.open(path)
            for pn, pp in enumerate(plumber_doc.pages):
                plumber_pages[pn] = pp
        except Exception as e:
            logger.warning(f"pdfplumber open error (non-fatal): {e}")

        for pn in range(metadata["num_pages"]):
            tier1_text = tier1_results.get(pn, "")
            page_text = tier1_text

            # Tier 2: pdfplumber fallback for sparse pages
            if len(tier1_text.strip()) < OCR_TRIGGER_CHARS and pn in plumber_pages:
                try:
                    t2 = plumber_pages[pn].extract_text() or ""
                    if len(t2.strip()) > len(tier1_text.strip()):
                        page_text = t2
                        logger.debug(f"  Page {pn+1}: used pdfplumber (tier-2)")
                except Exception as e:
                    logger.debug(f"  pdfplumber tier-2 page {pn+1}: {e}")

            # ── AI Page Classification (new in V3) ─────────────────────
            page_type = PageType.TEXT
            class_result = None
            if classifier is not None:
                try:
                    class_result = classifier.classify(
                        path, pn, page_text, plumber_pages.get(pn)
                    )
                    page_type = class_result.page_type
                    type_counts[page_type] = type_counts.get(page_type, 0) + 1
                    metadata["page_types"][pn + 1] = page_type
                    logger.debug(
                        f"  Page {pn+1}: type={page_type} "
                        f"conf={class_result.confidence:.2f} "
                        f"vision={class_result.used_vision_model}"
                    )
                except Exception as e:
                    logger.debug(f"  Page {pn+1} classification error: {e}")

            # Skip structural/boilerplate pages — REMOVED.
            # All pages are now kept so the inspector can report on them.
            # Pages classified as TOC/COVER/etc. still get their type tag
            # embedded; the inspector decides how to report them.
            # Only truly empty pages (no text at all) get a no_text tag and
            # produce zero chunks — they are NOT silently dropped.

            # ── Content-type-aware extraction ──────────────────────────
            if content_extractor is not None and class_result is not None:
                try:
                    extracted = content_extractor.extract(
                        page_type, path, pn, page_text, plumber_pages.get(pn)
                    )
                    if extracted:
                        page_text = extracted
                except Exception as e:
                    logger.debug(f"  Page {pn+1} content extraction error: {e}")
            else:
                # Legacy Tier 3: OCR fallback for empty pages
                if (len(page_text.strip()) < OCR_TRIGGER_CHARS
                        and ENABLE_OCR_FALLBACK and OCR_AVAILABLE):
                    try:
                        images = pdf_to_images(
                            path, first_page=pn + 1, last_page=pn + 1, dpi=200
                        )
                        if images:
                            ocr_text = pytesseract.image_to_string(images[0])
                            if len(ocr_text.strip()) > len(page_text.strip()):
                                page_text = ocr_text
                                logger.debug(f"  Page {pn+1}: used OCR (tier-3)")
                    except Exception as e:
                        logger.debug(f"  OCR tier-3 page {pn+1}: {e}")

            if not page_text.strip():
                # Tag the page as no_text so it appears in the report
                # rather than being silently dropped.
                metadata["page_types"][pn + 1] = "no_text"
                page_texts.append(f"\n[Page {pn + 1} | type=no_text]\n")
                continue

            # Embed page type annotation in the text so downstream chunkers
            # can carry it forward as metadata.
            page_tag = f"[Page {pn + 1} | type={page_type}]"
            page_entry = f"\n{page_tag}\n{page_text}\n"

            # Append structured table text (pdfplumber) only when the
            # content_extractor has NOT already done table extraction for this
            # page.  The content_extractor handles TABLE and MIXED pages when
            # the classifier is active; double-appending would duplicate every
            # table on those pages.
            classifier_handled = (
                content_extractor is not None
                and class_result is not None
                and page_type in (PageType.TABLE, PageType.MIXED)
            )
            if not classifier_handled and page_type not in (
                PageType.TABLE, PageType.IMAGE, PageType.DIAGRAM
            ):
                if pn in plumber_pages:
                    table_text = cls._extract_tables_as_text(plumber_pages[pn])
                    if table_text:
                        metadata["has_tables"] = True
                        metadata["tables_count"] += table_text.count("[Table ")
                        page_entry += table_text

            page_texts.append(page_entry)

        if plumber_doc is not None:
            try:
                plumber_doc.close()
            except Exception:
                pass

        if skipped_pages:
            logger.info(
                f"  ⊘ Skipped {len(skipped_pages)} TOC/boilerplate pages: "
                f"{skipped_pages[:10]}{'...' if len(skipped_pages) > 10 else ''}"
            )

        if type_counts:
            logger.info(
                f"  📄 Page type breakdown: "
                + ", ".join(f"{k}={v}" for k, v in sorted(type_counts.items()))
            )

        return "".join(page_texts), metadata


# ═══════════════════════════════════════════════════════════════════════════
# Document loader (unified: PDF + DOCX + TXT)
# ═══════════════════════════════════════════════════════════════════════════

class AdvancedDocumentLoader:
    @staticmethod
    def extract_metadata(path: str) -> Dict:
        st = os.stat(path)
        return {
            "file_size_bytes": st.st_size,
            "created_timestamp": st.st_ctime,
            "modified_timestamp": st.st_mtime,
            "file_extension": Path(path).suffix.lower(),
        }

    @classmethod
    def load(
        cls,
        path: str,
        classifier: "OllamaVisionClassifier" = None,
        content_extractor: "ContentTypeExtractor" = None,
    ) -> Tuple[Optional[str], Dict]:
        try:
            base_meta = cls.extract_metadata(path)
            if path.lower().endswith(".pdf"):
                # Extract bookmark ToC before loading text
                toc_entries = ToCExtractor.extract(path)
                base_meta["toc_entries"] = toc_entries
                text, doc_meta = EnhancedPDFLoader.load(
                    path,
                    classifier=classifier,
                    content_extractor=content_extractor,
                )
            elif path.lower().endswith(".docx"):
                base_meta["toc_entries"] = []
                text, doc_meta = cls._load_docx(path)
            elif path.lower().endswith(".txt"):
                base_meta["toc_entries"] = []
                text, doc_meta = cls._load_txt(path)
            else:
                return None, {}
            base_meta.update(doc_meta)
            return text, base_meta
        except Exception as e:
            logger.error(f"Error loading {path}: {e}")
            return None, {}

    @staticmethod
    def _load_docx(path: str) -> Tuple[str, Dict]:
        doc = docx.Document(path)
        parts: List[str] = []
        meta: Dict = {"num_paragraphs": 0, "has_tables": False, "tables_count": 0}
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
                meta["num_paragraphs"] += 1
        if doc.tables:
            meta["has_tables"] = True
            meta["tables_count"] = len(doc.tables)
            for tbl_idx, table in enumerate(doc.tables, 1):
                rows = []
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    parts.append(
                        f"\n[Table {tbl_idx}]\n" + "\n".join(rows) + "\n"
                    )
        return "\n".join(parts), meta

    @staticmethod
    def _load_txt(path: str) -> Tuple[str, Dict]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return text, {"num_lines": len(text.splitlines()), "char_count": len(text)}


# ═══════════════════════════════════════════════════════════════════════════
# Semantic Chunker
# ═══════════════════════════════════════════════════════════════════════════

class SemanticChunker:
    """
    Splits text at semantic boundaries detected via TF-IDF cosine similarity
    of sliding sentence windows.  Falls back to pure sentence-boundary chunking
    when sklearn is unavailable or when there are too few sentences.
    """

    def __init__(self, chunk_size: int = CHILD_CHUNK_SIZE,
                 overlap: int = CHILD_CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.overlap = overlap

    # ── Section detection ────────────────────────────────────────────────────

    _section_patterns = [re.compile(p, re.MULTILINE) for p in SECTION_PATTERNS]
    # Regex for identifying inline table markers (case-insensitive)
    _TABLE_TITLE_RE = re.compile(r'^\[Table ', re.IGNORECASE)
    # Regex for identifying inline image / diagram markers (uppercase, as emitted
    # by _extract_images_as_text).  No IGNORECASE flag needed since the generator
    # always uses all-caps [IMAGE N] and [DIAGRAM N].
    _IMAGE_TITLE_RE  = re.compile(r'^\[IMAGE \d+\]$')
    _DIAGRAM_TITLE_RE = re.compile(r'^\[DIAGRAM \d+\]$')

    @staticmethod
    def _section_type_for_title(title: str) -> str:
        """Return the section_type string for the given section title marker.
        Uses PageType constants to ensure consistent type strings."""
        if re.match(r'^\[Table ', title, re.IGNORECASE):
            return PageType.TABLE
        if re.match(r'^\[IMAGE \d+\]$', title, re.IGNORECASE):
            return PageType.IMAGE
        if re.match(r'^\[DIAGRAM \d+\]$', title, re.IGNORECASE):
            return PageType.DIAGRAM
        return PageType.TEXT

    # Regex to parse page-type annotations injected by EnhancedPDFLoader
    _PAGE_TAG_RE = re.compile(r'^\[Page\s+(\d+)\s*\|\s*type=(\w+)\]$')

    def detect_sections(
        self,
        text: str,
        smart_detector: Optional["SmartSectionDetector"] = None,
    ) -> List[DocumentSection]:
        """
        Detect DocumentSection objects from *text*.

        When *smart_detector* is provided, heading lines are determined by the
        three-tier SmartSectionDetector (ToC → font → LLM) and injected into
        the text stream before the regex fallback runs.  This means:

          • Headings found by Tier 1/3 override the regex entirely for those pages.
          • SECTION_PATTERNS regex still runs as a safety net for lines that
            look like headings but weren't captured by the upper tiers.

        The method still supports being called without a smart_detector for
        backwards compatibility (global chunking path).
        """
        sections: List[DocumentSection] = []
        lines = text.split("\n")
        cur: Dict = {"title": "Document", "content": "", "level": 0}
        stack = [cur]
        current_page: Optional[int] = None

        # Build a set of confirmed heading texts from the smart detector for
        # this page so we can quickly check whether a line is a known heading.
        # Maps normalised line text → level.
        smart_headings: Dict[str, int] = {}

        def _refresh_smart_headings(page_num: int, page_body: str) -> None:
            """Populate smart_headings for the given page."""
            smart_headings.clear()
            if smart_detector is None or page_num is None:
                return
            for h_text, h_level in smart_detector.headings_for_page(page_num, page_body):
                smart_headings[h_text.strip().lower()] = h_level

        # We need to peek ahead to get the page body for the LLM call.
        # Pre-split pages so we can pass body text to smart detector.
        _page_tag_re = self._PAGE_TAG_RE
        _page_bodies: Dict[int, str] = {}
        _cur_pg: Optional[int] = None
        _cur_body: List[str] = []
        for raw_line in lines:
            ls = raw_line.strip()
            m  = _page_tag_re.match(ls)
            if m:
                if _cur_pg is not None:
                    _page_bodies[_cur_pg] = "\n".join(_cur_body)
                _cur_pg   = int(m.group(1))
                _cur_body = []
            else:
                _cur_body.append(raw_line)
        if _cur_pg is not None:
            _page_bodies[_cur_pg] = "\n".join(_cur_body)

        for line in lines:
            ls = line.strip()
            if not ls:
                cur["content"] += "\n"
                continue

            # Parse [Page N | type=...] tags to track page number
            page_match = self._PAGE_TAG_RE.match(ls)
            if page_match:
                current_page = int(page_match.group(1))
                # Refresh smart headings for this page
                _refresh_smart_headings(
                    current_page,
                    _page_bodies.get(current_page, ""),
                )
                continue

            # Strip AUTOSAR repeating page headers (version banner lines)
            if _AUTOSAR_RUNNING_HEADER_RE.match(ls):
                continue

            # ── Heading detection ─────────────────────────────────────────
            is_hdr = False
            level  = 0

            # Priority 1: Smart detector confirmed heading
            norm_ls = ls.lower()
            if norm_ls in smart_headings:
                is_hdr = True
                level  = smart_headings[norm_ls]

            # Priority 2: Injected markers ([TABLE N], [IMAGE N], [DIAGRAM N])
            # Always process these regardless of smart detector.
            if not is_hdr:
                if self._TABLE_TITLE_RE.match(ls) or \
                   self._IMAGE_TITLE_RE.match(ls) or \
                   self._DIAGRAM_TITLE_RE.match(ls):
                    is_hdr = True
                    level  = 4

            # Priority 3: SECTION_PATTERNS regex fallback
            if not is_hdr:
                for pat in self._section_patterns:
                    if pat.match(ls):
                        is_hdr = True
                        if ls.startswith("#"):
                            level = len(ls) - len(ls.lstrip("#"))
                        elif self._TABLE_TITLE_RE.match(ls):
                            level = 4
                        elif self._IMAGE_TITLE_RE.match(ls) or self._DIAGRAM_TITLE_RE.match(ls):
                            level = 4
                        elif ls.isupper():
                            level = 1
                        elif re.match(r'^\d+(\.\d+)*\s', ls):
                            num_part = ls.split()[0].rstrip('.')
                            level    = num_part.count('.') + 1
                        else:
                            level = 2
                        break

            if is_hdr and len(ls) < 200:
                if cur["content"].strip():
                    hier  = [s["title"] for s in stack if s["title"]]
                    stype = self._section_type_for_title(cur["title"])
                    sections.append(DocumentSection(
                        title=cur["title"],
                        content=cur["content"].strip(),
                        level=cur["level"],
                        page_number=cur.get("page_number"),
                        section_type=stype,
                        section_hierarchy=list(hier),
                    ))
                while len(stack) > 1 and stack[-1]["level"] >= level:
                    stack.pop()
                cur = {
                    "title":       ls.strip("#: ").strip(),
                    "content":     "",
                    "level":       level,
                    "page_number": current_page,
                }
                stack.append(cur)
            else:
                cur["content"] += line + "\n"

        if cur["content"].strip():
            hier  = [s["title"] for s in stack if s["title"]]
            stype = self._section_type_for_title(cur["title"])
            sections.append(DocumentSection(
                title=cur["title"],
                content=cur["content"].strip(),
                level=cur["level"],
                page_number=cur.get("page_number"),
                section_type=stype,
                section_hierarchy=list(hier),
            ))
        return sections or [DocumentSection("Document", text, 0)]


    # ── Semantic boundary detection ──────────────────────────────────────

    def _find_semantic_splits(self, sentences: List[str]) -> List[int]:
        """
        Return a list of sentence indices that mark the START of a new semantic
        segment.  Always includes 0.
        """
        splits = [0]
        n = len(sentences)
        if n < 2 * SEMANTIC_WINDOW_SIZE + 1 or not SKLEARN_AVAILABLE:
            return splits

        w = SEMANTIC_WINDOW_SIZE
        try:
            vectorizer = TfidfVectorizer(
                stop_words="english",
                max_features=5000,
                min_df=1,
            )
            # Build window texts
            windows = []
            for i in range(n):
                start = max(0, i - w)
                end = min(n, i + w + 1)
                windows.append(" ".join(sentences[start:end]))
            vectorizer.fit(windows)
            vecs = vectorizer.transform(windows).toarray()

            for i in range(1, n):
                prev = vecs[i - 1].reshape(1, -1)
                curr = vecs[i].reshape(1, -1)
                sim = float(sk_cosine(prev, curr)[0][0])
                if sim < SEMANTIC_SPLIT_THRESHOLD:
                    splits.append(i)
        except Exception as e:
            logger.debug(f"Semantic splitting fallback: {e}")

        return splits

    # ── Core chunking ────────────────────────────────────────────────────

    def _split_long_sentence(self, sentence: str,
                             max_size: int) -> List[str]:
        if len(sentence) <= max_size:
            return [sentence]
        words = sentence.split()
        parts: List[str] = []
        current = ""
        for w in words:
            if len(current) + len(w) + 1 > max_size and current:
                parts.append(current.strip())
                current = w
            else:
                current = current + " " + w if current else w
        if current.strip():
            parts.append(current.strip())
        return parts

    def _sentences_to_chunks(
        self,
        sentences: List[str],
        section_title: str,
        section_hierarchy: List[str],
        page_number: Optional[int],
        chunk_type: str,
        target_size: int,
        overlap_chars: int,
    ) -> List[ChildChunk]:
        """
        Pack sentences into chunks respecting semantic boundaries and target size.
        """
        if not sentences:
            return []

        # Optionally find semantic split points
        if ENABLE_SEMANTIC_CHUNKING:
            split_points = set(self._find_semantic_splits(sentences))
        else:
            split_points = set()

        chunks: List[ChildChunk] = []
        current_sents: List[str] = []
        current_len = 0
        start_char = 0
        char_pos = 0

        def _flush(sents: List[str], sc: int) -> ChildChunk:
            text = " ".join(sents).strip()
            return ChildChunk(
                text=text,
                enriched_text=text,   # filled later by ContextualEnricher
                section_title=section_title,
                section_hierarchy=list(section_hierarchy),
                page_number=page_number,
                chunk_type=chunk_type,
                page_type=chunk_type if chunk_type in {
                    PageType.TEXT, PageType.LIST, PageType.TABLE, PageType.IMAGE,
                    PageType.DIAGRAM, PageType.EQUATION, PageType.MIXED,
                    PageType.COVER, PageType.TOC, PageType.REVISION_HISTORY, PageType.DISCLAIMER,
                    PageType.CODE_SNIPPET, PageType.CLASS_REFERENCE,
                    PageType.SPECIFICATION_ITEM, PageType.GLOSSARY, PageType.UNKNOWN,
                } else PageType.TEXT,
                word_count=len(word_tokenize(text)),
                sentence_count=len(sent_tokenize(text)),
                start_char=sc,
                end_char=sc + len(text),
                parent_id="",         # filled by parent-child builder
                child_index=0,        # filled by parent-child builder
            )

        for idx, sent in enumerate(sentences):
            at_boundary = idx in split_points and current_sents

            if current_len + len(sent) > target_size and current_sents:
                # Flush current chunk
                chunks.append(_flush(current_sents, start_char))

                # Keep overlap
                overlap_sents: List[str] = []
                overlap_len = 0
                for s in reversed(current_sents):
                    if overlap_len + len(s) <= overlap_chars:
                        overlap_sents.insert(0, s)
                        overlap_len += len(s)
                    else:
                        break

                start_char = char_pos - overlap_len
                current_sents = overlap_sents
                current_len = overlap_len

            elif at_boundary and current_len >= MIN_CHUNK_SIZE:
                chunks.append(_flush(current_sents, start_char))
                start_char = char_pos
                current_sents = []
                current_len = 0

            current_sents.append(sent)
            char_pos += len(sent) + 1
            current_len += len(sent) + 1

        if current_sents and " ".join(current_sents).strip():
            chunks.append(_flush(current_sents, start_char))

        return chunks

    # Minimum number of dash characters required to recognise a separator line
    # (e.g. "------") that separates a table header row from its data rows.
    _TABLE_SEPARATOR_MIN_DASHES = 3

    def _chunk_table_rows(
        self,
        text: str,
        section_title: str,
        section_hierarchy: List[str],
        page_number: Optional[int],
        target_size: int,
    ) -> List[ChildChunk]:
        """
        Split a table section into child chunks by grouping complete rows.

        Instead of sentence-tokenising (which arbitrarily breaks pipe-delimited
        rows at '.' characters), this method treats each non-empty line as one
        indivisible unit (a table row or header separator).

        The header row + separator are detected and prepended to every chunk
        after the first so that each child chunk is self-contained.
        """
        lines = [l for l in text.splitlines() if l.strip()]
        if not lines:
            return []

        # Detect header block: first row + optional dashes-only separator line.
        header_lines: List[str] = [lines[0]]
        body_start = 1
        if len(lines) > 1 and re.match(r'^[-\s|+]+$', lines[1]) and lines[1].count("-") >= self._TABLE_SEPARATOR_MIN_DASHES:
            header_lines.append(lines[1])
            body_start = 2

        header_text = "\n".join(header_lines)
        data_rows = lines[body_start:]

        def _make_chunk(rows: List[str], sc: int) -> ChildChunk:
            body = "\n".join(rows)
            chunk_text = f"{header_text}\n{body}" if rows else header_text
            chunk_text = chunk_text.strip()
            if section_title:
                chunk_text = f"[{section_title}]\n{chunk_text}"
            return ChildChunk(
                text=chunk_text,
                enriched_text=chunk_text,   # filled later by ContextualEnricher
                section_title=section_title,
                section_hierarchy=list(section_hierarchy),
                page_number=page_number,
                chunk_type="table",
                page_type=PageType.TABLE,
                word_count=len(chunk_text.split()),
                # sentence_count reused to store row count for table chunks
                # (sentence_count is the field name defined by ChildChunk; we
                # document here that for table sections it counts data rows).
                sentence_count=len(rows) + len(header_lines),
                start_char=sc,
                end_char=sc + len(chunk_text),
                parent_id="",
                child_index=0,
            )

        # If table fits entirely within target_size, return as a single chunk.
        if len(text.strip()) <= target_size or not data_rows:
            return [_make_chunk(data_rows, 0)]

        # Otherwise split data rows into chunks of approximately target_size.
        chunks: List[ChildChunk] = []
        current_rows: List[str] = []
        current_len = len(header_text) + 1
        char_offset = 0

        for row in data_rows:
            row_len = len(row) + 1   # +1 for newline
            if current_len + row_len > target_size and current_rows:
                c = _make_chunk(current_rows, char_offset)
                chunks.append(c)
                char_offset += len(c.text) + 1
                current_rows = []
                current_len = len(header_text) + 1
            current_rows.append(row)
            current_len += row_len

        if current_rows:
            chunks.append(_make_chunk(current_rows, char_offset))

        return chunks or [_make_chunk(data_rows, 0)]

    def chunk_text(
        self,
        text: str,
        section_title: str = "",
        section_hierarchy: Optional[List[str]] = None,
        page_number: Optional[int] = None,
        chunk_type: str = "text",
        target_size: int = CHILD_CHUNK_SIZE,
        overlap: int = CHILD_CHUNK_OVERLAP,
    ) -> List[ChildChunk]:
        if not text or not text.strip():
            return []
        if section_hierarchy is None:
            section_hierarchy = [section_title] if section_title else ["Document"]

        # Short content that is below MIN_CHUNK_SIZE individually but still non-empty
        # (common on class_reference / specification_item pages with many short rows).
        # Emit it as a single atomic chunk rather than silently dropping it.
        if len(text.strip()) < MIN_CHUNK_SIZE:
            stripped = text.strip()
            return [ChildChunk(
                text=stripped,
                enriched_text=stripped,
                section_title=section_title,
                section_hierarchy=list(section_hierarchy),
                page_number=page_number,
                chunk_type=chunk_type,
                page_type=chunk_type if chunk_type in {
                    PageType.TEXT, PageType.LIST, PageType.TABLE, PageType.IMAGE,
                    PageType.DIAGRAM, PageType.EQUATION, PageType.MIXED,
                    PageType.COVER, PageType.TOC, PageType.REVISION_HISTORY, PageType.DISCLAIMER,
                    PageType.CODE_SNIPPET, PageType.CLASS_REFERENCE,
                    PageType.SPECIFICATION_ITEM, PageType.GLOSSARY, PageType.UNKNOWN,
                } else PageType.TEXT,
                word_count=len(stripped.split()),
                sentence_count=1,
                start_char=0,
                end_char=len(stripped),
                parent_id="",
                child_index=0,
            )]

        # Table sections must not go through sent_tokenize: rows would be split
        # at '.' characters inside cell values.  Use row-based chunking instead.
        if chunk_type == "table":
            return self._chunk_table_rows(
                text, section_title, section_hierarchy, page_number, target_size
            )

        raw_sents = sent_tokenize(text)
        sentences: List[str] = []
        for s in raw_sents:
            sentences.extend(self._split_long_sentence(s, target_size))

        return self._sentences_to_chunks(
            sentences, section_title, section_hierarchy,
            page_number, chunk_type, target_size, overlap,
        )

    def chunk_sections(
        self,
        sections: List[DocumentSection],
        target_size: int = CHILD_CHUNK_SIZE,
        overlap: int = CHILD_CHUNK_OVERLAP,
    ) -> List[ChildChunk]:
        result: List[ChildChunk] = []
        for sec in sections:
            hier = sec.section_hierarchy or [sec.title]
            # Prepend the section heading to the content so that chunk.text
            # starts at the heading line.  This means _find_bbox_for_text
            # will locate the heading in the PDF word stream and the bounding
            # box will correctly enclose the heading + body instead of
            # starting below the heading.
            # Skip prepending for the synthetic root "Document" title and for
            # injected visual markers ([TABLE N], [IMAGE N], [DIAGRAM N]) which
            # are never present in the PDF word stream.
            skip_title = (
                not sec.title
                or sec.title == "Document"
                or re.match(r'^\[(TABLE|IMAGE|DIAGRAM)\s+\d+\]$', sec.title, re.IGNORECASE)
            )
            full_content = (
                sec.content
                if skip_title
                else f"{sec.title}\n{sec.content}"
            )
            chunks = self.chunk_text(
                full_content, sec.title, hier,
                sec.page_number, sec.section_type, target_size, overlap,
            )
            result.extend(chunks)
        return result


# ═══════════════════════════════════════════════════════════════════════════
# Parent-Child Builder
# ═══════════════════════════════════════════════════════════════════════════

class ParentChildBuilder:
    """
    Builds parent chunks and their child chunks from document sections.

    Workflow:
    1. Chunk the document at PARENT_CHUNK_SIZE → ParentChunk list
    2. For each parent, further chunk at CHILD_CHUNK_SIZE → ChildChunk list
    3. Set child.parent_id = parent.parent_id
    """

    def __init__(self, chunker: SemanticChunker):
        self.chunker = chunker

    def build(
        self,
        sections: List[DocumentSection],
        file_hash_value: str = "",
    ) -> Tuple[List[ParentChunk], List[ChildChunk]]:
        parents: List[ParentChunk] = []
        children: List[ChildChunk] = []

        # Build large parent chunks first
        raw_parents: List[ChildChunk] = self.chunker.chunk_sections(
            sections,
            target_size=PARENT_CHUNK_SIZE,
            overlap=PARENT_CHUNK_OVERLAP,
        )

        for raw_parent in raw_parents:
            pid = str(uuid.uuid5(
                uuid.NAMESPACE_DNS,
                # Include file_hash to prevent collisions across documents
                # that happen to share the same 64-char prefix at the same offset.
                f"{file_hash_value}_{raw_parent.start_char}_{raw_parent.text[:64]}"
            ))
            parent = ParentChunk(
                text=raw_parent.text,
                section_title=raw_parent.section_title,
                section_hierarchy=list(raw_parent.section_hierarchy),
                page_number=raw_parent.page_number,
                chunk_type=raw_parent.chunk_type,
                page_type=getattr(raw_parent, "page_type", PageType.TEXT),
                word_count=raw_parent.word_count,
                start_char=raw_parent.start_char,
                end_char=raw_parent.end_char,
                parent_id=pid,
            )
            parents.append(parent)

            # Now chunk the parent content into children
            child_raw = self.chunker.chunk_text(
                raw_parent.text,
                raw_parent.section_title,
                raw_parent.section_hierarchy,
                raw_parent.page_number,
                raw_parent.chunk_type,
                CHILD_CHUNK_SIZE,
                CHILD_CHUNK_OVERLAP,
            )

            for ci, child in enumerate(child_raw):
                child.parent_id = pid
                child.child_index = ci
                children.append(child)

        return parents, children


# ═══════════════════════════════════════════════════════════════════════════
# Contextual Enricher (Anthropic research)
# ═══════════════════════════════════════════════════════════════════════════


# ── End of inlined definitions ────────────────────────────────────────────────


# ═══════════════════════════════════════════════════════════════════════════
# Layout Analyser  (PyMuPDF-based spatial detection)
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class LayoutRegion:
    """One spatially-detected region on a PDF page."""
    kind:    str                              # "diagram" | "line_table" | "image"
    bbox:    Tuple[float, float, float, float]  # PDF pts (x0, y0, x1, y1)
    caption: str = ""
    label:   str = ""    # e.g. "Figure 2.1", "Table 3.10" if found near region


class LayoutAnalyser:
    """
    Performs spatial layout analysis using PyMuPDF (fitz).

    PyMuPDF exposes three object streams that pdfplumber / pypdfium2 miss:
      page.get_drawings()    → every vector path (lines, curves, shapes)
      page.get_images()      → raster image blocks with coordinates
      page.get_text("dict")  → text blocks with precise coordinates

    This class uses those streams to detect:

    1. LINE-BASED TABLES
       AUTOSAR tables are drawn as a grid of horizontal + vertical lines.
       pdfplumber.extract_tables() uses these lines internally but its
       heuristic sometimes splits one logical table into multiple objects
       or misses continuation pages.  We re-detect the grid independently
       so the viz can overlay the correct bounding box.

    2. VECTOR DIAGRAMS
       UML diagrams, flow charts, block diagrams (Figure 2.1, 2.2, …) are
       stored as vector paths — no raster pixels, just draw commands.
       They are completely invisible to pdfplumber word extraction.
       We cluster nearby drawing paths into diagram regions.

    3. RASTER IMAGES
       Embedded JPEG/PNG bitmaps detected via PyMuPDF's image info API.

    For each detected region the class also looks for a caption line
    directly below it (text starting with "Figure", "Table", or similar).
    """

    # ── Tuning constants ─────────────────────────────────────────────────
    # Minimum area (PDF pts²) for a drawing cluster to be a diagram
    MIN_DIAGRAM_AREA     : float = 1500.0
    # Minimum number of drawing objects in a cluster to call it a diagram
    MIN_DIAGRAM_OBJECTS  : int   = 3
    # Maximum stroke width for a path to be treated as a table border line
    TABLE_LINE_MAX_WIDTH : float = 2.0
    # A line is "table-like" when its aspect ratio is extreme (W/H or H/W >= this)
    TABLE_LINE_ASPECT    : float = 8.0
    # Minimum number of lines in a cluster to call it a table
    TABLE_MIN_LINES      : int   = 6
    # Proximity gap for clustering drawing objects (PDF pts)
    CLUSTER_GAP          : float = 24.0
    # Height of the caption search band below / above a region (pts)
    CAPTION_BAND         : float = 55.0
    # Caption prefixes that identify a figure or table label
    _CAPTION_RE = re.compile(
        r'^(figure|fig\.?|table|tbl\.?|diagram)\s*[\d.]',
        re.IGNORECASE,
    )

    # ── Public entry point ────────────────────────────────────────────────

    @classmethod
    def analyse(cls, fitz_page) -> List[LayoutRegion]:
        """
        Return all detected regions for one PyMuPDF page object.
        Empty list when PyMuPDF is unavailable or the page has no drawings.
        """
        if not PYMUPDF_AVAILABLE or fitz_page is None:
            return []

        try:
            drawings = fitz_page.get_drawings() or []
        except Exception:
            drawings = []

        table_paths, diagram_paths = cls._split_paths(drawings)

        regions: List[LayoutRegion] = []
        table_bboxes: List[Tuple] = []

        # Detect line-based tables first so diagram clustering can exclude them
        for r in cls._detect_line_tables(table_paths, fitz_page):
            regions.append(r)
            table_bboxes.append(r.bbox)

        # Detect vector diagram clusters
        for r in cls._detect_diagrams(diagram_paths, table_bboxes, fitz_page):
            regions.append(r)

        # Detect raster images
        for r in cls._detect_raster_images(fitz_page):
            regions.append(r)

        return regions

    # ── Path splitting ────────────────────────────────────────────────────

    @classmethod
    def _split_paths(cls, drawings):
        """Separate table-border lines from diagram drawing paths."""
        table_paths, diagram_paths = [], []
        for d in drawings:
            rect = d.get("rect")
            if rect is None:
                continue
            x0, y0, x1, y1 = float(rect[0]), float(rect[1]), float(rect[2]), float(rect[3])
            w = x1 - x0
            h = y1 - y0
            stroke_w = float(d.get("width") or 1.0)

            # Classify as a table-border line:
            #   thin in one dimension AND stroke width <= threshold
            is_h_line = h < 3.0 and w > 5.0
            is_v_line = w < 3.0 and h > 5.0
            if (is_h_line or is_v_line) and stroke_w <= cls.TABLE_LINE_MAX_WIDTH:
                table_paths.append((x0, y0, x1, y1))
            else:
                diagram_paths.append((x0, y0, x1, y1))
        return table_paths, diagram_paths

    # ── Line-based table detection ────────────────────────────────────────

    @classmethod
    def _detect_line_tables(cls, table_paths, fitz_page) -> List[LayoutRegion]:
        if not table_paths:
            return []

        clusters = cls._cluster_rects(table_paths, gap=cls.CLUSTER_GAP * 2)
        regions  = []
        for cluster in clusters:
            if len(cluster) < cls.TABLE_MIN_LINES:
                continue
            x0 = min(r[0] for r in cluster)
            y0 = min(r[1] for r in cluster)
            x1 = max(r[2] for r in cluster)
            y1 = max(r[3] for r in cluster)
            area = (x1 - x0) * (y1 - y0)
            if area < cls.MIN_DIAGRAM_AREA * 0.4:
                continue

            # Verify that there is actual text inside this region
            try:
                inner_text = fitz_page.get_text("text", clip=(x0, y0, x1, y1)).strip()
            except Exception:
                inner_text = "x"   # assume text present if query fails
            if not inner_text:
                continue

            caption = cls._find_caption_near(fitz_page, x0, y0, x1, y1)
            label   = cls._extract_label(caption)
            regions.append(LayoutRegion(
                kind="line_table",
                bbox=(x0, y0, x1, y1),
                caption=caption,
                label=label or "Table",
            ))
        return regions

    # ── Vector diagram detection ──────────────────────────────────────────

    @classmethod
    def _detect_diagrams(cls, diagram_paths, table_bboxes, fitz_page) -> List[LayoutRegion]:
        if not diagram_paths:
            return []

        # Exclude paths that fall entirely inside a detected table region
        filtered = []
        for r in diagram_paths:
            inside_table = any(
                r[0] >= tb[0] - 5 and r[1] >= tb[1] - 5
                and r[2] <= tb[2] + 5 and r[3] <= tb[3] + 5
                for tb in table_bboxes
            )
            if not inside_table:
                filtered.append(r)

        if not filtered:
            return []

        clusters = cls._cluster_rects(filtered, gap=cls.CLUSTER_GAP)
        regions  = []
        for cluster in clusters:
            if len(cluster) < cls.MIN_DIAGRAM_OBJECTS:
                continue
            x0 = min(r[0] for r in cluster)
            y0 = min(r[1] for r in cluster)
            x1 = max(r[2] for r in cluster)
            y1 = max(r[3] for r in cluster)
            area = (x1 - x0) * (y1 - y0)
            if area < cls.MIN_DIAGRAM_AREA:
                continue

            caption = cls._find_caption_near(fitz_page, x0, y0, x1, y1)
            label   = cls._extract_label(caption)
            regions.append(LayoutRegion(
                kind="diagram",
                bbox=(x0, y0, x1, y1),
                caption=caption,
                label=label or "Diagram",
            ))
        return regions

    # ── Raster image detection ────────────────────────────────────────────

    @classmethod
    def _detect_raster_images(cls, fitz_page) -> List[LayoutRegion]:
        regions = []
        try:
            for info in fitz_page.get_image_info(hashes=False):
                bbox = info.get("bbox")
                if not bbox:
                    continue
                x0, y0, x1, y1 = float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])
                if (x1 - x0) * (y1 - y0) < 400:   # skip tiny icons / bullets
                    continue
                caption = cls._find_caption_near(fitz_page, x0, y0, x1, y1)
                label   = cls._extract_label(caption)
                regions.append(LayoutRegion(
                    kind="image",
                    bbox=(x0, y0, x1, y1),
                    caption=caption,
                    label=label or "Image",
                ))
        except Exception:
            pass
        return regions

    # ── Caption finder ────────────────────────────────────────────────────

    @classmethod
    def _find_caption_near(cls, fitz_page, x0, y0, x1, y1) -> str:
        """Search for a caption line above or below a region."""
        pad = 10.0
        try:
            # Below the region
            below_clip = (x0 - pad, y1, x1 + pad, y1 + cls.CAPTION_BAND)
            words_below = fitz_page.get_text("words", clip=below_clip) or []
            if words_below:
                text = " ".join(w[4] for w in words_below[:30])
                if cls._CAPTION_RE.match(text.strip()):
                    return text.strip()
                # Even if it doesn't start with "Figure", collect it
                if text.strip():
                    return text.strip()[:120]

            # Above the region
            above_clip = (x0 - pad, y0 - cls.CAPTION_BAND, x1 + pad, y0)
            words_above = fitz_page.get_text("words", clip=above_clip) or []
            if words_above:
                text = " ".join(w[4] for w in words_above[:30])
                if cls._CAPTION_RE.match(text.strip()):
                    return text.strip()
        except Exception:
            pass
        return ""

    @staticmethod
    def _extract_label(caption: str) -> str:
        """Pull a 'Figure N.N' / 'Table N.N' label from a caption string."""
        m = re.match(
            r'((?:figure|fig\.?|table|tbl\.?|diagram)\s*[\d.]+)',
            caption, re.IGNORECASE,
        )
        return m.group(1).strip() if m else ""

    # ── Rectangle clustering ──────────────────────────────────────────────

    @classmethod
    def _cluster_rects(
        cls,
        rects: List[Tuple[float, float, float, float]],
        gap:   float = 20.0,
    ) -> List[List[Tuple[float, float, float, float]]]:
        """
        Single-linkage clustering: two rects belong to the same cluster when
        their padded bounding boxes overlap or are within *gap* pts of each other.
        """
        if not rects:
            return []

        # Use union-find for efficiency
        parent = list(range(len(rects)))

        def find(i):
            while parent[i] != i:
                parent[i] = parent[parent[i]]
                i = parent[i]
            return i

        def union(i, j):
            parent[find(i)] = find(j)

        for i in range(len(rects)):
            for j in range(i + 1, len(rects)):
                if cls._rects_near(rects[i], rects[j], gap):
                    union(i, j)

        from collections import defaultdict
        groups: Dict[int, List] = defaultdict(list)
        for i, r in enumerate(rects):
            groups[find(i)].append(r)

        return list(groups.values())

    @staticmethod
    def _rects_near(
        r1: Tuple[float, float, float, float],
        r2: Tuple[float, float, float, float],
        gap: float,
    ) -> bool:
        """True when two rects overlap or are separated by <= gap pts."""
        h_gap = max(0.0, max(r1[0], r2[0]) - min(r1[2], r2[2]))
        v_gap = max(0.0, max(r1[1], r2[1]) - min(r1[3], r2[3]))
        return h_gap <= gap and v_gap <= gap


# ═══════════════════════════════════════════════════════════════════════════
# Side-by-side page visualiser
# ═══════════════════════════════════════════════════════════════════════════

# ── Color palette (RGB) ──────────────────────────────────────────────────────
# Eight distinct parent colors — children of the same parent use a lighter tint.
_VIZ_PARENT_COLORS: List[Tuple[int, int, int]] = [
    (220,  53,  69),   # red
    (  0, 112, 220),   # blue
    ( 40, 167,  69),   # green
    (230, 126,  34),   # orange
    (111,  66, 193),   # purple
    ( 23, 162, 184),   # cyan
    (200, 173,  12),   # gold
    ( 32, 180, 135),   # teal
]

# Section divider color (thin horizontal line between sections)
_VIZ_SECTION_COLOR: Tuple[int, int, int] = (160, 160, 160)

# Caption bar height in pixels
_VIZ_CAPTION_H: int = 48

# Padding around bounding boxes (pixels at scale 3.0)
_VIZ_BOX_PAD: int = 6

# Colors for spatially-detected layout regions (RGB)
_VIZ_DIAGRAM_COLOR  : Tuple[int, int, int] = (255, 165,   0)   # orange — vector diagrams
_VIZ_TABLE_COLOR    : Tuple[int, int, int] = ( 70, 130, 180)   # steel-blue — line-based tables
_VIZ_IMAGE_COLOR    : Tuple[int, int, int] = ( 34, 139,  34)   # forest-green — raster images
# Alpha for layout region fills
_VIZ_REGION_ALPHA   : int = 35

# Default render scale for viz (can be overridden by --viz-scale)
VIZ_DEFAULT_SCALE: float = 3.0


def _tint(color: Tuple[int, int, int], alpha: float = 0.35) -> Tuple[int, int, int, int]:
    """Return an RGBA tuple: full color at `alpha` opacity for semi-transparent fills."""
    return (color[0], color[1], color[2], int(255 * alpha))


def _lighter(color: Tuple[int, int, int], factor: float = 0.55) -> Tuple[int, int, int]:
    """Mix color toward white by `factor` to produce a child-chunk tint."""
    return (
        int(color[0] + (255 - color[0]) * factor),
        int(color[1] + (255 - color[1]) * factor),
        int(color[2] + (255 - color[2]) * factor),
    )


def _words_to_charmap(
    words: List[Dict],
) -> Tuple[str, List[Tuple[int, int, float, float, float, float]]]:
    """
    Convert a list of pdfplumber word dicts into a normalised text string
    and a parallel character map.

    Returns: (norm_text, charmap)
      norm_text — space-joined, lower-cased, whitespace-collapsed word stream
      charmap   — list of (char_start, char_end, x0, top, x1, bottom)
                  where char_start/end are byte offsets into norm_text.

    Both the string and the offsets use the SAME normalisation so that
    _find_bbox_for_text can search norm_text and index directly into charmap
    without any offset skew.
    """
    norm_parts: List[str] = []
    charmap: List[Tuple[int, int, float, float, float, float]] = []
    pos = 0
    for w in words:
        raw = w.get("text", "")
        if not raw:
            continue
        norm = re.sub(r'\s+', ' ', raw.strip().lower())
        if not norm:
            continue
        end = pos + len(norm)
        charmap.append((pos, end, w["x0"], w["top"], w["x1"], w["bottom"]))
        norm_parts.append(norm)
        pos = end + 1   # +1 for the space separator

    norm_text = " ".join(norm_parts)
    return norm_text, charmap


def _clean_chunk_for_search(text: str) -> str:
    """
    Normalise chunk text so it can be matched against the pdfplumber word
    stream (built from extract_words(), which has no pipe separators or
    injected markers).

    Removes / normalises:
      - [[TABLE N]] / [TABLE N] / [IMAGE N] / [DIAGRAM N] prefix markers
      - Pipe cell separators from table rows  (extract_tables uses ' | ',
        extract_words produces plain space-separated words — they must agree)
      - Dash-only separator lines  (e.g. "------" between header and data rows)
      - Leading AUTOSAR running-header lines
    """
    text = text.strip()

    # ── Strip injected visual-content markers ─────────────────────────────
    text = re.sub(
        r'^\[?\[(?:TABLE|IMAGE|DIAGRAM)\s+\d+\]?\]\s*',
        '', text, flags=re.IGNORECASE,
    ).strip()

    # ── Strip pipe cell separators from table rows ────────────────────────
    # extract_tables() joins cells with " | "; extract_words() has no pipes.
    # Replace every " | " with a space so the search key matches the word
    # stream.  Also strip a leading/trailing pipe that some rows have.
    if '|' in text:
        # Process line-by-line to also remove dashes-only separator lines
        cleaned_lines = []
        for line in text.splitlines():
            stripped = line.strip()
            # Drop pure separator lines like "---", "---|---", "---+---"
            if stripped and re.match(r'^[-|\s+]+$', stripped) and stripped.count('-') >= 3:
                continue
            # Strip pipes from data/header rows
            no_pipes = re.sub(r'\s*\|\s*', ' ', stripped).strip()
            if no_pipes:
                cleaned_lines.append(no_pipes)
        text = '\n'.join(cleaned_lines).strip()

    if not text:
        return ''

    # ── Strip leading AUTOSAR running-header lines ────────────────────────
    lines      = text.splitlines()
    drop_until = 0
    for i, line in enumerate(lines):
        ls = line.strip()
        if not ls:
            drop_until = i + 1
            continue
        if _AUTOSAR_RUNNING_HEADER_RE.match(ls):
            drop_until = i + 1
            continue
        break

    result = '\n'.join(lines[drop_until:]).strip()
    return result if result else text


def _find_bbox_for_text(
    chunk_text: str,
    norm_page: str,                                               # pre-normalised word-stream text
    charmap: List[Tuple[int, int, float, float, float, float]],  # aligned with norm_page
    pdf_w: float,
    pdf_h: float,
    img_w: int,
    img_h: int,
    norm_page_alt: Optional[str] = None,  # fallback: pdfplumber extract_text() normalised
) -> Optional[Tuple[int, int, int, int]]:
    """
    Locate *chunk_text* inside *norm_page* (the normalised pdfplumber word
    stream) and return a pixel bounding box (left, top, right, bottom).

    Fixes vs. the original:
      1. norm_page and charmap share the same offsets (built by _words_to_charmap).
      2. Injected chunk markers and repeating headers are stripped before searching.
      3. Chunk tail is used to find the true end position on the page.
      4. Multi-attempt search: first 60 chars, then first 30, then first 15.
      5. Falls back to norm_page_alt (pdfplumber extract_text) when word-stream
         search fails — handles pypdfium2 vs pdfplumber ligature differences.
         The charmap still comes from extract_words() so coordinates are correct.
    """
    if not charmap or not chunk_text.strip():
        return None

    sx = img_w / pdf_w
    sy = img_h / pdf_h

    def _norm(s: str) -> str:
        return re.sub(r'\s+', ' ', s.strip().lower())

    clean = _clean_chunk_for_search(chunk_text)
    nc    = _norm(clean)
    if not nc:
        return None

    # For table chunks the search space may still have stray pipes if it was
    # built from extract_text() (the alt source).  Normalise both sides.
    is_table_chunk = '|' in chunk_text
    if is_table_chunk:
        norm_page_search     = re.sub(r'\s*\|\s*', ' ', norm_page)
        norm_page_search     = re.sub(r'\s+', ' ', norm_page_search).strip()
        norm_page_alt_search = (
            re.sub(r'\s+', ' ', re.sub(r'\s*\|\s*', ' ', norm_page_alt)).strip()
            if norm_page_alt else norm_page_alt
        )
    else:
        norm_page_search     = norm_page
        norm_page_alt_search = norm_page_alt

    # ── Find start position — try word-stream first, then alt ────────────
    def _try_find(search_space: str) -> int:
        for prefix_len in (60, 40, 30, 20, 15):
            if len(nc) >= prefix_len:
                hit = search_space.find(nc[:prefix_len])
                if hit != -1:
                    return hit
        return -1

    idx = _try_find(norm_page_search)
    # Fallback: try pdfplumber extract_text() source
    if idx == -1 and norm_page_alt_search and norm_page_alt_search != norm_page_search:
        idx = _try_find(norm_page_alt_search)
        if idx != -1:
            ratio = idx / max(len(norm_page_alt_search), 1)
            idx   = int(ratio * len(norm_page_search))

    if idx == -1:
        return None

    # If we used a pipe-stripped search space, remap idx proportionally back
    # to the original norm_page offsets (which align with charmap).
    if is_table_chunk and norm_page_search is not norm_page:
        ratio = idx / max(len(norm_page_search), 1)
        idx   = int(ratio * len(norm_page))

    # ── Find end position using chunk tail ────────────────────────────────
    tail     = _norm(clean[-80:]) if len(clean) > 80 else nc
    tail_key = tail[-40:] if len(tail) >= 40 else tail
    if is_table_chunk:
        tail_key = re.sub(r'\s*\|\s*', ' ', tail_key)
        tail_key = re.sub(r'\s+', ' ', tail_key).strip()

    end_candidate = norm_page_search.find(tail_key, idx)
    if end_candidate != -1:
        raw_end = end_candidate + len(tail_key)
        if is_table_chunk and norm_page_search is not norm_page:
            ratio   = raw_end / max(len(norm_page_search), 1)
            end_idx = int(ratio * len(norm_page))
        else:
            end_idx = raw_end
    else:
        end_idx = idx + len(nc)

    # ── Collect words from original charmap (no pipes — always safe) ──────
    x0s, tops, x1s, bots = [], [], [], []
    for (cs, ce, wx0, wtop, wx1, wbot) in charmap:
        if ce > idx and cs < end_idx:
            x0s.append(wx0)
            tops.append(wtop)
            x1s.append(wx1)
            bots.append(wbot)

    if not x0s:
        return None

    # ── Convert PDF points → pixels with padding ──────────────────────────
    pad    = _VIZ_BOX_PAD
    left   = max(0,           int(min(x0s)  * sx) - pad)
    top    = max(0,           int(min(tops) * sy) - pad)
    right  = min(img_w - 1,  int(max(x1s)  * sx) + pad)
    bottom = min(img_h - 1,  int(max(bots) * sy) + pad)
    return (left, top, right, bottom)


def _draw_box(
    draw: "PILImageDraw.ImageDraw",
    overlay: "PILImage.Image",
    bbox: Tuple[int, int, int, int],
    color: Tuple[int, int, int],
    label: str,
    line_width: int = 3,
    font=None,
) -> None:
    """Draw a filled semi-transparent rectangle + solid border + label."""
    left, top, right, bottom = bbox

    # Semi-transparent fill on the overlay layer
    fill_layer = PILImage.new("RGBA", overlay.size, (0, 0, 0, 0))
    fd = PILImageDraw.Draw(fill_layer)
    fd.rectangle([left, top, right, bottom], fill=_tint(color, 0.22))
    overlay.alpha_composite(fill_layer)

    # Solid border
    draw.rectangle([left, top, right, bottom],
                   outline=color, width=line_width)

    # Label badge (small filled rect + text)
    if label:
        try:
            tw = draw.textlength(label, font=font)
        except Exception:
            tw = len(label) * 7
        th = 14
        bx0 = left
        by0 = max(0, top - th - 2)
        bx1 = bx0 + int(tw) + 6
        by1 = by0 + th + 2
        draw.rectangle([bx0, by0, bx1, by1], fill=color)
        draw.text((bx0 + 3, by0 + 1), label, fill=(255, 255, 255), font=font)


def _add_caption(
    img: "PILImage.Image",
    text: str,
    bg: Tuple[int, int, int] = (30, 30, 30),
    fg: Tuple[int, int, int] = (240, 240, 240),
) -> "PILImage.Image":
    """Append a dark caption strip below `img`."""
    cap = PILImage.new("RGB", (img.width, _VIZ_CAPTION_H), bg)
    d   = PILImageDraw.Draw(cap)
    # Center the text vertically
    try:
        tw = d.textlength(text)
    except Exception:
        tw = len(text) * 7
    x = max(4, (img.width - int(tw)) // 2)
    d.text((x, (_VIZ_CAPTION_H - 14) // 2), text, fill=fg)

    combined = PILImage.new("RGB", (img.width, img.height + _VIZ_CAPTION_H))
    combined.paste(img, (0, 0))
    combined.paste(cap, (0, img.height))
    return combined


def build_viz_page(
    pdf_path: str,
    page_index: int,          # 0-based
    page_num: int,            # 1-based (for labels)
    page_type: str,
    plumber_page,             # pdfplumber page object
    sections: List,           # List[DocumentSection]
    parents: List,            # List[ParentChunk]
    children: List,           # List[ChildChunk]
    draw_children: bool = True,
    scale: float = VIZ_DEFAULT_SCALE,
    fitz_page=None,           # PyMuPDF page object (optional, for layout analysis)
) -> Optional["PILImage.Image"]:
    """
    Render a side-by-side comparison image for one PDF page.

    Left  — clean page render
    Right — same render annotated with:
              • Layout regions (diagrams, tables, images) from PyMuPDF spatial
                analysis — drawn as semi-transparent colored overlays
              • Chunk text bounding boxes — located via word-stream text search
              • Section divider lines

    Returns a PIL Image (RGB), or None if rendering is unavailable.
    """
    if not PILLOW_AVAILABLE or plumber_page is None:
        return None

    # ── 1. Render the page ───────────────────────────────────────────────
    try:
        doc = pdfium.PdfDocument(pdf_path)
        try:
            pdfium_page = doc[page_index]
            bitmap      = pdfium_page.render(scale=scale)
            raw_img     = bitmap.to_pil().convert("RGB")
            pdfium_page.close()
        finally:
            doc.close()
    except Exception:
        return None

    img_w, img_h = raw_img.size
    pdf_w = float(plumber_page.width)
    pdf_h = float(plumber_page.height)
    sx    = img_w / pdf_w
    sy    = img_h / pdf_h

    # ── 2. Spatial layout analysis via PyMuPDF ────────────────────────────
    layout_regions: List[LayoutRegion] = []
    if PYMUPDF_AVAILABLE and fitz_page is not None:
        try:
            layout_regions = LayoutAnalyser.analyse(fitz_page)
        except Exception as _le:
            logger.debug(f"LayoutAnalyser page {page_num}: {_le}")

    # ── 3. Build word charmap from pdfplumber ─────────────────────────────
    try:
        words = plumber_page.extract_words(
            x_tolerance=3, y_tolerance=3, keep_blank_chars=False
        ) or []
    except Exception:
        words = []

    norm_page_text, charmap = _words_to_charmap(words)

    try:
        _plumber_full_text      = plumber_page.extract_text() or ""
        norm_page_text_plumber  = re.sub(r'\s+', ' ', _plumber_full_text.lower()).strip()
    except Exception:
        norm_page_text_plumber  = norm_page_text

    # ── 4. Build annotated right panel ────────────────────────────────────
    right   = raw_img.copy().convert("RGBA")
    overlay = PILImage.new("RGBA", right.size, (0, 0, 0, 0))
    draw    = PILImageDraw.Draw(right)

    try:
        font       = PILImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 12)
        small_font = PILImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 10)
    except Exception:
        font = small_font = None

    # ── 4a. Layout region overlays (diagrams / tables / images) ───────────
    _region_color_map = {
        "diagram"   : _VIZ_DIAGRAM_COLOR,
        "line_table": _VIZ_TABLE_COLOR,
        "image"     : _VIZ_IMAGE_COLOR,
    }

    region_layer = PILImage.new("RGBA", right.size, (0, 0, 0, 0))
    region_draw  = PILImageDraw.Draw(region_layer)

    for reg in layout_regions:
        color = _region_color_map.get(reg.kind, _VIZ_DIAGRAM_COLOR)
        rx0   = max(0,         int(reg.bbox[0] * sx) - 4)
        ry0   = max(0,         int(reg.bbox[1] * sy) - 4)
        rx1   = min(img_w - 1, int(reg.bbox[2] * sx) + 4)
        ry1   = min(img_h - 1, int(reg.bbox[3] * sy) + 4)

        # Semi-transparent fill
        region_draw.rectangle(
            [rx0, ry0, rx1, ry1],
            fill=(*color, _VIZ_REGION_ALPHA),
        )
        # Bold outline (3px)
        for _lw in range(3):
            region_draw.rectangle(
                [rx0 + _lw, ry0 + _lw, rx1 - _lw, ry1 - _lw],
                outline=(*color, 220),
            )

        # Label tag at top-left of region
        tag  = reg.label or reg.kind.replace("_", " ").title()
        tag_x, tag_y = rx0 + 3, ry0 + 3
        # Dark shadow for readability
        region_draw.text((tag_x + 1, tag_y + 1), tag, fill=(0, 0, 0, 200), font=font)
        region_draw.text((tag_x, tag_y),          tag, fill=(*color, 255),  font=font)

    right.alpha_composite(region_layer)

    # ── 4b. Section divider lines ──────────────────────────────────────────
    sy_scale = img_h / pdf_h

    for sec in sections:
        if not sec.content.strip():
            continue
        # Search for section title first, then content anchor
        anchor = re.sub(r'\s+', ' ', sec.title.strip().lower())[:50] \
                 if sec.title and sec.title != "Document" \
                 else re.sub(r'\s+', ' ', sec.content.strip().lower())[:30]
        idx = norm_page_text.find(anchor)
        if idx == -1 and norm_page_text_plumber:
            alt_idx = norm_page_text_plumber.find(anchor)
            if alt_idx != -1:
                ratio = alt_idx / max(len(norm_page_text_plumber), 1)
                idx   = int(ratio * len(norm_page_text))
        if idx == -1:
            continue
        for (cs, ce, wx0, wtop, wx1, wbot) in charmap:
            if cs >= idx:
                line_y = max(0, int(wtop * sy_scale) - 3)
                draw.line([(0, line_y), (img_w - 1, line_y)],
                          fill=_VIZ_SECTION_COLOR, width=1)
                break

    # ── 4c. Parent chunk text bounding boxes ───────────────────────────────
    for pi, parent in enumerate(parents):
        color = _VIZ_PARENT_COLORS[pi % len(_VIZ_PARENT_COLORS)]
        label = f"P{pi + 1}"
        bbox  = _find_bbox_for_text(
            parent.text, norm_page_text, charmap,
            pdf_w, pdf_h, img_w, img_h,
            norm_page_alt=norm_page_text_plumber,
        )
        if bbox:
            _draw_box(draw, overlay, bbox, color, label, line_width=3, font=font)

    right.alpha_composite(overlay)

    # ── 4d. Child chunk boxes (thinner, lighter) ────────────────────────────
    if draw_children:
        child_overlay = PILImage.new("RGBA", right.size, (0, 0, 0, 0))
        child_draw    = PILImageDraw.Draw(right)
        for ci, child in enumerate(children):
            pi_match = 0
            for pi, parent in enumerate(parents):
                if getattr(child, "parent_id", "") == getattr(parent, "parent_id", None):
                    pi_match = pi
                    break
            base_color = _VIZ_PARENT_COLORS[pi_match % len(_VIZ_PARENT_COLORS)]
            color      = _lighter(base_color, 0.4)
            label      = f"C{pi_match + 1}{chr(ord('a') + (ci % 26))}"
            bbox       = _find_bbox_for_text(
                child.text, norm_page_text, charmap,
                pdf_w, pdf_h, img_w, img_h,
                norm_page_alt=norm_page_text_plumber,
            )
            if bbox:
                _draw_box(child_draw, child_overlay, bbox, color, label,
                          line_width=2, font=small_font)
        right.alpha_composite(child_overlay)

    # Single final RGB conversion after all compositing
    right = right.convert("RGB")

    # ── 5. Captions ───────────────────────────────────────────────────────
    emoji      = _PAGE_TYPE_EMOJI.get(page_type, page_type)
    diag_count = sum(1 for r in layout_regions if r.kind == "diagram")
    tbl_count  = sum(1 for r in layout_regions if r.kind == "line_table")
    img_count  = sum(1 for r in layout_regions if r.kind == "image")

    layout_note = ""
    parts = []
    if diag_count:
        parts.append(f"{diag_count}D")
    if tbl_count:
        parts.append(f"{tbl_count}T")
    if img_count:
        parts.append(f"{img_count}I")
    if parts:
        layout_note = f"  ·  layout[{'|'.join(parts)}]"

    left_cap  = f"PAGE {page_num}  ·  {emoji}  (original)"
    right_cap = (
        f"PAGE {page_num}  ·  {len(sections)} sec  ·  "
        f"{len(parents)} P  ·  {len(children)} C"
        f"{layout_note}"
        f"  ·  🟠=diag 🔵=tbl 🟢=img"
    )
    left_panel  = _add_caption(raw_img, left_cap)
    right_panel = _add_caption(right,   right_cap)

    # ── 6. Combine side by side with a divider ────────────────────────────
    divider_w = 4
    total_w   = left_panel.width + divider_w + right_panel.width
    total_h   = max(left_panel.height, right_panel.height)

    combined = PILImage.new("RGB", (total_w, total_h), (200, 200, 200))
    combined.paste(left_panel,  (0, 0))
    combined.paste(right_panel, (left_panel.width + divider_w, 0))

    return combined


# ═══════════════════════════════════════════════════════════════════════════
# Display helpers
# ═══════════════════════════════════════════════════════════════════════════

_PAGE_TYPE_EMOJI: Dict[str, str] = {
    PageType.TEXT:               "📝 text",
    PageType.LIST:               "📋 list",
    PageType.TABLE:              "📊 table",
    PageType.IMAGE:              "🖼️  image",
    PageType.DIAGRAM:            "📐 diagram",
    PageType.EQUATION:           "🔢 equation",
    PageType.MIXED:              "🔀 mixed",
    PageType.COVER:              "📄 cover",
    PageType.TOC:                "📑 toc",
    PageType.REVISION_HISTORY:   "🕓 revision_history",
    PageType.DISCLAIMER:         "⚖️  disclaimer",
    PageType.CODE_SNIPPET:       "💻 code_snippet",
    PageType.CLASS_REFERENCE:    "🗂️  class_reference",
    PageType.SPECIFICATION_ITEM: "📌 specification_item",
    PageType.GLOSSARY:           "📖 glossary",
    PageType.UNKNOWN:            "❓ unknown",
}

_W = 90  # report width


def _bar(label: str, value: int, total: int, width: int = 30) -> str:
    if total == 0:
        pct, filled = 0.0, 0
    else:
        pct = value / total
        filled = round(pct * width)
    bar = "█" * filled + "░" * (width - filled)
    return f"  {label:<15} [{bar}] {value:>4}  ({pct*100:5.1f} %)"


def _divider(char: str = "─", width: int = _W) -> str:
    return char * width


def _header(title: str, char: str = "═", width: int = _W) -> str:
    pad = max(0, width - len(title) - 4)
    left = pad // 2
    right = pad - left
    return f"{char * (left + 2)}  {title}  {char * (right + 2)}"


# ═══════════════════════════════════════════════════════════════════════════
# Parse per-page blocks from the concatenated text produced by EnhancedPDFLoader
# ═══════════════════════════════════════════════════════════════════════════

def _parse_page_blocks(full_text: str) -> List[Tuple[int, str, str]]:
    """
    Split the concatenated page text into (page_number, page_type, page_body) triples.
    EnhancedPDFLoader.load() embeds  \\n[Page N | type=<type>]\\n  before each page.
    """
    tag_re = re.compile(r'\n\[Page\s+(\d+)\s*\|\s*type=(\w+)\]')
    parts  = tag_re.split(full_text)
    # layout: [preamble, pn, pt, body, pn, pt, body, …]
    blocks: List[Tuple[int, str, str]] = []
    i = 1
    while i + 2 < len(parts):
        blocks.append((int(parts[i]), parts[i + 1], parts[i + 2]))
        i += 3
    return blocks


# ═══════════════════════════════════════════════════════════════════════════
# Main inspection logic
# ═══════════════════════════════════════════════════════════════════════════

def inspect_pdf(
    pdf_path: str,
    output_path: str,
    use_vision: bool = True,
    ollama_url: str = OLLAMA_URL,
    vision_model: str = OLLAMA_VISION_MODEL,
    vision_desc_model: str = OLLAMA_VISION_DESC_MODEL,
    top_k_preview: int = 300,
    save_viz: bool = False,
    viz_children: bool = True,
    viz_dir: Optional[str] = None,
    viz_scale: float = VIZ_DEFAULT_SCALE,
) -> None:
    """
    Run the full V3 pipeline on *pdf_path* and write a detailed report to
    *output_path*.  Nothing is written to Qdrant.

    If *save_viz* is True, a side-by-side comparison PNG is saved for every
    non-skipped page into *viz_dir* (default: <pdf-stem>_viz/ next to the PDF).
    Left panel shows the raw page render; right panel shows the same render
    annotated with section dividers, parent chunk boxes, and (optionally) child
    chunk boxes.
    """
    filename = Path(pdf_path).name
    lines: List[str] = []

    def emit(text: str = "") -> None:
        lines.append(text)
        print(text)

    t_start = time.time()

    # ── Viz directory setup ───────────────────────────────────────────────
    if save_viz:
        if viz_dir is None:
            # Called directly (not via main()), so we build the path here.
            stem    = Path(pdf_path).stem
            viz_dir = str(Path(output_path).parent / "viz")
        # Ensure the directory exists (main() already created it; this is a
        # safety net for direct calls to inspect_pdf()).
        os.makedirs(viz_dir, exist_ok=True)
        if not PILLOW_AVAILABLE:
            save_viz = False
            emit("  ⚠ Pillow not installed — --save-viz disabled.")

    # ── Open pdfplumber for word-level bboxes (used by viz) ───────────────
    _plumber_doc = None
    _plumber_pages: Dict[int, object] = {}   # page_index (0-based) → plumber page
    _fitz_doc    = None
    _fitz_pages: Dict[int, object] = {}      # page_index (0-based) → fitz page
    if save_viz:
        try:
            import pdfplumber as _pdfplumber
            _plumber_doc = _pdfplumber.open(pdf_path)
            for _pi, _pp in enumerate(_plumber_doc.pages):
                _plumber_pages[_pi] = _pp
        except Exception as _e:
            emit(f"  ⚠ Could not open pdfplumber for viz ({_e}) — skipping viz.")

        if PYMUPDF_AVAILABLE:
            try:
                _fitz_doc = _fitz.open(pdf_path)
                for _pi in range(len(_fitz_doc)):
                    _fitz_pages[_pi] = _fitz_doc[_pi]
                emit(f"  ✓ PyMuPDF loaded — vector diagram + table overlays enabled")
            except Exception as _e:
                emit(f"  ⚠ PyMuPDF unavailable for viz ({_e}) — layout overlays disabled")
        else:
            emit("  ℹ PyMuPDF (fitz) not installed — install with: pip install pymupdf"
                 " — layout overlays disabled")

    # ── Header ────────────────────────────────────────────────────────────
    emit(_header(f"PDF CHUNK INSPECTOR — {filename}"))
    emit(f"  File    : {os.path.abspath(pdf_path)}")
    emit(f"  Report  : {os.path.abspath(output_path)}")
    emit(f"  Vision  : {'ENABLED — ' + vision_model if use_vision else 'DISABLED (heuristics only)'}")
    if save_viz:
        emit(f"  Viz dir : {os.path.abspath(viz_dir)}  (children={'yes' if viz_children else 'no'})")
    emit()
    emit("  Per-content-type chunk sizes (child tier):")
    for ct in [PageType.TEXT, PageType.LIST, PageType.TABLE, PageType.IMAGE,
               PageType.DIAGRAM, PageType.EQUATION, PageType.MIXED,
               PageType.CODE_SNIPPET, PageType.CLASS_REFERENCE,
               PageType.SPECIFICATION_ITEM, PageType.GLOSSARY]:
        c_sz = CHILD_CHUNK_SIZES.get(ct, CHILD_CHUNK_SIZE)
        p_sz = PARENT_CHUNK_SIZES.get(ct, PARENT_CHUNK_SIZE)
        emit(f"    {ct:<10} child={c_sz:>5} chars  |  parent={p_sz:>5} chars")
    emit(f"  Overlap : child={CHILD_CHUNK_OVERLAP} chars  |  parent={PARENT_CHUNK_OVERLAP} chars")
    emit()

    # ── Initialise classifier (optional) ─────────────────────────────────
    classifier        = None
    content_extractor = None
    if use_vision and ENABLE_PAGE_CLASSIFICATION:
        emit("▶  Initialising vision classifier …")
        try:
            classifier        = OllamaVisionClassifier(
                ollama_url     = ollama_url,
                classify_model = vision_model,
                describe_model = vision_desc_model,
            )
            content_extractor = ContentTypeExtractor(
                ollama_url     = ollama_url,
                classify_model = vision_model,
                describe_model = vision_desc_model,
            )
            emit(f"   ✓ Connected to Ollama at {ollama_url}")
        except Exception as exc:
            emit(f"   ⚠ Classifier unavailable ({exc}) — heuristics only.")
            classifier = content_extractor = None
    else:
        emit("▶  Heuristic classifier only (--no-vision).")

    # ── Load the full PDF ─────────────────────────────────────────────────
    emit()
    emit(f"▶  Loading & classifying: {filename} …")
    t0 = time.time()
    full_text, doc_meta = AdvancedDocumentLoader.load(
        pdf_path,
        classifier        = classifier,
        content_extractor = content_extractor,
    )
    load_elapsed = time.time() - t0

    if not full_text:
        emit("  ✗ No text could be extracted from this PDF. Aborting.")
        _write_report(output_path, lines)
        return

    total_pages         = doc_meta.get("num_pages", 0)
    page_types_meta: Dict[int, str] = doc_meta.get("page_types", {})

    emit(f"   ✓ Done in {load_elapsed:.1f}s  |  "
         f"{total_pages} pages  |  "
         f"{len(full_text):,} chars total")
    emit()

    # ── Parse into per-page blocks ────────────────────────────────────────
    page_blocks = _parse_page_blocks(full_text)

    if not page_blocks:
        emit("  ⚠ Could not parse per-page blocks from extracted text.")
        emit("    (The PDF may use a non-standard text layout.)")
        emit("    Proceeding with global-only analysis.")

    # ── Initialise SmartSectionDetector ──────────────────────────────────
    toc_entries: List[Tuple[str, int, int]] = doc_meta.get("toc_entries", [])
    emit()
    if toc_entries:
        emit(f"▶  ToC bookmark tree: {len(toc_entries)} entries found — using as primary section signal")
    else:
        emit("▶  ToC bookmark tree: none found — falling back to font analysis + LLM")

    # Build pdfplumber page dict for font detector (open once, reuse)
    _sd_plumber_pages: Dict[int, object] = {}
    _sd_plumber_doc   = None
    try:
        _sd_plumber_doc = pdfplumber.open(pdf_path)
        for _pi, _pp in enumerate(_sd_plumber_doc.pages):
            _sd_plumber_pages[_pi] = _pp
    except Exception as _e:
        emit(f"  ⚠ pdfplumber for section detection unavailable: {_e}")

    font_detector = FontSectionDetector()
    llm_detector  = LLMSectionDetector(ollama_url=ollama_url, model=OLLAMA_SECTION_MODEL)
    smart_detector = SmartSectionDetector(
        toc_entries   = toc_entries,
        font_detector = font_detector,
        llm_detector  = llm_detector,
        plumber_pages = _sd_plumber_pages,
    )

    # ── Initialise chunker & builder ──────────────────────────────────────
    chunker = SemanticChunker(chunk_size=CHILD_CHUNK_SIZE, overlap=CHILD_CHUNK_OVERLAP)
    builder = ParentChildBuilder(chunker)

    # ═══════════════════════════════════════════════════════════════════════
    # PER-PAGE ANALYSIS
    # ═══════════════════════════════════════════════════════════════════════
    emit(_header("PER-PAGE CLASSIFICATION & CHUNKING BREAKDOWN"))
    emit()

    # Accumulate totals for the summary
    type_counter: Dict[str, int]  = {}
    skipped_pages: List[int]      = []
    total_per_page_parents: int   = 0
    total_per_page_children: int  = 0

    # Collect per-page data for the cross-table
    cross_table: List[Tuple[int, str, int, int, bool]] = []

    for page_num, page_type, page_body in page_blocks:
        emoji_label = _PAGE_TYPE_EMOJI.get(page_type, "❓ " + page_type)
        type_counter[page_type] = type_counter.get(page_type, 0) + 1

        # ── Page header ───────────────────────────────────────────────────
        emit(_divider())
        emit(f"PAGE {page_num:>4}  |  {emoji_label.upper()}")
        emit(_divider())

        # ── Classification details ────────────────────────────────────────
        classified_by_vision = use_vision and (page_num in page_types_meta)
        has_tables_text  = "[Table " in page_body
        has_image_text   = "[IMAGE " in page_body
        has_diagram_text = "[DIAGRAM " in page_body
        emit(f"  Classifier result : {page_type}  "
             f"({'vision LLM' if classified_by_vision else 'heuristic'})")
        emit(f"  Text extracted    : {len(page_body):,} chars")
        emit(f"  Embedded tables   : {'YES — pdfplumber pipe-delimited' if has_tables_text else 'none'}")
        emit(f"  Embedded images   : {'YES' if has_image_text else 'none'}")
        emit(f"  Embedded diagrams : {'YES' if has_diagram_text else 'none'}")

        # Show inherited section hierarchy from ToC
        toc_hier = smart_detector.hierarchy_at_page(page_num)
        emit(f"  ToC hierarchy     : {' > '.join(toc_hier)}")

        if not page_body.strip():
            emit()
            emit("  No text content — no chunks produced.")
            cross_table.append((page_num, page_type, 0, 0, False))
            emit()
            continue

        # ── Run chunking on this page's text only ─────────────────────────
        sections = chunker.detect_sections(page_body, smart_detector=smart_detector)
        parents, children = builder.build(sections, file_hash_value=f"{filename}_p{page_num}")

        total_per_page_parents  += len(parents)
        total_per_page_children += len(children)
        cross_table.append((page_num, page_type, len(parents), len(children), False))

        # ── Section breakdown ─────────────────────────────────────────────
        emit()
        emit(f"  SECTIONS DETECTED:  {len(sections)}")
        for si, sec in enumerate(sections, 1):
            hier = " > ".join(sec.section_hierarchy) if sec.section_hierarchy else sec.title
            emit(f"    [{si:>3}] type={sec.section_type:<8}  "
                 f"level={sec.level}  "
                 f"len={len(sec.content):>5} chars")
            emit(f"          hierarchy : {hier[:80]}")
            emit(f"          preview   : {repr(sec.content[:100])}")

        # ── Parent chunks ─────────────────────────────────────────────────
        emit()
        emit(f"  PARENT CHUNKS:  {len(parents)}  "
             "(large context window — stored in _parents collection)")
        for pi, p in enumerate(parents, 1):
            hier = " > ".join(p.section_hierarchy) if p.section_hierarchy else p.section_title
            emit(f"    Parent {pi:>3}:  "
                 f"chars [{p.start_char:>6} … {p.end_char:>6}]  "
                 f"({p.end_char - p.start_char:>5} chars, {p.word_count} words)")
            emit(f"      section   : {hier[:75]}")
            emit(f"      page_type : {p.page_type}  |  chunk_type: {p.chunk_type}")
            emit(f"      parent_id : {p.parent_id}")
            emit(f"      preview   : {repr(p.text[:top_k_preview])}")

        # ── Child chunks ──────────────────────────────────────────────────
        emit()
        emit(f"  CHILD CHUNKS:  {len(children)}  "
             "(precision retrieval units — stored in _children collection)")
        for ci, c in enumerate(children, 1):
            hier = " > ".join(c.section_hierarchy) if c.section_hierarchy else c.section_title
            emit(f"    Child {ci:>4}:  "
                 f"chars [{c.start_char:>6} … {c.end_char:>6}]  "
                 f"({c.end_char - c.start_char:>5} chars)  "
                 f"words={c.word_count}  sents={c.sentence_count}  "
                 f"chunk_type={c.chunk_type}")
            emit(f"      parent_id : {c.parent_id[:32]}…")
            emit(f"      section   : {hier[:75]}")
            emit(f"      preview   : {repr(c.text[:top_k_preview])}")

        # ── Side-by-side visualisation ────────────────────────────────────
        if save_viz:
            page_index_0 = page_num - 1   # pdfplumber uses 0-based indices
            plumber_pg   = _plumber_pages.get(page_index_0)
            try:
                viz_img = build_viz_page(
                    pdf_path    = pdf_path,
                    page_index  = page_index_0,
                    page_num    = page_num,
                    page_type   = page_type,
                    plumber_page= plumber_pg,
                    sections    = sections,
                    parents     = parents,
                    children    = children,
                    draw_children = viz_children,
                    scale       = viz_scale,
                    fitz_page   = _fitz_pages.get(page_index_0),
                )
                if viz_img is not None:
                    viz_fname = f"page_{page_num:04d}_comparison.png"
                    viz_fpath = os.path.join(viz_dir, viz_fname)
                    viz_img.save(viz_fpath, format="PNG", optimize=False)
                    emit()
                    emit(f"  🖼  viz saved → {viz_fpath}")
                else:
                    emit()
                    emit(f"  ⚠ viz: could not render page {page_num} (pypdfium2/Pillow unavailable?)")
            except Exception as _viz_err:
                emit()
                emit(f"  ⚠ viz error page {page_num}: {_viz_err}")

        emit()

    # ── Close pdfplumber docs if we opened them ───────────────────────────
    if _plumber_doc is not None:
        try:
            _plumber_doc.close()
        except Exception:
            pass
    if _sd_plumber_doc is not None:
        try:
            _sd_plumber_doc.close()
        except Exception:
            pass
    if _fitz_doc is not None:
        try:
            _fitz_doc.close()
        except Exception:
            pass

    # ═══════════════════════════════════════════════════════════════════════
    # GLOBAL RUN (whole document processed as one unit — mirrors real ingestion)
    # ═══════════════════════════════════════════════════════════════════════
    emit()
    emit(_header("GLOBAL CHUNKING (mirrors real ingestion — chunks may span pages)"))
    emit()
    emit("  Running section detection and parent-child build on the full document …")
    t0 = time.time()
    global_sections = chunker.detect_sections(full_text, smart_detector=smart_detector)
    global_parents, global_children = builder.build(global_sections, file_hash_value=filename)
    global_elapsed  = time.time() - t0

    emit(f"  Done in {global_elapsed:.1f}s")
    emit()
    emit(f"  Sections detected (global)  : {len(global_sections)}")
    emit(f"  Parent chunks (global)      : {len(global_parents)}"
         f"  → stored in _parents collection")
    emit(f"  Child  chunks (global)      : {len(global_children)}"
         f"  → stored in _children collection")
    emit()

    if global_parents:
        sizes = [p.end_char - p.start_char for p in global_parents]
        emit(f"  Parent size  min={min(sizes):>5}  max={max(sizes):>5}  "
             f"avg={sum(sizes)//len(sizes):>5} chars")
    if global_children:
        sizes = [c.end_char - c.start_char for c in global_children]
        emit(f"  Child  size  min={min(sizes):>5}  max={max(sizes):>5}  "
             f"avg={sum(sizes)//len(sizes):>5} chars")

    emit()
    emit("  Global section list (title | type | level | chars):")
    for si, sec in enumerate(global_sections, 1):
        hier = " > ".join(sec.section_hierarchy) if sec.section_hierarchy else sec.title
        emit(f"    [{si:>4}]  type={sec.section_type:<8}  "
             f"lvl={sec.level}  "
             f"len={len(sec.content):>6} chars  "
             f"│ {hier[:60]}")

    # ═══════════════════════════════════════════════════════════════════════
    # PAGE × CHUNK CROSS-TABLE
    # ═══════════════════════════════════════════════════════════════════════
    emit()
    emit(_header("PAGE × CHUNK COUNT  (per-page independent analysis)"))
    emit()
    col = [6, 14, 10, 10, 14]
    emit(f"  {'Page':>{col[0]}}  "
         f"{'Type':<{col[1]}}  "
         f"{'Parents':>{col[2]}}  "
         f"{'Children':>{col[3]}}  "
         f"{'Note':<{col[4]}}")
    emit("  " + _divider("-", sum(col) + 8))
    for pn, pt, np_, nc, sk in cross_table:
        note = "— no text" if (np_ == 0 and nc == 0 and pt == "no_text") else ("— 0 chunks" if (np_ == 0 and nc == 0) else "")
        emit(f"  {pn:>{col[0]}}  "
             f"{pt:<{col[1]}}  "
             f"{np_:>{col[2]}}  "
             f"{nc:>{col[3]}}  "
             f"{note:<{col[4]}}")
    emit()
    emit(f"  TOTAL (per-page analysis)  parents={total_per_page_parents}  "
         f"children={total_per_page_children}")
    emit(f"  TOTAL (global  analysis)   parents={len(global_parents)}  "
         f"children={len(global_children)}")
    emit()
    emit("  Note: totals differ because global chunking allows chunks to span page")
    emit("  boundaries, while per-page analysis treats each page independently.")

    # ═══════════════════════════════════════════════════════════════════════
    # SUMMARY
    # ═══════════════════════════════════════════════════════════════════════
    emit()
    emit(_header("SUMMARY"))
    emit()
    emit(f"  PDF                   : {filename}")
    emit(f"  Total pages           : {total_pages}")
    emit(f"  All pages indexed     : YES — no pages skipped")
    emit(f"  ToC bookmark entries  : {len(toc_entries)}"
         + (f"  (used as primary section signal)" if toc_entries else "  (font+LLM fallback used)"))
    emit(f"  Indexed pages         : {len(page_blocks)}")
    emit()
    emit("  Page-type breakdown:")
    total_classified = sum(type_counter.values())
    for pt in [PageType.TEXT, PageType.LIST, PageType.TABLE, PageType.IMAGE, PageType.DIAGRAM,
               PageType.EQUATION, PageType.MIXED, PageType.CODE_SNIPPET, PageType.CLASS_REFERENCE,
               PageType.SPECIFICATION_ITEM, PageType.GLOSSARY,
               PageType.COVER, PageType.TOC, PageType.REVISION_HISTORY, PageType.DISCLAIMER,
               PageType.UNKNOWN]:
        count = type_counter.get(pt, 0)
        if count:
            emit(_bar(_PAGE_TYPE_EMOJI.get(pt, pt), count, total_classified))
    emit()
    emit(f"  Sections (global)     : {len(global_sections)}")
    emit(f"  Parent chunks (global): {len(global_parents)}  "
         f"(stored in _parents collection, ~{PARENT_CHUNK_SIZE} chars each)")
    emit(f"  Child  chunks (global): {len(global_children)}  "
         f"(stored in _children collection, ~{CHILD_CHUNK_SIZE} chars each)")

    # ═══════════════════════════════════════════════════════════════════════
    # ANOMALY CHECK
    # ═══════════════════════════════════════════════════════════════════════
    emit()
    emit(_header("ANOMALY CHECK"))
    emit()
    anomalies: List[str] = []
    for pn, pt, np_, nc, _sk in cross_table:
        # All pages are now indexed; no_text pages produce 0 chunks which is expected
        if pt == "no_text":
            continue
        body     = next((b for n, t, b in page_blocks if n == pn), "")
        body_len = len(body.strip())
        if nc == 0 and body_len > 200:
            anomalies.append(
                f"  ⚠  Page {pn:>3} ({pt}): {body_len} chars extracted but 0 child chunks."
                f"  Check section detection or content extraction for this page."
            )
        if pt in (PageType.TABLE, PageType.MIXED) and np_ == 0:
            anomalies.append(
                f"  ⚠  Page {pn:>3} ({pt}): table/mixed page produced 0 parent chunks."
            )
        if pt in (PageType.IMAGE, PageType.DIAGRAM) and body_len < 50:
            anomalies.append(
                f"  ℹ  Page {pn:>3} ({pt}): very short text ({body_len} chars) — "
                f"vision description may be needed for useful retrieval."
            )
    if anomalies:
        for a in anomalies:
            emit(a)
    else:
        emit("  ✓  No anomalies detected.")
    emit()

    # ── Footer ─────────────────────────────────────────────────────────────
    total_elapsed = time.time() - t_start
    emit(_header(f"Inspection complete in {total_elapsed:.1f}s"))
    emit()

    _write_report(output_path, lines)


def _write_report(output_path: str, lines: List[str]) -> None:
    """Write the collected report lines to a UTF-8 text file."""
    parent_dir = os.path.dirname(os.path.abspath(output_path))
    os.makedirs(Path(output_path).parent, exist_ok=True)
    if parent_dir:
        os.makedirs(parent_dir, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as fh:
        fh.write("\n".join(lines))
    print(f"\n✓ Report written to: {output_path}")


# ═══════════════════════════════════════════════════════════════════════════
# CLI entry-point
# ═══════════════════════════════════════════════════════════════════════════

def main() -> None:
    parser = argparse.ArgumentParser(
        prog="PDF_Chunk_Inspector",
        description=(
            "Inspect how Qdrant_Database_Generation_V3.py would classify and chunk "
            "a given PDF — without writing anything to Qdrant."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "Examples:\n"
            "  python PDF_Chunk_Inspector.py --pdf document.pdf\n"
            "  python PDF_Chunk_Inspector.py --pdf document.pdf --no-vision\n"
            "  python PDF_Chunk_Inspector.py --pdf document.pdf --output report.txt\n"
            "  python PDF_Chunk_Inspector.py --pdf document.pdf --top-k-preview 500\n"
            "  python PDF_Chunk_Inspector.py --pdf document.pdf --save-viz\n"
            "  python PDF_Chunk_Inspector.py --pdf document.pdf --save-viz --no-viz-children\n"
            "  python PDF_Chunk_Inspector.py --pdf document.pdf --save-viz --viz-scale 2.0\n"
        ),
    )
    parser.add_argument(
        "--pdf", required=True, metavar="PATH",
        help="Path to the PDF file to inspect.",
    )
    parser.add_argument(
        "--output", default=None, metavar="PATH",
        help=(
            "Path for the output report (.txt). "
            "Default: <pdf-stem>_inspection.txt in the same directory as the PDF."
        ),
    )
    parser.add_argument(
        "--no-vision", action="store_true",
        help="Disable Ollama vision classifier; use heuristics only (much faster).",
    )
    parser.add_argument(
        "--ollama-url", default=OLLAMA_URL, metavar="URL",
        help=f"Ollama base URL (default: {OLLAMA_URL}).",
    )
    parser.add_argument(
        "--vision-model", default=OLLAMA_VISION_MODEL, metavar="MODEL",
        help=f"Ollama vision model for page classification (default: {OLLAMA_VISION_MODEL}).",
    )
    parser.add_argument(
        "--vision-desc-model", default=OLLAMA_VISION_DESC_MODEL, metavar="MODEL",
        help=f"Ollama model for image/diagram descriptions (default: {OLLAMA_VISION_DESC_MODEL}).",
    )
    parser.add_argument(
        "--top-k-preview", type=int, default=300, metavar="N",
        help="Number of characters to preview for each chunk in the report (default: 300).",
    )
    parser.add_argument(
        "--save-viz", action="store_true",
        help=(
            "Generate a side-by-side comparison PNG for every non-skipped page. "
            "Left panel: original render. Right panel: annotated with section lines, "
            "parent boxes, and child boxes."
        ),
    )
    parser.add_argument(
        "--no-viz-children", action="store_true",
        help="When --save-viz is active, omit child chunk boxes (show parent boxes only).",
    )
    parser.add_argument(
        "--viz-dir", default=None, metavar="PATH",
        help=(
            "Directory to write visualisation PNGs into. "
            "Default: <pdf-stem>_viz/ next to the report file."
        ),
    )
    parser.add_argument(
        "--viz-scale", type=float, default=VIZ_DEFAULT_SCALE, metavar="FLOAT",
        help=f"Render scale for visualisation images (default: {VIZ_DEFAULT_SCALE}). "
             "Higher = sharper but larger files.",
    )
    args = parser.parse_args()

    # ── Validate input PDF first ──────────────────────────────────────────
    if not os.path.isfile(args.pdf):
        sys.exit(f"ERROR: File not found: {args.pdf}")
    if not args.pdf.lower().endswith(".pdf"):
        sys.exit(f"ERROR: Expected a .pdf file, got: {args.pdf}")

    pdf_path = Path(args.pdf).resolve()
    pdf_name = pdf_path.stem

    # ── Build output directory paths ─────────────────────────────────────
    # Output is always stored next to the script, NOT next to the PDF.
    # Structure: <script_dir>/<pdf_stem>/
    #                 ├── <pdf_stem>_report.txt
    #                 └── viz/
    #                     ├── page_0001_comparison.png
    #                     └── ...
    script_dir  = Path(__file__).resolve().parent
    base_output = Path(args.output).resolve() if args.output else script_dir
    output_dir  = base_output / pdf_name

    # Always recreate output_dir cleanly so stale files never accumulate.
    # On a fresh run (no prior output) this is equivalent to mkdir.
    # On a rerun it wipes the previous report + any stale viz PNGs.
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True, exist_ok=False)

    output_path = output_dir / f"{pdf_name}_report.txt"

    # viz dir is always inside output_dir
    viz_dir = output_dir / "viz"
    viz_dir.mkdir(parents=True, exist_ok=False)

    inspect_pdf(
        pdf_path          = str(pdf_path),
        output_path       = str(output_path),
        use_vision        = not args.no_vision,
        ollama_url        = args.ollama_url,
        vision_model      = args.vision_model,
        vision_desc_model = args.vision_desc_model,
        top_k_preview     = args.top_k_preview,
        save_viz          = args.save_viz,
        viz_children      = not args.no_viz_children,
        viz_dir           = str(viz_dir),
        viz_scale         = args.viz_scale,
    )


if __name__ == "__main__":
    main()
