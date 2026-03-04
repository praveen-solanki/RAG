"""
ADVANCED RAG INGESTION SYSTEM
===============================
Features:
- Section-aware chunking with context preservation
- Hybrid search (Dense embeddings + BM25 sparse vectors)
- Rich metadata extraction
- BGE-M3 embeddings support via Ollama
- Multi-vector storage in Qdrant
- Semantic chunking with sentence boundaries
"""
import argparse
import json
import numpy as np
import os
import re
import hashlib
import uuid
import time
from typing import Optional, List, Dict, Tuple
from dataclasses import dataclass
from pathlib import Path
import logging

import pdfplumber
import docx
import requests
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import (
    VectorParams,
    Distance,
    Filter,
    FieldCondition,
    MatchValue,
    PointStruct,
    SparseVector,
    SparseVectorParams,
    SparseIndexParams,
)
from rank_bm25 import BM25Okapi
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize

# Download NLTK data if needed
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    try:
        nltk.download('punkt_tab', quiet=True)
    except Exception:
        nltk.download('punkt', quiet=True)

# ================= CONFIG =================

DATA_DIR = r"/home/olj3kor/praveen/Image_dataset_generation/pdfs/standards"
COLLECTION = "test_63pdf"
QDRANT_URL = "http://localhost:7333"

# Embedding options
USE_OLLAMA_BGE_M3 = True  # Set to True to use Ollama BGE-M3
OLLAMA_URL = "http://localhost:11434"
OLLAMA_MODEL = "bge-m3:latest"

# Fallback to SentenceTransformer if Ollama unavailable
FALLBACK_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

# Chunking parameters
CHUNK_SIZE = 512  # Increased for better context
CHUNK_OVERLAP = 128
MIN_CHUNK_SIZE = 100
MAX_CHUNK_SIZE = 1024

# Section detection
ENABLE_SECTION_AWARE = True
SECTION_PATTERNS = [
    r'^#{1,6}\s+(.+)$',  # Markdown headers
    r'^([A-Z][^.!?]*):$',  # Title case with colon
    r'^\d+\.\s+([A-Z].+)$',  # Numbered sections
    r'^([A-Z\s]{3,})$',  # All caps headers (min 3 chars)
]

BM25_OUTPUT = "bm25_index.json"

# =========================================

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class DocumentSection:
    """Represents a document section with hierarchy"""
    title: str
    content: str
    level: int
    page_number: Optional[int] = None
    section_type: str = "text"  # text, table, list, code
    section_hierarchy: Optional[List[str]] = None

    def __post_init__(self):
        if self.section_hierarchy is None:
            self.section_hierarchy = [self.title]


@dataclass
class EnrichedChunk:
    """Chunk with rich metadata"""
    text: str
    section_title: str
    section_hierarchy: List[str]
    page_number: Optional[int]
    chunk_type: str
    word_count: int
    sentence_count: int
    start_char: int
    end_char: int


class OllamaBGEM3Embedder:
    """BGE-M3 embedder using Ollama"""
    
    def __init__(self, base_url: str = "http://localhost:11434", model: str = "bge-m3"):
        self.base_url = base_url
        self.model = model
        self.dimension = 1024  # BGE-M3 dimension
        self._test_connection()
    
    def _test_connection(self):
        """Test Ollama connection and verify model availability"""
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=5)
            if response.status_code == 200:
                logger.info(f"✓ Connected to Ollama at {self.base_url}")
                # Verify the requested model is available
                available_models = [m.get("name", "") for m in response.json().get("models", [])]
                if self.model not in available_models:
                    logger.warning(
                        f"Model '{self.model}' not found in Ollama. "
                        f"Available: {available_models}"
                    )
            else:
                raise ConnectionError("Ollama not responding")
        except Exception as e:
            logger.error(f"✗ Cannot connect to Ollama: {e}")
            raise
    
    def encode(self, texts: List[str], batch_size: int = 8, show_progress_bar: bool = False) -> List[List[float]]:
        """Encode texts using Ollama BGE-M3 with retry logic"""
        embeddings = []
        
        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            
            for text in batch:
                embedding = None
                for attempt in range(3):
                    try:
                        response = requests.post(
                            f"{self.base_url}/api/embeddings",
                            json={"model": self.model, "prompt": text},
                            timeout=30
                        )
                        
                        if response.status_code == 200:
                            embedding = response.json()["embedding"]
                            break
                        else:
                            logger.warning(f"Embedding attempt {attempt+1} failed with status {response.status_code}")
                    except Exception as e:
                        logger.warning(f"Embedding attempt {attempt+1} error: {e}")
                    if attempt < 2:
                        time.sleep(2 ** attempt)

                if embedding is None:
                    logger.warning(f"Failed to get embedding after 3 attempts, skipping chunk")
                embeddings.append(embedding)
            
            if show_progress_bar and (i // batch_size) % 10 == 0:
                logger.info(f"Encoded {min(i + batch_size, len(texts))}/{len(texts)} texts")
        
        return embeddings


class SectionAwareChunker:
    """Advanced chunker that respects document structure"""
    
    def __init__(self, chunk_size: int = 512, overlap: int = 128):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.section_patterns = [re.compile(p, re.MULTILINE) for p in SECTION_PATTERNS]
    
    def detect_sections(self, text: str) -> List[DocumentSection]:
        """Detect document sections and hierarchy"""
        sections = []
        lines = text.split('\n')
        current_section = {"title": "Introduction", "content": "", "level": 0}
        section_stack = [current_section]
        
        for line_num, line in enumerate(lines):
            line_stripped = line.strip()
            
            if not line_stripped:
                current_section["content"] += "\n"
                continue
            
            # Check if line is a section header
            is_header = False
            header_level = 0
            
            for pattern in self.section_patterns:
                match = pattern.match(line_stripped)
                if match:
                    is_header = True
                    # Determine level based on pattern type
                    if line_stripped.startswith('#'):
                        header_level = len(line_stripped) - len(line_stripped.lstrip('#'))
                    elif line_stripped.isupper():
                        header_level = 1
                    else:
                        header_level = 2
                    break
            
            if is_header and len(line_stripped) < 200:  # Headers shouldn't be too long
                # Save previous section
                if current_section["content"].strip():
                    # Build hierarchy from stack
                    hierarchy = [s["title"] for s in section_stack if s["title"]]
                    # Detect table section type
                    section_type = "table" if current_section["title"].startswith("[Table ") else "text"
                    sections.append(DocumentSection(
                        title=current_section["title"],
                        content=current_section["content"].strip(),
                        level=current_section["level"],
                        section_type=section_type,
                        section_hierarchy=list(hierarchy),
                    ))
                
                # Pop stack entries at same or deeper level
                while len(section_stack) > 1 and section_stack[-1]["level"] >= header_level:
                    section_stack.pop()
                
                # Start new section
                current_section = {
                    "title": line_stripped.strip('#: ').strip(),
                    "content": "",
                    "level": header_level
                }
                section_stack.append(current_section)
            else:
                current_section["content"] += line + "\n"
        
        # Add final section
        if current_section["content"].strip():
            hierarchy = [s["title"] for s in section_stack if s["title"]]
            section_type = "table" if current_section["title"].startswith("[Table ") else "text"
            sections.append(DocumentSection(
                title=current_section["title"],
                content=current_section["content"].strip(),
                level=current_section["level"],
                section_type=section_type,
                section_hierarchy=list(hierarchy),
            ))
        
        return sections if sections else [DocumentSection("Document", text, 0)]
    
    def _split_long_sentence(self, sentence: str) -> List[str]:
        """Split a sentence exceeding MAX_CHUNK_SIZE at word boundaries"""
        if len(sentence) <= MAX_CHUNK_SIZE:
            return [sentence]
        words = sentence.split()
        parts = []
        current = ""
        for word in words:
            if len(current) + len(word) + 1 > MAX_CHUNK_SIZE and current:
                parts.append(current.strip())
                current = word
            else:
                current = current + " " + word if current else word
        if current.strip():
            parts.append(current.strip())
        return parts

    def chunk_with_sentences(self, text: str, section_title: str = "", section_hierarchy: Optional[List[str]] = None) -> List[EnrichedChunk]:
        """Chunk text respecting sentence boundaries"""
        if not text or len(text) < MIN_CHUNK_SIZE:
            return []
        
        if section_hierarchy is None:
            section_hierarchy = [section_title]

        raw_sentences = sent_tokenize(text)
        # Enforce MAX_CHUNK_SIZE by splitting long sentences
        sentences = []
        for s in raw_sentences:
            sentences.extend(self._split_long_sentence(s))

        chunks = []
        current_chunk = ""
        current_start = 0
        
        for sentence in sentences:
            # If adding sentence exceeds chunk size
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                # Save current chunk
                chunks.append(EnrichedChunk(
                    text=current_chunk.strip(),
                    section_title=section_title,
                    section_hierarchy=section_hierarchy,
                    page_number=None,
                    chunk_type="text",
                    word_count=len(word_tokenize(current_chunk)),
                    sentence_count=len(sent_tokenize(current_chunk)),
                    start_char=current_start,
                    end_char=current_start + len(current_chunk)
                ))
                
                # Start new chunk with overlap
                old_chunk_len = len(current_chunk)
                overlap_text = current_chunk[-self.overlap:] if len(current_chunk) > self.overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
                current_start += old_chunk_len - len(overlap_text)
            else:
                current_chunk += " " + sentence if current_chunk else sentence
        
        # Add final chunk
        if current_chunk.strip():
            chunks.append(EnrichedChunk(
                text=current_chunk.strip(),
                section_title=section_title,
                section_hierarchy=section_hierarchy,
                page_number=None,
                chunk_type="text",
                word_count=len(word_tokenize(current_chunk)),
                sentence_count=len(sent_tokenize(current_chunk)),
                start_char=current_start,
                end_char=current_start + len(current_chunk)
            ))
        
        return chunks
    
    def chunk_sections(self, sections: List[DocumentSection]) -> List[EnrichedChunk]:
        """Chunk document sections while preserving context"""
        all_chunks = []
        
        for section in sections:
            # Build hierarchy path from section level
            hierarchy = section.section_hierarchy if hasattr(section, 'section_hierarchy') else [section.title]
            section_chunks = self.chunk_with_sentences(section.content, section.title, hierarchy)
            
            # Add section metadata
            for chunk in section_chunks:
                chunk.page_number = section.page_number
                chunk.chunk_type = section.section_type
            
            all_chunks.extend(section_chunks)
        
        return all_chunks


class BM25Index:
    """BM25 sparse vector index"""
    
    def __init__(self):
        self.bm25 = None
        self.tokenized_corpus = []
        self.vocabulary = {}
        self.token_idf = {}
    
    def fit(self, texts: List[str]):
        """Build BM25 index"""
        self.tokenized_corpus = [self._tokenize(t) for t in texts]
        self.bm25 = BM25Okapi(self.tokenized_corpus)
        
        # Build vocabulary from ALL tokens across ALL documents
        all_tokens = sorted({t for doc in self.tokenized_corpus for t in doc})
        self.vocabulary = {token: idx for idx, token in enumerate(all_tokens)}
        
        # Compute IDF per token correctly
        N = len(self.tokenized_corpus)
        for token in self.vocabulary:
            df = sum(1 for doc in self.tokenized_corpus if token in doc)
            self.token_idf[token] = np.log((N - df + 0.5) / (df + 0.5) + 1)
        
        logger.info(f"  ✓ BM25 vocabulary size: {len(self.vocabulary)}")
    
    def _tokenize(self, text: str) -> List[str]:
        """Tokenize text for BM25"""
        return [token.lower() for token in word_tokenize(text) if token.isalnum()]
    
    def get_sparse_vector(self, text: str) -> SparseVector:
        """Get BM25 sparse vector for text"""
        tokens = self._tokenize(text)
        total = len(tokens)
        token_counts = {}
        
        for token in tokens:
            if token in self.vocabulary:
                token_counts[token] = token_counts.get(token, 0) + 1
        
        indices = []
        values = []
        
        for token, count in token_counts.items():
            tf = count / total if total else 0
            idf = self.token_idf.get(token, 1.0)
            indices.append(self.vocabulary[token])
            values.append(float(tf * idf))
        
        return SparseVector(indices=indices, values=values)

class AdvancedDocumentLoader:
    """Enhanced document loader with metadata extraction"""
    
    @staticmethod
    def extract_metadata(path: str) -> Dict:
        """Extract document metadata"""
        file_stat = os.stat(path)
        
        return {
            "file_size_bytes": file_stat.st_size,
            "created_timestamp": file_stat.st_ctime,
            "modified_timestamp": file_stat.st_mtime,
            "file_extension": Path(path).suffix.lower(),
        }
    
    @staticmethod
    def load_pdf(path: str) -> Tuple[str, Dict]:
        """Load PDF with enhanced metadata"""
        text = ""
        metadata = {"num_pages": 0, "has_tables": False, "tables_count": 0}
        
        with pdfplumber.open(path) as pdf:
            metadata["num_pages"] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, start=1):
                # Extract text
                page_text = page.extract_text()
                if page_text:
                    text += f"\n[Page {page_num}]\n{page_text}\n"
                
                # Extract tables
                tables = page.extract_tables()
                if tables:
                    metadata["has_tables"] = True
                    metadata["tables_count"] += len(tables)
                    
                    for table_idx, table in enumerate(tables, start=1):
                        if table:
                            text += f"\n[Table {table_idx} on Page {page_num}]\n"
                            for row in table:
                                text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                            text += "\n"
        
        return text, metadata
    
    @staticmethod
    def load_docx(path: str) -> Tuple[str, Dict]:
        """Load DOCX with metadata"""
        doc = docx.Document(path)
        text_parts = []
        metadata = {"num_paragraphs": 0, "has_tables": False, "tables_count": 0}
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                metadata["num_paragraphs"] += 1
        
        # Extract tables
        if doc.tables:
            metadata["has_tables"] = True
            metadata["tables_count"] = len(doc.tables)
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)
        
        return "\n".join(text_parts), metadata
    
    @staticmethod
    def load_txt(path: str) -> Tuple[str, Dict]:
        """Load TXT with metadata"""
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        
        metadata = {
            "num_lines": len(text.split('\n')),
            "char_count": len(text)
        }
        
        return text, metadata
    
    @classmethod
    def load(cls, path: str) -> Tuple[Optional[str], Dict]:
        """Universal loader with metadata"""
        try:
            base_metadata = cls.extract_metadata(path)
            
            if path.lower().endswith(".pdf"):
                text, doc_metadata = cls.load_pdf(path)
            elif path.lower().endswith(".docx"):
                text, doc_metadata = cls.load_docx(path)
            elif path.lower().endswith(".txt"):
                text, doc_metadata = cls.load_txt(path)
            else:
                return None, {}
            
            # Merge metadata
            base_metadata.update(doc_metadata)
            
            return text, base_metadata
            
        except Exception as e:
            logger.error(f"Error loading {path}: {e}")
            return None, {}


def file_hash(path: str) -> str:
    """Generate SHA256 hash of file using block-wise reading"""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def already_indexed(client: QdrantClient, collection: str, file_hash_value: str) -> bool:
    """Check if file already indexed"""
    try:
        filt = Filter(
            must=[FieldCondition(key="file_hash", match=MatchValue(value=file_hash_value))]
        )
        points, _ = client.scroll(collection_name=collection, scroll_filter=filt, limit=1)
        return len(points) > 0
    except Exception as e:
        logger.warning(f"Could not check if file is already indexed: {e}")
        return False


def main():
    """Main ingestion pipeline"""

    parser = argparse.ArgumentParser(description="Advanced RAG Ingestion System")
    parser.add_argument('--data-dir', type=str, default=DATA_DIR, help='Directory containing documents')
    parser.add_argument('--collection', type=str, default=COLLECTION, help='Qdrant collection name')
    parser.add_argument('--qdrant-url', type=str, default=QDRANT_URL, help='Qdrant server URL')
    parser.add_argument('--ollama-url', type=str, default=OLLAMA_URL, help='Ollama server URL')
    parser.add_argument('--ollama-model', type=str, default=OLLAMA_MODEL, help='Ollama embedding model name')
    parser.add_argument('--chunk-size', type=int, default=CHUNK_SIZE, help='Chunk size in characters')
    parser.add_argument('--chunk-overlap', type=int, default=CHUNK_OVERLAP, help='Chunk overlap in characters')
    parser.add_argument('--bm25-output', type=str, default=BM25_OUTPUT, help='Path to save BM25 index JSON')
    args = parser.parse_args()

    data_dir = args.data_dir
    collection = args.collection
    qdrant_url = args.qdrant_url
    ollama_url = args.ollama_url
    ollama_model = args.ollama_model
    chunk_size = args.chunk_size
    chunk_overlap = args.chunk_overlap
    bm25_output = args.bm25_output

    logger.info("="*80)
    logger.info("ADVANCED RAG INGESTION SYSTEM")
    logger.info("="*80)
    
    # Initialize Qdrant client
    client = QdrantClient(url=qdrant_url)
    
    # Initialize embedding model
    if USE_OLLAMA_BGE_M3:
        try:
            logger.info(f"Using Ollama BGE-M3 model at {ollama_url}")
            embedder = OllamaBGEM3Embedder(ollama_url, ollama_model)
            embedding_dim = embedder.dimension
        except Exception as e:
            logger.warning(f"Ollama unavailable, falling back to {FALLBACK_MODEL}")
            embedder = SentenceTransformer(FALLBACK_MODEL)
            embedding_dim = 384
    else:
        logger.info(f"Using SentenceTransformer: {FALLBACK_MODEL}")
        embedder = SentenceTransformer(FALLBACK_MODEL)
        embedding_dim = 384
    
    # Create collection with hybrid search support, checking for dimension mismatch
    existing = [c.name for c in client.get_collections().collections]
    if collection not in existing:
        client.create_collection(
            collection_name=collection,
            vectors_config={
                "dense": VectorParams(
                    size=embedding_dim,
                    distance=Distance.COSINE,
                )
            },
            sparse_vectors_config={
                "bm25": SparseVectorParams(
                    index=SparseIndexParams(
                        on_disk=False,
                    )
                )
            },
        )
        logger.info(f"✓ Created collection: {collection}")
    else:
        # Check for dimension mismatch
        col_info = client.get_collection(collection)
        existing_dim = col_info.config.params.vectors.get("dense").size if hasattr(col_info.config.params.vectors, "get") else None
        if existing_dim is not None and existing_dim != embedding_dim:
            logger.error(
                f"✗ Dimension mismatch: collection '{collection}' has {existing_dim}-dim vectors, "
                f"but current embedder produces {embedding_dim}-dim vectors. Aborting."
            )
            return
        logger.info(f"✓ Collection exists: {collection}")
    
    # Initialize components
    chunker = SectionAwareChunker(chunk_size, chunk_overlap)
    bm25_index = BM25Index()
    loader = AdvancedDocumentLoader()
    
    # Process files
    total_chunks = 0
    total_files = 0
    all_points = []
    
    logger.info("\nPhase 1: Loading and chunking documents...")
    
    for root, _, files in os.walk(data_dir):
        for file in files:
            if not file.lower().endswith((".pdf", ".docx", ".txt")):
                continue
            
            path = os.path.join(root, file)
            logger.info(f"\n→ Processing: {file}")
            
            try:
                h = file_hash(path)
                
                if already_indexed(client, collection, h):
                    logger.info("  ⊘ Already indexed")
                    continue
                
                # Load document
                text, doc_metadata = loader.load(path)
                if not text or len(text.strip()) < MIN_CHUNK_SIZE:
                    logger.info("  ⊘ Empty or too short")
                    continue
                
                # Section-aware chunking
                if ENABLE_SECTION_AWARE:
                    sections = chunker.detect_sections(text)
                    chunks = chunker.chunk_sections(sections)
                    logger.info(f"  ✓ Detected {len(sections)} sections → {len(chunks)} chunks")
                else:
                    chunks = chunker.chunk_with_sentences(text, "Document")
                    logger.info(f"  ✓ Created {len(chunks)} chunks")
                
                if not chunks:
                    continue
                
                chunk_texts = [chunk.text for chunk in chunks]
                
                # Generate embeddings
                logger.info(f"  ⚡ Generating embeddings...")
                if isinstance(embedder, OllamaBGEM3Embedder):
                    embeddings = embedder.encode(chunk_texts, batch_size=8)
                else:
                    embeddings = embedder.encode(
                        chunk_texts,
                        batch_size=32,
                        show_progress_bar=False,
                        convert_to_numpy=True
                    )
                    embeddings = [emb.tolist() for emb in embeddings]
                
                # Create points — only for chunks with successful embeddings
                for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                    if embedding is None:
                        logger.warning(f"  ⚠ Skipping chunk {i} — embedding failed")
                        continue
                    point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{h}_{i}"))
                    
                    point = {
                        "id": point_id,
                        "vector": embedding,
                        "payload": {
                            "content": chunk.text,
                            "source_path": path,
                            "filename": file,
                            "folder": os.path.relpath(root, data_dir),
                            "file_type": Path(file).suffix.lower(),
                            "file_hash": h,
                            "chunk_id": i,
                            "total_chunks": len(chunks),
                            "chunk_length": len(chunk.text),
                            "word_count": chunk.word_count,
                            "sentence_count": chunk.sentence_count,
                            "section_title": chunk.section_title,
                            "section_hierarchy": chunk.section_hierarchy,
                            "chunk_type": chunk.chunk_type,
                            **doc_metadata
                        }
                    }
                    all_points.append((point, chunk.text))
                    total_chunks += 1
                
                total_files += 1
                
            except Exception as e:
                logger.error(f"  ✗ Error: {e}")
                continue
    
    if not all_points:
        logger.warning("No documents to index!")
        return
    
    # Collect BM25 texts only from successfully-embedded chunks
    all_texts_for_bm25 = [text for _, text in all_points]

    # Build BM25 index
    logger.info(f"\nPhase 2: Building BM25 index for {len(all_texts_for_bm25)} chunks...")
    bm25_index.fit(all_texts_for_bm25)
    
    # Save BM25 vocabulary and IDF for retrieval (convert numpy floats to Python floats)
    bm25_data = {
        "vocabulary": bm25_index.vocabulary,
        "token_idf": {k: float(v) for k, v in bm25_index.token_idf.items()},
    }
    Path(bm25_output).parent.mkdir(parents=True, exist_ok=True)
    with open(bm25_output, "w") as f:
        json.dump(bm25_data, f)
    logger.info(f"  ✓ Saved BM25 index to {bm25_output}")
    logger.info("\nPhase 3: Adding sparse vectors and uploading to Qdrant...")
    
    points_to_upload = []
    for point_data, text in all_points:
        sparse_vector = bm25_index.get_sparse_vector(text)
        
        point_struct = PointStruct(
            id=point_data["id"],
            vector={
                "dense": point_data["vector"],  # Dense vector
                "bm25": sparse_vector           # Sparse vector
            },
            payload=point_data["payload"]
        )
        points_to_upload.append(point_struct)
    
    # Batch upload
    batch_size = 100
    for i in range(0, len(points_to_upload), batch_size):
        batch = points_to_upload[i:i + batch_size]
        client.upsert(collection_name=collection, points=batch, wait=True)
        logger.info(f"  ✓ Uploaded batch {i//batch_size + 1}/{(len(points_to_upload)-1)//batch_size + 1}")
    
    # Summary
    logger.info("\n" + "="*80)
    logger.info("INGESTION COMPLETE")
    logger.info("="*80)
    logger.info(f"✓ Files processed: {total_files}")
    logger.info(f"✓ Total chunks: {total_chunks}")
    logger.info(f"✓ Avg chunks/file: {total_chunks/total_files if total_files > 0 else 0:.1f}")
    logger.info(f"✓ Collection: {collection}")
    logger.info(f"✓ Embedding dimension: {embedding_dim}")
    logger.info(f"✓ BM25 vocabulary size: {len(bm25_index.vocabulary)}")
    logger.info("="*80)


if __name__ == "__main__":
    main()