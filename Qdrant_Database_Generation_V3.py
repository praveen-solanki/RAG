"""
ADVANCED RAG INGESTION SYSTEM V3
=================================
New features over V2:

1. Parent-Child Collection Architecture
   - Small child chunks (≤512 chars) go into `{collection}_children` for precise retrieval
   - Large parent chunks (≤2048 chars) go into `{collection}_parents` for rich LLM context
   - Each child stores parent_id; retrieval returns parent text to the LLM

2. Enhanced PDF Extraction (three-tier pipeline)
   - Tier 1: pypdfium2 (word-spacing-correct, fast)
   - Tier 2: pdfplumber fallback for pages where pypdfium2 yields sparse text
   - Tier 3: pytesseract OCR for pages that remain empty after Tiers 1+2
             (requires pdf2image + tesseract installed; skipped gracefully if absent)
   - Tables extracted as pipe-delimited text (preserves table content, not just count)

3. Semantic Chunking
   - Sentences are grouped by TF-IDF cosine similarity of sliding windows
   - Splits only at semantic breakpoints (topic transitions), not arbitrary char counts
   - Falls back to sentence-boundary chunking when sklearn is unavailable

4. Contextual Chunk Enrichment (Anthropic research — 49% retrieval improvement)
   - Two modes: "llm" (calls Ollama generation model) or "template" (fast, no extra LLM call)
   - Template mode prepends document title, section hierarchy, and neighbouring sentence
     context to every child chunk before embedding; the raw text is preserved separately
   - LLM mode uses Ollama to generate a precise situational description of each chunk

5. MinHash LSH Deduplication (replaces word-level Jaccard sliding window)
   - 128-permutation MinHash signatures using mmh3 (already in requirements)
   - 32-band LSH with 4 rows per band → ~83% threshold collision probability
   - Global deduplication across ALL files processed in a single run (not per-file)
   - O(1) lookup vs O(window × chunk_text_length) for Jaccard

6. Additional improvements
   - Per-file statistics (accepted/skipped/dedup'ed chunks logged)
   - Better BM25 tokenisation (stop-word-aware)
   - Graceful degradation: every new feature can be disabled via CONFIG flags
"""

import argparse
import json
import math
import time
import os
import re
import hashlib
import uuid
from collections import defaultdict
from typing import Optional, List, Dict, Tuple, Set
from dataclasses import dataclass, field
from pathlib import Path
import logging

import numpy as np
import pdfplumber
import pypdfium2 as pdfium
import docx
import requests
import mmh3
from qdrant_client import QdrantClient
from qdrant_client.models import (
    VectorParams,
    Distance,
    Filter,
    FieldCondition,
    MatchValue,
    PointStruct,
    SparseVector,
    SparseVectorParams,
    SparseIndexParams,
)
from rank_bm25 import BM25Okapi
from sentence_transformers import SentenceTransformer
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize

# Optional OCR imports (graceful degradation if not installed)
try:
    from pdf2image import convert_from_path as pdf_to_images
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# Optional sklearn for semantic chunking
try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity as sk_cosine
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False

# ─── NLTK bootstrap ─────────────────────────────────────────────────────────
for _resource in ("tokenizers/punkt", "tokenizers/punkt_tab",
                  "corpora/stopwords"):
    try:
        nltk.data.find(_resource)
    except LookupError:
        _name = _resource.split("/")[-1]
        try:
            nltk.download(_name, quiet=True)
        except Exception:
            pass

try:
    from nltk.corpus import stopwords as _sw
    _STOP_WORDS: Set[str] = set(_sw.words("english"))
except Exception:
    _STOP_WORDS: Set[str] = set()


# ═══════════════════════════════════════════════════════════════════════════
# CONFIG
# ═══════════════════════════════════════════════════════════════════════════

DATA_DIR       = os.environ.get(
    "RAG_DATA_DIR",
    r"/home/olj3kor/praveen/Image_dataset_generation/pdfs/standards",
)
COLLECTION     = "rag_v3"          # base name; _children / _parents suffixes added automatically
QDRANT_URL     = "http://localhost:7333"

# Embedding
USE_OLLAMA_BGE_M3 = True
OLLAMA_URL        = "http://localhost:11434"
OLLAMA_MODEL      = "bge-m3:latest"
FALLBACK_MODEL    = "sentence-transformers/all-MiniLM-L6-v2"

# Generation model for LLM-based contextual enrichment
# Set to None or "" to use template-based enrichment only
CONTEXT_LLM_MODEL = ""             # e.g. "llama3:8b" if you have it in Ollama
CONTEXT_MODE      = "template"     # "template" | "llm"

# Chunking parameters
CHILD_CHUNK_SIZE    = 512          # child chunk target size (chars)
CHILD_CHUNK_OVERLAP = 256
PARENT_CHUNK_SIZE   = 2048         # parent chunk target size (chars)
PARENT_CHUNK_OVERLAP= 256
MIN_CHUNK_SIZE      = 80

# Semantic chunking
ENABLE_SEMANTIC_CHUNKING   = True
SEMANTIC_SPLIT_THRESHOLD   = 0.35  # cosine sim below this → semantic boundary
SEMANTIC_WINDOW_SIZE       = 3     # sentences per window for sim computation

# Contextual enrichment
ENABLE_CONTEXT_ENRICHMENT  = True
CONTEXT_NEIGHBOUR_SENTS    = 1     # sentences from prev/next chunk to include in template context

# MinHash dedup
ENABLE_MINHASH_DEDUP  = True
MINHASH_NUM_PERM      = 128        # number of permutations
MINHASH_BANDS         = 32         # LSH bands (threshold ≈ (1/bands)^(1/rows))
MINHASH_ROWS          = 4          # rows per band  → threshold ≈ 0.83
MINHASH_SHINGLE_SIZE  = 3          # character n-grams for shingling

# PDF extraction
ENABLE_OCR_FALLBACK   = True       # use pytesseract if tier-1/2 yield sparse text
OCR_TRIGGER_CHARS     = 80         # page chars below this triggers tier-2 / tier-3

# TOC detection (inherited from V2)
TOC_LINE_RATIO        = 0.50
TOC_MIN_CONTENT_CHARS = 50
TOC_MIN_LINE_COUNT    = 5
TOC_MAX_CONTENT_CHARS = 800

# Embedding safety
MAX_EMBED_CHARS = 4000

# Section detection
ENABLE_SECTION_AWARE = True
SECTION_PATTERNS = [
    r'^#{1,6}\s+(.+)$',
    r'^([A-Z][^.!?]*):$',
    r'^\d+\.\s+([A-Z].+)$',
    r'^([A-Z\s]{3,})$',
]

BM25_OUTPUT = "bm25_index_v3.json"

# ═══════════════════════════════════════════════════════════════════════════
# Logging
# ═══════════════════════════════════════════════════════════════════════════

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════
# Data classes
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class DocumentSection:
    title: str
    content: str
    level: int
    page_number: Optional[int] = None
    section_type: str = "text"
    section_hierarchy: Optional[List[str]] = None

    def __post_init__(self):
        if self.section_hierarchy is None:
            self.section_hierarchy = [self.title]


@dataclass
class ParentChunk:
    """Large context window stored in the parents collection."""
    text: str
    section_title: str
    section_hierarchy: List[str]
    page_number: Optional[int]
    chunk_type: str
    word_count: int
    start_char: int
    end_char: int
    parent_id: str = field(default_factory=lambda: str(uuid.uuid4()))


@dataclass
class ChildChunk:
    """Small precision chunk stored in the children collection."""
    text: str
    enriched_text: str          # text with contextual prefix (used for embedding)
    section_title: str
    section_hierarchy: List[str]
    page_number: Optional[int]
    chunk_type: str
    word_count: int
    sentence_count: int
    start_char: int
    end_char: int
    parent_id: str              # FK → ParentChunk.parent_id
    child_index: int            # position within parent


# ═══════════════════════════════════════════════════════════════════════════
# Helper utilities
# ═══════════════════════════════════════════════════════════════════════════

def file_hash(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def already_indexed(client: QdrantClient, collection: str,
                    file_hash_value: str) -> bool:
    try:
        filt = Filter(must=[
            FieldCondition(key="file_hash", match=MatchValue(value=file_hash_value))
        ])
        points, _ = client.scroll(collection_name=collection,
                                  scroll_filter=filt, limit=1)
        return len(points) > 0
    except Exception as e:
        logger.warning(f"Could not check existing index: {e}")
        return False


# ═══════════════════════════════════════════════════════════════════════════
# MinHash LSH deduplicator
# ═══════════════════════════════════════════════════════════════════════════

class MinHashDeduplicator:
    """
    Global near-duplicate detector using MinHash + LSH.

    Compared to the V2 word-level Jaccard sliding window this implementation:
    - Works across ALL files in a single run (not per-file)
    - Runs in O(num_perm) per chunk instead of O(window × len(text))
    - Approximates Jaccard similarity on character shingles (more robust for
      short technical texts that share vocabulary but differ in order)
    """

    def __init__(self,
                 num_perm: int = MINHASH_NUM_PERM,
                 bands: int = MINHASH_BANDS,
                 rows: int = MINHASH_ROWS,
                 shingle_size: int = MINHASH_SHINGLE_SIZE):
        self.num_perm = num_perm
        self.bands = bands
        self.rows = rows
        self.shingle_size = shingle_size
        # LSH buckets: band_index → bucket_key → list[text_id]
        self._buckets: List[Dict[int, List[int]]] = [
            defaultdict(list) for _ in range(bands)
        ]
        self._signatures: List[np.ndarray] = []
        self._ids: List[str] = []

    # ── Internal helpers ────────────────────────────────────────────────

    def _shingle(self, text: str) -> Set[str]:
        """Return set of character n-grams."""
        t = text.lower()
        k = self.shingle_size
        return {t[i:i + k] for i in range(max(0, len(t) - k + 1))} or {t}

    def _minhash_signature(self, shingles: Set[str]) -> np.ndarray:
        """Compute MinHash signature using mmh3 seeds."""
        sig = np.full(self.num_perm, np.iinfo(np.int64).max, dtype=np.int64)
        for shingle in shingles:
            for seed in range(self.num_perm):
                h = mmh3.hash(shingle, seed=seed, signed=True)
                if h < sig[seed]:
                    sig[seed] = h
        return sig

    def _lsh_add(self, sig: np.ndarray, idx: int):
        """Insert signature into LSH buckets."""
        for b in range(self.bands):
            start = b * self.rows
            band_key = hash(sig[start:start + self.rows].tobytes())
            self._buckets[b][band_key].append(idx)

    def _candidates(self, sig: np.ndarray) -> Set[int]:
        """Return candidate duplicate indices from LSH buckets."""
        candidates: Set[int] = set()
        for b in range(self.bands):
            start = b * self.rows
            band_key = hash(sig[start:start + self.rows].tobytes())
            for idx in self._buckets[b].get(band_key, []):
                candidates.add(idx)
        return candidates

    def _estimate_jaccard(self, sig_a: np.ndarray,
                          sig_b: np.ndarray) -> float:
        return float(np.mean(sig_a == sig_b))

    # ── Public API ───────────────────────────────────────────────────────

    # Maximum number of LSH candidates to verify before giving up.
    # Prevents O(n) comparisons when many chunks hash to the same band bucket
    # (e.g. highly repetitive corpus).  Set to 0 to disable the cap.
    _MAX_CANDIDATES = 200

    def is_duplicate(self, text: str, threshold: float = 0.80) -> bool:
        """
        Return True if text is a near-duplicate of something already seen.
        If not a duplicate, add text to the index so future calls can detect it.
        """
        shingles = self._shingle(text)
        if not shingles:
            return False

        sig = self._minhash_signature(shingles)
        idx = len(self._signatures)

        candidates = self._candidates(sig)
        if self._MAX_CANDIDATES:
            candidates = set(list(candidates)[: self._MAX_CANDIDATES])

        for cand_idx in candidates:
            if self._estimate_jaccard(sig, self._signatures[cand_idx]) >= threshold:
                return True

        # Not a duplicate — add to index
        self._signatures.append(sig)
        self._ids.append(str(idx))
        self._lsh_add(sig, idx)
        return False

    def __len__(self) -> int:
        return len(self._signatures)


# ═══════════════════════════════════════════════════════════════════════════
# Embedding models
# ═══════════════════════════════════════════════════════════════════════════

class OllamaBGEM3Embedder:
    """BGE-M3 embedder via Ollama with retry + NaN/Inf patching."""

    def __init__(self, base_url: str = "http://localhost:11434",
                 model: str = "bge-m3"):
        self.base_url = base_url
        self.model = model
        self.dimension = 1024
        self._test_connection()

    def _test_connection(self):
        try:
            r = requests.get(f"{self.base_url}/api/tags", timeout=5)
            if r.status_code == 200:
                logger.info(f"✓ Connected to Ollama at {self.base_url}")
                available = [m.get("name", "")
                             for m in r.json().get("models", [])]
                if self.model not in available:
                    logger.warning(
                        f"Model '{self.model}' not found. Available: {available}"
                    )
            else:
                raise ConnectionError("Ollama not responding")
        except Exception as e:
            logger.error(f"✗ Cannot connect to Ollama: {e}")
            raise

    def encode(self, texts: List[str], batch_size: int = 8,
               show_progress_bar: bool = False) -> List[Optional[List[float]]]:
        embeddings: List[Optional[List[float]]] = []
        skipped = 0
        patched = 0

        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]
            for text in batch:
                safe = text if len(text) <= MAX_EMBED_CHARS else text[:MAX_EMBED_CHARS]
                embedding = None
                for attempt in range(3):
                    try:
                        r = requests.post(
                            f"{self.base_url}/api/embeddings",
                            json={"model": self.model, "prompt": safe},
                            timeout=60,
                        )
                        if r.status_code == 200:
                            data = r.json()
                            if "embedding" in data:
                                raw = data["embedding"]
                                if not all(math.isfinite(x) for x in raw):
                                    arr = np.array(raw, dtype=np.float64)
                                    mask = ~np.isfinite(arr)
                                    logger.warning(
                                        f"Embedding attempt {attempt+1}: "
                                        f"{mask.sum()} non-finite component(s) → 0.0"
                                    )
                                    arr[mask] = 0.0
                                    raw = arr.tolist()
                                    patched += 1
                                embedding = raw
                                break
                            else:
                                logger.warning(
                                    f"Attempt {attempt+1}: no 'embedding' key: "
                                    f"{r.text[:100]}"
                                )
                        else:
                            body = r.text
                            logger.warning(
                                f"Attempt {attempt+1} HTTP {r.status_code}: "
                                f"{body[:100]}"
                            )
                            if "unsupported value" in body:
                                break
                    except Exception as e:
                        logger.warning(f"Attempt {attempt+1} error: {e}")
                    if attempt < 2:
                        time.sleep(2 ** attempt)

                if embedding is None:
                    skipped += 1
                    logger.warning(
                        f"Failed to embed chunk after 3 attempts "
                        f"({len(safe)} chars). Skipped so far: {skipped}"
                    )
                embeddings.append(embedding)

            if show_progress_bar and (i // batch_size) % 10 == 0:
                logger.info(
                    f"Encoded {min(i + batch_size, len(texts))}/{len(texts)}"
                )

        if patched:
            logger.warning(
                f"⚠  {patched}/{len(texts)} chunks had NaN/Inf components patched."
            )
        if skipped:
            logger.warning(
                f"⚠  {skipped}/{len(texts)} chunks could not be embedded and "
                f"will be missing from the index."
            )
        return embeddings


# ═══════════════════════════════════════════════════════════════════════════
# Enhanced PDF Loader (3-tier)
# ═══════════════════════════════════════════════════════════════════════════

class EnhancedPDFLoader:
    """
    Three-tier PDF text extraction:
    Tier 1 — pypdfium2 (fast, correct word spacing)
    Tier 2 — pdfplumber (better for some text-layer PDFs; used when Tier 1
              yields fewer than OCR_TRIGGER_CHARS per page)
    Tier 3 — pytesseract OCR (scanned PDFs; used when Tiers 1+2 both fail)
    Tables are extracted with pdfplumber and formatted as pipe-delimited text.
    """

    _TOC_LINE_RE = re.compile(r'(\.\s*){2,}.*\d+\s*$')

    @staticmethod
    def _is_toc_page(text: str) -> bool:
        stripped = text.strip()
        if len(stripped) < TOC_MIN_CONTENT_CHARS:
            return True
        if len(stripped) > TOC_MAX_CONTENT_CHARS:
            return False
        lines = [l for l in stripped.splitlines() if l.strip()]
        if len(lines) < TOC_MIN_LINE_COUNT:
            return False
        toc = sum(
            1 for l in lines
            if EnhancedPDFLoader._TOC_LINE_RE.search(l)
        )
        return (toc / len(lines)) >= TOC_LINE_RATIO

    @staticmethod
    def _extract_tables_as_text(plumber_page) -> str:
        """Extract tables from a pdfplumber page as pipe-delimited text."""
        parts: List[str] = []
        try:
            tables = plumber_page.extract_tables()
            for tbl_idx, table in enumerate(tables):
                rows = []
                for row in table:
                    cells = [str(c).strip() if c else "" for c in row]
                    rows.append(" | ".join(cells))
                if rows:
                    parts.append(
                        f"\n[Table {tbl_idx + 1}]\n" + "\n".join(rows) + "\n"
                    )
        except Exception as e:
            logger.debug(f"Table extraction error (non-fatal): {e}")
        return "".join(parts)

    @classmethod
    def load(cls, path: str) -> Tuple[str, Dict]:
        """Load PDF using 3-tier extraction; returns (full_text, metadata)."""
        metadata: Dict = {"num_pages": 0, "has_tables": False, "tables_count": 0}
        page_texts: List[str] = []
        skipped_pages: List[int] = []

        # ── Tier 1: pypdfium2 ──────────────────────────────────────────
        pdf = pdfium.PdfDocument(path)
        tier1_results: Dict[int, str] = {}
        try:
            metadata["num_pages"] = len(pdf)
            for pn in range(len(pdf)):
                page = pdf[pn]
                try:
                    tp = page.get_textpage()
                    try:
                        tier1_results[pn] = tp.get_text_bounded()
                    finally:
                        tp.close()
                except Exception as e:
                    logger.warning(f"pypdfium2 page {pn+1} error: {e}")
                    tier1_results[pn] = ""
                finally:
                    page.close()
        finally:
            pdf.close()

        # ── Tier 2 + 3 with pdfplumber (also handles tables) ──────────
        plumber_pages: Dict[int, object] = {}
        try:
            plumber_doc = pdfplumber.open(path)
            for pn, pp in enumerate(plumber_doc.pages):
                plumber_pages[pn] = pp
        except Exception as e:
            logger.warning(f"pdfplumber open error (non-fatal): {e}")
            plumber_doc = None

        for pn in range(metadata["num_pages"]):
            tier1_text = tier1_results.get(pn, "")
            page_text = tier1_text

            # Tier 2: pdfplumber fallback for sparse pages
            if len(tier1_text.strip()) < OCR_TRIGGER_CHARS and pn in plumber_pages:
                try:
                    t2 = plumber_pages[pn].extract_text() or ""
                    if len(t2.strip()) > len(tier1_text.strip()):
                        page_text = t2
                        logger.debug(f"  Page {pn+1}: used pdfplumber (tier-2)")
                except Exception as e:
                    logger.debug(f"  pdfplumber tier-2 page {pn+1}: {e}")

            # Tier 3: OCR for truly empty pages
            if (len(page_text.strip()) < OCR_TRIGGER_CHARS
                    and ENABLE_OCR_FALLBACK and OCR_AVAILABLE):
                try:
                    images = pdf_to_images(
                        path, first_page=pn + 1, last_page=pn + 1, dpi=200
                    )
                    if images:
                        ocr_text = pytesseract.image_to_string(images[0])
                        if len(ocr_text.strip()) > len(page_text.strip()):
                            page_text = ocr_text
                            logger.debug(f"  Page {pn+1}: used OCR (tier-3)")
                except Exception as e:
                    logger.debug(f"  OCR tier-3 page {pn+1}: {e}")

            if not page_text.strip():
                continue

            if cls._is_toc_page(page_text):
                skipped_pages.append(pn + 1)
                continue

            # Append text
            page_entry = f"\n[Page {pn + 1}]\n{page_text}\n"

            # Append table text (pdfplumber)
            if pn in plumber_pages:
                table_text = cls._extract_tables_as_text(plumber_pages[pn])
                if table_text:
                    metadata["has_tables"] = True
                    metadata["tables_count"] += table_text.count("[Table ")
                    page_entry += table_text

            page_texts.append(page_entry)

        if plumber_doc is not None:
            try:
                plumber_doc.close()
            except Exception:
                pass

        if skipped_pages:
            logger.info(
                f"  ⊘ Skipped {len(skipped_pages)} TOC/boilerplate pages: "
                f"{skipped_pages[:10]}{'...' if len(skipped_pages) > 10 else ''}"
            )

        return "".join(page_texts), metadata


# ═══════════════════════════════════════════════════════════════════════════
# Document loader (unified: PDF + DOCX + TXT)
# ═══════════════════════════════════════════════════════════════════════════

class AdvancedDocumentLoader:
    @staticmethod
    def extract_metadata(path: str) -> Dict:
        st = os.stat(path)
        return {
            "file_size_bytes": st.st_size,
            "created_timestamp": st.st_ctime,
            "modified_timestamp": st.st_mtime,
            "file_extension": Path(path).suffix.lower(),
        }

    @classmethod
    def load(cls, path: str) -> Tuple[Optional[str], Dict]:
        try:
            base_meta = cls.extract_metadata(path)
            if path.lower().endswith(".pdf"):
                text, doc_meta = EnhancedPDFLoader.load(path)
            elif path.lower().endswith(".docx"):
                text, doc_meta = cls._load_docx(path)
            elif path.lower().endswith(".txt"):
                text, doc_meta = cls._load_txt(path)
            else:
                return None, {}
            base_meta.update(doc_meta)
            return text, base_meta
        except Exception as e:
            logger.error(f"Error loading {path}: {e}")
            return None, {}

    @staticmethod
    def _load_docx(path: str) -> Tuple[str, Dict]:
        doc = docx.Document(path)
        parts: List[str] = []
        meta: Dict = {"num_paragraphs": 0, "has_tables": False, "tables_count": 0}
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
                meta["num_paragraphs"] += 1
        if doc.tables:
            meta["has_tables"] = True
            meta["tables_count"] = len(doc.tables)
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts), meta

    @staticmethod
    def _load_txt(path: str) -> Tuple[str, Dict]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        return text, {"num_lines": len(text.splitlines()), "char_count": len(text)}


# ═══════════════════════════════════════════════════════════════════════════
# Semantic Chunker
# ═══════════════════════════════════════════════════════════════════════════

class SemanticChunker:
    """
    Splits text at semantic boundaries detected via TF-IDF cosine similarity
    of sliding sentence windows.  Falls back to pure sentence-boundary chunking
    when sklearn is unavailable or when there are too few sentences.
    """

    def __init__(self, chunk_size: int = CHILD_CHUNK_SIZE,
                 overlap: int = CHILD_CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.overlap = overlap

    # ── Section detection (unchanged from V2) ───────────────────────────

    _section_patterns = [re.compile(p, re.MULTILINE) for p in SECTION_PATTERNS]

    def detect_sections(self, text: str) -> List[DocumentSection]:
        sections: List[DocumentSection] = []
        lines = text.split("\n")
        cur: Dict = {"title": "Introduction", "content": "", "level": 0}
        stack = [cur]

        for line in lines:
            ls = line.strip()
            if not ls:
                cur["content"] += "\n"
                continue
            is_hdr, level = False, 0
            for pat in self._section_patterns:
                if pat.match(ls):
                    is_hdr = True
                    level = (
                        len(ls) - len(ls.lstrip("#")) if ls.startswith("#")
                        else (1 if ls.isupper() else 2)
                    )
                    break
            if is_hdr and len(ls) < 200:
                if cur["content"].strip():
                    hier = [s["title"] for s in stack if s["title"]]
                    stype = "table" if cur["title"].startswith("[Table ") else "text"
                    sections.append(DocumentSection(
                        title=cur["title"],
                        content=cur["content"].strip(),
                        level=cur["level"],
                        section_type=stype,
                        section_hierarchy=list(hier),
                    ))
                while len(stack) > 1 and stack[-1]["level"] >= level:
                    stack.pop()
                cur = {"title": ls.strip("#: ").strip(), "content": "", "level": level}
                stack.append(cur)
            else:
                cur["content"] += line + "\n"

        if cur["content"].strip():
            hier = [s["title"] for s in stack if s["title"]]
            stype = "table" if cur["title"].startswith("[Table ") else "text"
            sections.append(DocumentSection(
                title=cur["title"],
                content=cur["content"].strip(),
                level=cur["level"],
                section_type=stype,
                section_hierarchy=list(hier),
            ))
        return sections or [DocumentSection("Document", text, 0)]

    # ── Semantic boundary detection ──────────────────────────────────────

    def _find_semantic_splits(self, sentences: List[str]) -> List[int]:
        """
        Return a list of sentence indices that mark the START of a new semantic
        segment.  Always includes 0.
        """
        splits = [0]
        n = len(sentences)
        if n < 2 * SEMANTIC_WINDOW_SIZE + 1 or not SKLEARN_AVAILABLE:
            return splits

        w = SEMANTIC_WINDOW_SIZE
        try:
            vectorizer = TfidfVectorizer(
                stop_words="english",
                max_features=5000,
                min_df=1,
            )
            # Build window texts
            windows = []
            for i in range(n):
                start = max(0, i - w)
                end = min(n, i + w + 1)
                windows.append(" ".join(sentences[start:end]))
            vectorizer.fit(windows)
            vecs = vectorizer.transform(windows).toarray()

            for i in range(1, n):
                prev = vecs[i - 1].reshape(1, -1)
                curr = vecs[i].reshape(1, -1)
                sim = float(sk_cosine(prev, curr)[0][0])
                if sim < SEMANTIC_SPLIT_THRESHOLD:
                    splits.append(i)
        except Exception as e:
            logger.debug(f"Semantic splitting fallback: {e}")

        return splits

    # ── Core chunking ────────────────────────────────────────────────────

    def _split_long_sentence(self, sentence: str,
                             max_size: int) -> List[str]:
        if len(sentence) <= max_size:
            return [sentence]
        words = sentence.split()
        parts: List[str] = []
        current = ""
        for w in words:
            if len(current) + len(w) + 1 > max_size and current:
                parts.append(current.strip())
                current = w
            else:
                current = current + " " + w if current else w
        if current.strip():
            parts.append(current.strip())
        return parts

    def _sentences_to_chunks(
        self,
        sentences: List[str],
        section_title: str,
        section_hierarchy: List[str],
        page_number: Optional[int],
        chunk_type: str,
        target_size: int,
        overlap_chars: int,
    ) -> List[ChildChunk]:
        """
        Pack sentences into chunks respecting semantic boundaries and target size.
        """
        if not sentences:
            return []

        # Optionally find semantic split points
        if ENABLE_SEMANTIC_CHUNKING:
            split_points = set(self._find_semantic_splits(sentences))
        else:
            split_points = set()

        chunks: List[ChildChunk] = []
        current_sents: List[str] = []
        current_len = 0
        start_char = 0
        char_pos = 0

        def _flush(sents: List[str], sc: int) -> ChildChunk:
            text = " ".join(sents).strip()
            return ChildChunk(
                text=text,
                enriched_text=text,   # filled later by ContextualEnricher
                section_title=section_title,
                section_hierarchy=list(section_hierarchy),
                page_number=page_number,
                chunk_type=chunk_type,
                word_count=len(word_tokenize(text)),
                sentence_count=len(sent_tokenize(text)),
                start_char=sc,
                end_char=sc + len(text),
                parent_id="",         # filled by parent-child builder
                child_index=0,        # filled by parent-child builder
            )

        for idx, sent in enumerate(sentences):
            at_boundary = idx in split_points and current_sents

            if current_len + len(sent) > target_size and current_sents:
                # Flush current chunk
                chunks.append(_flush(current_sents, start_char))

                # Keep overlap
                overlap_sents: List[str] = []
                overlap_len = 0
                for s in reversed(current_sents):
                    if overlap_len + len(s) <= overlap_chars:
                        overlap_sents.insert(0, s)
                        overlap_len += len(s)
                    else:
                        break

                start_char = char_pos - overlap_len
                current_sents = overlap_sents
                current_len = overlap_len

            elif at_boundary and current_len >= MIN_CHUNK_SIZE:
                chunks.append(_flush(current_sents, start_char))
                start_char = char_pos
                current_sents = []
                current_len = 0

            current_sents.append(sent)
            char_pos += len(sent) + 1
            current_len += len(sent) + 1

        if current_sents and " ".join(current_sents).strip():
            chunks.append(_flush(current_sents, start_char))

        return chunks

    def chunk_text(
        self,
        text: str,
        section_title: str = "",
        section_hierarchy: Optional[List[str]] = None,
        page_number: Optional[int] = None,
        chunk_type: str = "text",
        target_size: int = CHILD_CHUNK_SIZE,
        overlap: int = CHILD_CHUNK_OVERLAP,
    ) -> List[ChildChunk]:
        if not text or len(text.strip()) < MIN_CHUNK_SIZE:
            return []
        if section_hierarchy is None:
            section_hierarchy = [section_title] if section_title else ["Document"]

        raw_sents = sent_tokenize(text)
        sentences: List[str] = []
        for s in raw_sents:
            sentences.extend(self._split_long_sentence(s, target_size))

        return self._sentences_to_chunks(
            sentences, section_title, section_hierarchy,
            page_number, chunk_type, target_size, overlap,
        )

    def chunk_sections(
        self,
        sections: List[DocumentSection],
        target_size: int = CHILD_CHUNK_SIZE,
        overlap: int = CHILD_CHUNK_OVERLAP,
    ) -> List[ChildChunk]:
        result: List[ChildChunk] = []
        for sec in sections:
            hier = sec.section_hierarchy or [sec.title]
            chunks = self.chunk_text(
                sec.content, sec.title, hier,
                sec.page_number, sec.section_type, target_size, overlap,
            )
            result.extend(chunks)
        return result


# ═══════════════════════════════════════════════════════════════════════════
# Parent-Child Builder
# ═══════════════════════════════════════════════════════════════════════════

class ParentChildBuilder:
    """
    Builds parent chunks and their child chunks from document sections.

    Workflow:
    1. Chunk the document at PARENT_CHUNK_SIZE → ParentChunk list
    2. For each parent, further chunk at CHILD_CHUNK_SIZE → ChildChunk list
    3. Set child.parent_id = parent.parent_id
    """

    def __init__(self, chunker: SemanticChunker):
        self.chunker = chunker

    def build(
        self,
        sections: List[DocumentSection],
        file_hash_value: str = "",
    ) -> Tuple[List[ParentChunk], List[ChildChunk]]:
        parents: List[ParentChunk] = []
        children: List[ChildChunk] = []

        # Build large parent chunks first
        raw_parents: List[ChildChunk] = self.chunker.chunk_sections(
            sections,
            target_size=PARENT_CHUNK_SIZE,
            overlap=PARENT_CHUNK_OVERLAP,
        )

        for raw_parent in raw_parents:
            pid = str(uuid.uuid5(
                uuid.NAMESPACE_DNS,
                # Include file_hash to prevent collisions across documents
                # that happen to share the same 64-char prefix at the same offset.
                f"{file_hash_value}_{raw_parent.start_char}_{raw_parent.text[:64]}"
            ))
            parent = ParentChunk(
                text=raw_parent.text,
                section_title=raw_parent.section_title,
                section_hierarchy=list(raw_parent.section_hierarchy),
                page_number=raw_parent.page_number,
                chunk_type=raw_parent.chunk_type,
                word_count=raw_parent.word_count,
                start_char=raw_parent.start_char,
                end_char=raw_parent.end_char,
                parent_id=pid,
            )
            parents.append(parent)

            # Now chunk the parent content into children
            child_raw = self.chunker.chunk_text(
                raw_parent.text,
                raw_parent.section_title,
                raw_parent.section_hierarchy,
                raw_parent.page_number,
                raw_parent.chunk_type,
                CHILD_CHUNK_SIZE,
                CHILD_CHUNK_OVERLAP,
            )

            for ci, child in enumerate(child_raw):
                child.parent_id = pid
                child.child_index = ci
                children.append(child)

        return parents, children


# ═══════════════════════════════════════════════════════════════════════════
# Contextual Enricher (Anthropic research)
# ═══════════════════════════════════════════════════════════════════════════

class ContextualEnricher:
    """
    Prepends a short situational context to each child chunk before embedding.

    Two modes:
    - "template": Fast, no extra LLM call.  Builds context from filename,
      section hierarchy, and a snippet of the previous and next chunk.
    - "llm": Calls Ollama generation model to produce a precise description.
      Falls back to template mode if the LLM call fails.

    The original chunk text (child.text) is preserved unchanged; only
    child.enriched_text is modified.  This means the stored payload always
    holds the true source text while the embedding captures rich context.
    """

    _TEMPLATE = (
        "Document: {filename}\n"
        "Section: {section}\n"
        "Context: {context}\n\n"
        "{chunk_text}"
    )

    def __init__(self,
                 mode: str = CONTEXT_MODE,
                 ollama_url: str = OLLAMA_URL,
                 llm_model: str = CONTEXT_LLM_MODEL or ""):
        self.mode = mode if ENABLE_CONTEXT_ENRICHMENT else "none"
        self.ollama_url = ollama_url
        self.llm_model = llm_model

    def enrich(
        self,
        children: List[ChildChunk],
        filename: str,
    ) -> None:
        """Enrich all children in-place (modifies child.enriched_text)."""
        if self.mode == "none":
            return

        n = len(children)
        for i, child in enumerate(children):
            section = " > ".join(child.section_hierarchy) if child.section_hierarchy else ""

            # Build neighbouring context snippet
            prev_snippet = ""
            next_snippet = ""
            if i > 0:
                prev_sents = sent_tokenize(children[i - 1].text)
                prev_snippet = " ".join(
                    prev_sents[-CONTEXT_NEIGHBOUR_SENTS:]
                )
            if i < n - 1:
                next_sents = sent_tokenize(children[i + 1].text)
                next_snippet = " ".join(
                    next_sents[:CONTEXT_NEIGHBOUR_SENTS]
                )

            context_parts = []
            if prev_snippet:
                context_parts.append(f"...{prev_snippet}")
            if next_snippet:
                context_parts.append(f"{next_snippet}...")
            context_hint = " | ".join(context_parts) if context_parts else "start of section"

            if self.mode == "llm" and self.llm_model:
                enriched = self._llm_context(child.text, filename, section, context_hint)
            else:
                enriched = None

            if enriched is None:
                # Template fallback
                enriched = self._TEMPLATE.format(
                    filename=filename,
                    section=section or "—",
                    context=context_hint,
                    chunk_text=child.text,
                )

            child.enriched_text = enriched

    def _llm_context(
        self,
        chunk_text: str,
        filename: str,
        section: str,
        context_hint: str,
    ) -> Optional[str]:
        """
        Ask Ollama to write a short situational context for the chunk, then
        prepend it to the chunk text (Anthropic recipe).
        """
        prompt = (
            f"Document: {filename}\n"
            f"Section: {section or 'unknown'}\n"
            f"Neighbouring text snippet: {context_hint}\n\n"
            f"Please write 1–3 short sentences that explain what the following "
            f"passage is about and how it relates to the document above. "
            f"Do NOT repeat the passage itself.\n\n"
            f"Passage:\n{chunk_text[:1500]}\n\n"
            f"Context description:"
        )
        try:
            r = requests.post(
                f"{self.ollama_url}/api/generate",
                json={
                    "model": self.llm_model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {"temperature": 0.0, "num_predict": 120},
                },
                timeout=30,
            )
            if r.status_code == 200:
                context_desc = r.json().get("response", "").strip()
                if context_desc:
                    return f"{context_desc}\n\n{chunk_text}"
        except Exception as e:
            logger.debug(f"LLM context generation error: {e}")
        return None


# ═══════════════════════════════════════════════════════════════════════════
# BM25 Index (identical to V2, placed here for self-containment)
# ═══════════════════════════════════════════════════════════════════════════

class BM25Index:
    def __init__(self):
        self.bm25 = None
        self.tokenized_corpus: List[List[str]] = []
        self.vocabulary: Dict[str, int] = {}
        self.token_idf: Dict[str, float] = {}

    def fit(self, texts: List[str]):
        self.tokenized_corpus = [self._tokenize(t) for t in texts]
        self.bm25 = BM25Okapi(self.tokenized_corpus)
        all_tokens = sorted({t for doc in self.tokenized_corpus for t in doc})
        self.vocabulary = {t: i for i, t in enumerate(all_tokens)}
        N = len(self.tokenized_corpus)
        for token in self.vocabulary:
            df = sum(1 for doc in self.tokenized_corpus if token in doc)
            self.token_idf[token] = np.log((N - df + 0.5) / (df + 0.5) + 1)
        logger.info(f"  ✓ BM25 vocabulary: {len(self.vocabulary)} tokens")

    def _tokenize(self, text: str) -> List[str]:
        tokens = [t.lower() for t in word_tokenize(text) if t.isalnum()]
        # Remove stop words for better BM25 quality.
        # If the entire token list consists of stop words (rare edge case),
        # return an empty list rather than keeping noisy stop-word tokens.
        filtered = [t for t in tokens if t not in _STOP_WORDS]
        return filtered

    def get_sparse_vector(self, text: str) -> SparseVector:
        tokens = self._tokenize(text)
        total = len(tokens)
        counts: Dict[str, int] = {}
        for t in tokens:
            if t in self.vocabulary:
                counts[t] = counts.get(t, 0) + 1
        indices, values = [], []
        for t, cnt in counts.items():
            tf = cnt / total if total else 0.0
            indices.append(self.vocabulary[t])
            values.append(float(tf * self.token_idf.get(t, 1.0)))
        return SparseVector(indices=indices, values=values)


# ═══════════════════════════════════════════════════════════════════════════
# Collection bootstrap helpers
# ═══════════════════════════════════════════════════════════════════════════

def _ensure_collection(
    client: QdrantClient,
    name: str,
    embedding_dim: int,
    with_sparse: bool = True,
) -> None:
    existing = {c.name for c in client.get_collections().collections}
    if name not in existing:
        vectors_cfg: Dict = {
            "dense": VectorParams(size=embedding_dim, distance=Distance.COSINE)
        }
        sparse_cfg = (
            {"bm25": SparseVectorParams(index=SparseIndexParams(on_disk=False))}
            if with_sparse
            else {}
        )
        client.create_collection(
            collection_name=name,
            vectors_config=vectors_cfg,
            sparse_vectors_config=sparse_cfg,
        )
        logger.info(f"✓ Created collection: {name}")
    else:
        # Dimension-mismatch guard
        col_info = client.get_collection(name)
        try:
            vcfg = col_info.config.params.vectors
            dense_cfg = vcfg.get("dense") if hasattr(vcfg, "get") else None
            if dense_cfg and dense_cfg.size != embedding_dim:
                raise RuntimeError(
                    f"Dimension mismatch in '{name}': "
                    f"collection={dense_cfg.size}, embedder={embedding_dim}"
                )
        except AttributeError:
            pass
        logger.info(f"✓ Collection exists: {name}")


# ═══════════════════════════════════════════════════════════════════════════
# Main pipeline
# ═══════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Advanced RAG Ingestion System V3"
    )
    parser.add_argument("--data-dir",       default=DATA_DIR)
    parser.add_argument("--collection",     default=COLLECTION)
    parser.add_argument("--qdrant-url",     default=QDRANT_URL)
    parser.add_argument("--ollama-url",     default=OLLAMA_URL)
    parser.add_argument("--ollama-model",   default=OLLAMA_MODEL)
    parser.add_argument("--context-mode",   default=CONTEXT_MODE,
                        choices=["template", "llm", "none"])
    parser.add_argument("--context-llm",    default=CONTEXT_LLM_MODEL or "")
    parser.add_argument("--bm25-output",    default=BM25_OUTPUT)
    args = parser.parse_args()

    data_dir       = args.data_dir
    base_collection = args.collection
    children_col   = f"{base_collection}_children"
    parents_col    = f"{base_collection}_parents"
    qdrant_url     = args.qdrant_url
    bm25_output    = args.bm25_output

    logger.info("=" * 80)
    logger.info("ADVANCED RAG INGESTION SYSTEM V3")
    logger.info(f"  Parent collection : {parents_col}")
    logger.info(f"  Children collection: {children_col}")
    logger.info(f"  Semantic chunking : {ENABLE_SEMANTIC_CHUNKING}")
    logger.info(f"  Context enrichment: {ENABLE_CONTEXT_ENRICHMENT} ({args.context_mode})")
    logger.info(f"  MinHash dedup     : {ENABLE_MINHASH_DEDUP}")
    logger.info(f"  OCR fallback      : {ENABLE_OCR_FALLBACK} "
                f"({'available' if OCR_AVAILABLE else 'NOT installed'})")
    logger.info("=" * 80)

    # ── Qdrant client ────────────────────────────────────────────────────
    client = QdrantClient(url=qdrant_url)

    # ── Embedding model ──────────────────────────────────────────────────
    if USE_OLLAMA_BGE_M3:
        try:
            logger.info(f"Using Ollama BGE-M3 at {args.ollama_url}")
            embedder = OllamaBGEM3Embedder(args.ollama_url, args.ollama_model)
            embedding_dim = embedder.dimension
        except Exception:
            logger.warning(f"Ollama unavailable → fallback to {FALLBACK_MODEL}")
            embedder = SentenceTransformer(FALLBACK_MODEL)
            embedding_dim = 384
    else:
        embedder = SentenceTransformer(FALLBACK_MODEL)
        embedding_dim = 384
    logger.info(f"Embedding dimension: {embedding_dim}")

    # ── Collections ──────────────────────────────────────────────────────
    _ensure_collection(client, children_col, embedding_dim, with_sparse=True)
    _ensure_collection(client, parents_col,  embedding_dim, with_sparse=False)

    # ── Pipeline components ──────────────────────────────────────────────
    chunker     = SemanticChunker(CHILD_CHUNK_SIZE, CHILD_CHUNK_OVERLAP)
    pc_builder  = ParentChildBuilder(chunker)
    enricher    = ContextualEnricher(
        mode=args.context_mode,
        ollama_url=args.ollama_url,
        llm_model=args.context_llm,
    )
    deduplicator = MinHashDeduplicator() if ENABLE_MINHASH_DEDUP else None
    bm25_index   = BM25Index()

    # ── Accumulators ─────────────────────────────────────────────────────
    total_files    = 0
    total_children = 0
    total_parents  = 0
    all_child_points: List[Tuple[Dict, str]] = []   # (point_dict, raw_text)
    all_parent_points: List[Tuple[Dict, str]] = []

    logger.info("\nPhase 1: Loading and chunking documents...")

    for root, _, files in os.walk(data_dir):
        for file in files:
            if not file.lower().endswith((".pdf", ".docx", ".txt")):
                continue

            path = os.path.join(root, file)
            logger.info(f"\n→ Processing: {file}")

            try:
                h = file_hash(path)

                # Skip already-indexed files (check children collection)
                if already_indexed(client, children_col, h):
                    logger.info("  ⊘ Already indexed — skipping")
                    continue

                # Load document
                text, doc_metadata = AdvancedDocumentLoader.load(path)
                if not text or len(text.strip()) < MIN_CHUNK_SIZE:
                    logger.info("  ⊘ Empty or too short after extraction")
                    continue

                # Detect sections
                sections = chunker.detect_sections(text) if ENABLE_SECTION_AWARE else [
                    DocumentSection("Document", text, 0)
                ]
                logger.info(f"  ✓ {len(sections)} sections detected")

                # Build parent-child hierarchy
                parents, children = pc_builder.build(sections, file_hash_value=h)
                logger.info(
                    f"  ✓ {len(parents)} parent chunks → {len(children)} child chunks"
                )

                if not children:
                    continue

                # Contextual enrichment
                if ENABLE_CONTEXT_ENRICHMENT:
                    enricher.enrich(children, file)
                    logger.info(f"  ✓ Context enrichment applied ({args.context_mode} mode)")

                # MinHash deduplication (global, cross-file)
                accepted_children: List[ChildChunk] = []
                accepted_parents: Set[str] = set()
                dedup_removed = 0

                for child in children:
                    if deduplicator is not None:
                        if deduplicator.is_duplicate(child.text):
                            dedup_removed += 1
                            continue
                    accepted_children.append(child)
                    accepted_parents.add(child.parent_id)

                if dedup_removed:
                    logger.info(f"  ⊘ Removed {dedup_removed} near-duplicate child chunks")

                # Filter parents: only keep those that have accepted children
                accepted_parent_objs = [
                    p for p in parents if p.parent_id in accepted_parents
                ]

                logger.info(
                    f"  ✓ Accepted: {len(accepted_children)} child chunks, "
                    f"{len(accepted_parent_objs)} parent chunks"
                )

                if not accepted_children:
                    continue

                # ── Embed children (use enriched text for embedding) ─────
                child_texts_embed = [c.enriched_text for c in accepted_children]
                logger.info(
                    f"  ⚡ Embedding {len(child_texts_embed)} child chunks..."
                )
                if isinstance(embedder, OllamaBGEM3Embedder):
                    child_embeddings = embedder.encode(child_texts_embed, batch_size=8)
                else:
                    raw = embedder.encode(
                        child_texts_embed,
                        batch_size=32,
                        show_progress_bar=False,
                        convert_to_numpy=True,
                    )
                    child_embeddings = [e.tolist() for e in raw]

                # ── Embed parents ─────────────────────────────────────────
                parent_texts = [p.text for p in accepted_parent_objs]
                logger.info(
                    f"  ⚡ Embedding {len(parent_texts)} parent chunks..."
                )
                if isinstance(embedder, OllamaBGEM3Embedder):
                    parent_embeddings = embedder.encode(parent_texts, batch_size=8)
                else:
                    raw = embedder.encode(
                        parent_texts,
                        batch_size=32,
                        show_progress_bar=False,
                        convert_to_numpy=True,
                    )
                    parent_embeddings = [e.tolist() for e in raw]

                # ── Build child point records ─────────────────────────────
                folder = os.path.relpath(root, data_dir)
                file_type = Path(file).suffix.lower()

                for i, (child, emb) in enumerate(
                        zip(accepted_children, child_embeddings)):
                    if emb is None:
                        logger.warning(f"  ⚠ Skipping child {i} — embedding failed")
                        continue
                    cid = str(uuid.uuid5(
                        uuid.NAMESPACE_DNS, f"{h}_child_{i}"
                    ))
                    point: Dict = {
                        "id": cid,
                        "vector": emb,
                        "payload": {
                            "content": child.text,            # raw text for display
                            "enriched_content": child.enriched_text,  # embedded text
                            "parent_id": child.parent_id,
                            "child_index": child.child_index,
                            "source_path": path,
                            "filename": file,
                            "folder": folder,
                            "file_type": file_type,
                            "file_hash": h,
                            "chunk_type": child.chunk_type,
                            "section_title": child.section_title,
                            "section_hierarchy": child.section_hierarchy,
                            "word_count": child.word_count,
                            "sentence_count": child.sentence_count,
                            "start_char": child.start_char,
                            "end_char": child.end_char,
                            **doc_metadata,
                        },
                    }
                    all_child_points.append((point, child.text))
                    total_children += 1

                # ── Build parent point records ────────────────────────────
                pid_to_emb = {
                    p.parent_id: emb
                    for p, emb in zip(accepted_parent_objs, parent_embeddings)
                }
                for parent in accepted_parent_objs:
                    emb = pid_to_emb.get(parent.parent_id)
                    if emb is None:
                        continue
                    point = {
                        "id": parent.parent_id,
                        "vector": emb,
                        "payload": {
                            "content": parent.text,
                            "source_path": path,
                            "filename": file,
                            "folder": folder,
                            "file_type": file_type,
                            "file_hash": h,
                            "chunk_type": parent.chunk_type,
                            "section_title": parent.section_title,
                            "section_hierarchy": parent.section_hierarchy,
                            "word_count": parent.word_count,
                            "start_char": parent.start_char,
                            "end_char": parent.end_char,
                            **doc_metadata,
                        },
                    }
                    all_parent_points.append((point, parent.text))
                    total_parents += 1

                total_files += 1

            except Exception as e:
                logger.error(f"  ✗ Error processing {file}: {e}", exc_info=True)
                continue

    if not all_child_points:
        logger.warning("No documents to index!")
        return

    # ── Phase 2: BM25 on child chunks ─────────────────────────────────────
    all_child_texts = [text for _, text in all_child_points]
    logger.info(
        f"\nPhase 2: Building BM25 index for {len(all_child_texts)} child chunks..."
    )
    bm25_index.fit(all_child_texts)
    bm25_data = {
        "vocabulary": bm25_index.vocabulary,
        "token_idf": {k: float(v) for k, v in bm25_index.token_idf.items()},
    }
    Path(bm25_output).parent.mkdir(parents=True, exist_ok=True)
    with open(bm25_output, "w") as f:
        json.dump(bm25_data, f)
    logger.info(f"  ✓ Saved BM25 index → {bm25_output}")

    # ── Phase 3: Upload children to Qdrant ───────────────────────────────
    logger.info(
        f"\nPhase 3: Uploading {len(all_child_points)} child chunks "
        f"to '{children_col}'..."
    )
    child_structs: List[PointStruct] = []
    for pd, chunk_text in all_child_points:
        sv = bm25_index.get_sparse_vector(chunk_text)
        child_structs.append(PointStruct(
            id=pd["id"],
            vector={"dense": pd["vector"], "bm25": sv},
            payload=pd["payload"],
        ))

    batch_size = 100
    for i in range(0, len(child_structs), batch_size):
        client.upsert(
            collection_name=children_col,
            points=child_structs[i:i + batch_size],
            wait=True,
        )
        logger.info(
            f"  ✓ Children batch "
            f"{i // batch_size + 1}/{math.ceil(len(child_structs) / batch_size)}"
        )

    # ── Phase 4: Upload parents to Qdrant ────────────────────────────────
    logger.info(
        f"\nPhase 4: Uploading {len(all_parent_points)} parent chunks "
        f"to '{parents_col}'..."
    )
    parent_structs: List[PointStruct] = []
    for pd, _ in all_parent_points:
        parent_structs.append(PointStruct(
            id=pd["id"],
            vector={"dense": pd["vector"]},
            payload=pd["payload"],
        ))

    for i in range(0, len(parent_structs), batch_size):
        client.upsert(
            collection_name=parents_col,
            points=parent_structs[i:i + batch_size],
            wait=True,
        )
        logger.info(
            f"  ✓ Parents batch "
            f"{i // batch_size + 1}/{math.ceil(len(parent_structs) / batch_size)}"
        )

    # ── Summary ──────────────────────────────────────────────────────────
    logger.info("\n" + "=" * 80)
    logger.info("INGESTION COMPLETE")
    logger.info("=" * 80)
    logger.info(f"✓ Files processed       : {total_files}")
    logger.info(f"✓ Parent chunks indexed : {total_parents}")
    logger.info(f"✓ Child chunks indexed  : {total_children}")
    if total_files > 0:
        logger.info(
            f"✓ Avg children/file     : {total_children / total_files:.1f}"
        )
    logger.info(f"✓ Children collection   : {children_col}")
    logger.info(f"✓ Parents collection    : {parents_col}")
    logger.info(f"✓ Embedding dimension   : {embedding_dim}")
    logger.info(f"✓ BM25 vocabulary       : {len(bm25_index.vocabulary)}")
    if deduplicator is not None:
        logger.info(f"✓ MinHash index size    : {len(deduplicator)}")
    logger.info("=" * 80)


if __name__ == "__main__":
    main()
