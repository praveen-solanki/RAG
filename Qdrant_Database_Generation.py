##
##File
## ?
##FileTypeDetector
## ?
##LoaderFactory
## ?
##[ PDF | DOCX | TXT | HTML | CSV | MD ] Loader
## ?
##Text (string)
## ?
##Chunk ? Embed ? Store (unchanged)
##
#
#
#import os
#import hashlib
#import uuid
#
#import pdfplumber
#import docx
#from sentence_transformers import SentenceTransformer
#from qdrant_client import QdrantClient
#from qdrant_client.models import (
#    VectorParams,
#    Distance,
#    Filter,
#    FieldCondition,
#    MatchValue,
#)
#
## ================= CONFIG =================
#
#DATA_DIR = "DATASET"
#COLLECTION = "rag_database_384_10"
#QDRANT_URL = "http://localhost:7333"
#EMBEDDING_MODEL = "all-MiniLM-L6-v2"
#
#CHUNK_SIZE = 400
#CHUNK_OVERLAP = 60
#
#assert CHUNK_OVERLAP < CHUNK_SIZE, "CHUNK_OVERLAP must be smaller than CHUNK_SIZE"
#
## =========================================
#
## ---------- Qdrant client ----------
#client = QdrantClient(url=QDRANT_URL)
#
## ---------- Create collection (safe to re-run) ----------
#existing = [c.name for c in client.get_collections().collections]
#if COLLECTION not in existing:
#    client.create_collection(
#        collection_name=COLLECTION,
#        vectors_config=VectorParams(  # CHANGED: No dictionary wrapper here
#            size=384,
#            distance=Distance.COSINE,
#        ),
#    )
#
#
#
#    print(f"? Created collection: {COLLECTION}")
#else:
#    print(f"?? Collection already exists: {COLLECTION}")
#
## ---------- Embedding model ----------
#model = SentenceTransformer(EMBEDDING_MODEL)
#
## ---------- Utilities ----------
#def file_hash(path: str) -> str:
#    with open(path, "rb") as f:
#        return hashlib.sha256(f.read()).hexdigest()
#
#
#def already_indexed(file_hash_value: str) -> bool:
#    filt = Filter(
#        must=[
#            FieldCondition(
#                key="file_hash",
#                match=MatchValue(value=file_hash_value),
#            )
#        ]
#    )
#
#    points, _ = client.scroll(
#        collection_name=COLLECTION,
#        scroll_filter=filt,
#        limit=1,
#    )
#    return len(points) > 0
#
#
#def chunk_text(text: str, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
#    chunks = []
#    start = 0
#    length = len(text)
#
#    while start < length:
#        end = min(start + size, length)
#        chunks.append(text[start:end])
#        start += size - overlap
#
#    return chunks
#
#
## ---------- File loaders ----------
#def load_text(path: str):
#    if path.lower().endswith(".pdf"):
#        text = ""
#        with pdfplumber.open(path) as pdf:
#            for page in pdf.pages:
#                page_text = page.extract_text()
#                if page_text:
#                    text += page_text + "\n"
#        return text
#
#    if path.lower().endswith(".docx"):
#        d = docx.Document(path)
#        return "\n".join(p.text for p in d.paragraphs)
#
#    if path.lower().endswith(".txt"):
#        with open(path, "r", encoding="utf-8", errors="ignore") as f:
#            return f.read()
#
#    return None
#
#
## ---------- Ingestion ----------
#total_chunks = 0
#
#for root, _, files in os.walk(DATA_DIR):
#    for file in files:
#        if not file.lower().endswith((".pdf", ".docx", ".txt")):
#            continue
#
#        path = os.path.join(root, file)
#        print(f"\n?? Processing: {path}")
#
#        h = file_hash(path)
#
#        if already_indexed(h):
#            print("??  Skipped (already indexed)")
#            continue
#
#        text = load_text(path)
#        if not text or len(text.strip()) < 50:
#            print("??  Skipped (empty or invalid content)")
#            continue
#
#        chunks = chunk_text(text)
#        embeddings = model.encode(
#            chunks,
#            batch_size=32,
#            show_progress_bar=False,
#        )
#
#        points = []
#        for i, vector in enumerate(embeddings):
#            point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{h}_{i}"))
#
#            points.append(
#                {
#                    "id": point_id,
#                    "vector": vector.tolist(), # CHANGED: No {"embedding": ...} wrapper
#                    "payload": {
#                        "source_path": path,
#                        "folder": os.path.relpath(root, DATA_DIR),
#                        "file_type": file.split(".")[-1],
#                        "file_hash": h,
#                        "chunk_id": i,
#                        "content": chunks[i], # CHANGED: Renamed "text" -> "content"
#                    },
#                }
#            )
#
#
#        client.upsert(
#            collection_name=COLLECTION,
#            points=points,
#            wait=True,
#        )
#
#        total_chunks += len(points)
#        print(f"? Indexed {len(points)} chunks")
#
#print("\n?? Ingestion complete.")
#print(f"?? Total chunks indexed: {total_chunks}")
#print("?? Qdrant is ready.")
#
#
#
#




"""
IMPROVED INGESTION SCRIPT
Fixes and enhancements to your original code
"""

import os
import hashlib
import uuid
from typing import Optional, List

import pdfplumber
import docx
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import (
    VectorParams,
    Distance,
    Filter,
    FieldCondition,
    MatchValue,
    PointStruct,
)

# ================= CONFIG =================

DATA_DIR = "DATASET"
COLLECTION = "rag_database_384_10"
QDRANT_URL = "http://localhost:7333"
EMBEDDING_MODEL = "all-MiniLM-L6-v2"

CHUNK_SIZE = 400
CHUNK_OVERLAP = 60

assert CHUNK_OVERLAP < CHUNK_SIZE, "CHUNK_OVERLAP must be smaller than CHUNK_SIZE"

# =========================================

# ---------- Qdrant client ----------
client = QdrantClient(url=QDRANT_URL)

# ---------- Create collection (safe to re-run) ----------
existing = [c.name for c in client.get_collections().collections]
if COLLECTION not in existing:
    client.create_collection(
        collection_name=COLLECTION,
        vectors_config=VectorParams(
            size=384,
            distance=Distance.COSINE,
        ),
    )
    print(f"? Created collection: {COLLECTION}")
else:
    print(f"? Collection already exists: {COLLECTION}")

# ---------- Embedding model ----------
print(f"Loading embedding model: {EMBEDDING_MODEL}...")
model = SentenceTransformer(EMBEDDING_MODEL)
print("? Model loaded")

# ---------- Utilities ----------
def file_hash(path: str) -> str:
    """Generate SHA256 hash of file content"""
    with open(path, "rb") as f:
        return hashlib.sha256(f.read()).hexdigest()


def already_indexed(file_hash_value: str) -> bool:
    """Check if file with this hash is already indexed"""
    filt = Filter(
        must=[
            FieldCondition(
                key="file_hash",
                match=MatchValue(value=file_hash_value),
            )
        ]
    )

    points, _ = client.scroll(
        collection_name=COLLECTION,
        scroll_filter=filt,
        limit=1,
    )
    return len(points) > 0


def chunk_text(text: str, size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
    """
    Split text into overlapping chunks
    IMPROVEMENT: Better handling of empty text and edge cases
    """
    if not text or len(text.strip()) < size:
        return [text.strip()] if text.strip() else []
    
    chunks = []
    start = 0
    length = len(text)

    while start < length:
        end = min(start + size, length)
        chunk = text[start:end].strip()  # ADDED: Strip whitespace
        if chunk:  # ADDED: Only add non-empty chunks
            chunks.append(chunk)
        start += size - overlap
        
        # ADDED: Prevent infinite loop if overlap >= size
        if start == 0:
            start = size

    return chunks


# ---------- File loaders ----------
def load_text(path: str) -> Optional[str]:
    """
    Load text from various file formats
    IMPROVEMENT: Better error handling and table extraction for PDFs
    """
    try:
        if path.lower().endswith(".pdf"):
            text = ""
            with pdfplumber.open(path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text += f"[Page {page_num}]\n{page_text}\n\n"
                    
                    # ADDED: Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table_idx, table in enumerate(tables, start=1):
                            if table:
                                text += f"[Table {table_idx} on Page {page_num}]\n"
                                for row in table:
                                    text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                                text += "\n"
            return text

        if path.lower().endswith(".docx"):
            doc = docx.Document(path)
            # IMPROVEMENT: Better text extraction including tables
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables from docx
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)
            
            return "\n".join(text_parts)

        if path.lower().endswith(".txt"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

    except Exception as e:
        print(f"  ?? Error loading {path}: {e}")
        return None

    return None


# ---------- Ingestion ----------
total_chunks = 0
total_files = 0
skipped_files = 0
error_files = 0

print("\n" + "="*80)
print("STARTING INGESTION")
print("="*80)

for root, _, files in os.walk(DATA_DIR):
    for file in files:
        if not file.lower().endswith((".pdf", ".docx", ".txt")):
            continue

        path = os.path.join(root, file)
        print(f"\n?? Processing: {file}")

        try:
            h = file_hash(path)

            if already_indexed(h):
                print("  ??  Skipped (already indexed)")
                skipped_files += 1
                continue

            text = load_text(path)
            if not text or len(text.strip()) < 50:
                print("  ??  Skipped (empty or invalid content)")
                skipped_files += 1
                continue

            chunks = chunk_text(text)
            if not chunks:
                print("  ??  Skipped (no chunks generated)")
                skipped_files += 1
                continue

            # IMPROVEMENT: Batch embedding with progress indicator
            print(f"  ? Generating embeddings for {len(chunks)} chunks...")
            embeddings = model.encode(
                chunks,
                batch_size=32,
                show_progress_bar=False,
                convert_to_numpy=True,
            )

            # IMPROVEMENT: Use PointStruct for better type safety
            points = []
            for i, (chunk, vector) in enumerate(zip(chunks, embeddings)):
                point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{h}_{i}"))

                point = PointStruct(
                    id=point_id,
                    vector=vector.tolist(),
                    payload={
                        "source_path": path,
                        "filename": file,  # ADDED: Just filename for easier filtering
                        "folder": os.path.relpath(root, DATA_DIR),
                        "file_type": file.split(".")[-1],
                        "file_hash": h,
                        "chunk_id": i,
                        "total_chunks": len(chunks),  # ADDED: Useful metadata
                        "content": chunk,
                        "chunk_length": len(chunk),  # ADDED: For analysis
                    },
                )
                points.append(point)

            # IMPROVEMENT: Batch upsert with error handling
            print(f"  ? Uploading to Qdrant...")
            client.upsert(
                collection_name=COLLECTION,
                points=points,
                wait=True,
            )

            total_chunks += len(points)
            total_files += 1
            print(f"  ? Indexed {len(points)} chunks")

        except Exception as e:
            print(f"  ? Error processing {file}: {e}")
            error_files += 1
            continue

print("\n" + "="*80)
print("INGESTION COMPLETE")
print("="*80)
print(f"? Files processed: {total_files}")
print(f"??  Files skipped: {skipped_files}")
print(f"? Files with errors: {error_files}")
print(f"?? Total chunks indexed: {total_chunks}")
print(f"?? Average chunks per file: {total_chunks/total_files if total_files > 0 else 0:.1f}")
print("="*80)